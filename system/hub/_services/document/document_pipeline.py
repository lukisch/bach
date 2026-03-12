#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# SPDX-License-Identifier: MIT
"""
Copyright (c) 2026 BACH Contributors

Permission is hereby granted, free of charge, to any person obtaining a copy
of this software and associated documentation files (the "Software"), to deal
in the Software without restriction, including without limitation the rights
to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
copies of the Software, and to permit persons to whom the Software is
furnished to do so, subject to the following conditions:

The above copyright notice and this permission notice shall be included in all
copies or substantial portions of the Software.

THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
SOFTWARE.
"""

"""
DocumentPipeline — Vereinter Dokumenten-Verarbeitungs-Service
==============================================================

Ersetzt die redundante Logik aus:
  - document_collector.py  (Scan, Kategorisierung, Text-Extraktion, TextBundle)
  - file_access_hook.py    (Scan, Text-Extraktion aus Bundles, Profil-Cache)

Gemeinsamer Kern beider Dateien:
  - Ordner scannen (rglob, Dateityp-Filter)
  - Dokument-Typ erkennen (pattern-matching auf Dateinamen)
  - Text extrahieren (DOCX, PDF, TXT, XLSX, MSG, EML)
  - Anonymisierung anwenden (optional)

WICHTIG: Die asynchronen / proaktiven Teile aus file_access_hook.py
(scan_for_new_clients, AnonymizedWorkflow-Integration) bleiben in
file_access_hook.py und werden in Phase 2 migriert.

Dieses Modul stellt SYNC-Interfaces bereit.

Verwendung (Use-Case 1 — Foerderplaner, ohne Anonymisierung):
    from hub._services.document.document_pipeline import DocumentPipeline

    pipeline = DocumentPipeline()
    result = pipeline.scan_folder("/path/to/client/folder")
    bundle = pipeline.extract_bundle(result.documents)

Verwendung (Use-Case 2 — mit Anonymisierung, fuer Bundles):
    from hub._services.document.document_pipeline import DocumentPipeline
    from hub._services.document.anonymizer_service import AnonymProfile

    pipeline = DocumentPipeline()
    result = pipeline.scan_folder("/path/to/client/folder")
    bundle = pipeline.extract_bundle(result.documents, anonym_profile=profile)
    pipeline.save_bundle(bundle, "/path/to/output.txt")

Verwendung (Use-Case 3 — Bundle-Verzeichnis lesen, fuer FileAccessHook):
    pipeline = DocumentPipeline()
    content = pipeline.read_bundle_dir(Path("/path/to/bundles/K_XXX"))

Version: 1.0.0
Erstellt: 2026-02-22
Referenz: SQ078 Phase 1
Ersetzt: document_collector.py (gemeinsame Scan/Extraktion-Logik)
         file_access_hook._extract_text() (identischer Code)
"""

import re
import sys
from dataclasses import dataclass, field
from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict, List, Optional, Tuple, TYPE_CHECKING
from enum import Enum

if TYPE_CHECKING:
    from .anonymizer_service import AnonymProfile


# ═══════════════════════════════════════════════════════════════
# Datenklassen (uebernommen aus document_collector.py)
# ═══════════════════════════════════════════════════════════════

class DocumentCategory(Enum):
    """Dokument-Kategorien nach Prioritaet."""
    CORE = "core"           # Immer an LLM senden
    STUFE2 = "stufe2"       # Nach CORE senden
    EXTENDED = "extended"   # Nur auf Anfrage
    SKIP = "skip"           # Nicht verarbeiten


# Prioritaets-Reihenfolge fuer Sortierung (niedrig = hoeher priorisiert)
CATEGORY_PRIORITY = {
    DocumentCategory.CORE: 0,
    DocumentCategory.STUFE2: 1,
    DocumentCategory.EXTENDED: 2,
    DocumentCategory.SKIP: 3,
}

# Dokument-Typ Prioritaet innerhalb derselben Kategorie
# Hierarchie: Aktuelle Protokolle > Hilfeplan/Bewilligung > proAutismus-Bericht
#              > Arztberichte > Schulberichte > Mails > Rest
DOC_TYPE_PRIORITY = {
    "protokoll": 0,
    "aktendeckblatt": 1,
    "hilfeplan": 2,
    "proautismus_bericht": 3,
    "arztbericht": 4,
    "schulbericht": 5,
    "mail": 6,
    "root_dokument": 0,  # Root-Dateien sind oft das Hauptprotokoll (Verlaufsprotokoll)
    "unknown": 8,
}


@dataclass
class DocumentInfo:
    """Informationen zu einem Dokument."""
    path: Path
    filename: str
    suffix: str
    category: DocumentCategory
    doc_type: str           # z.B. "protokoll", "hilfeplan", "arztbericht"
    date_hint: Optional[datetime] = None  # Erkanntes Datum aus Dateinamen
    size_bytes: int = 0
    text_extracted: bool = False
    text_length: int = 0


@dataclass
class ScanResult:
    """Ergebnis der Dokumenten-Sammlung."""
    documents: List[DocumentInfo] = field(default_factory=list)
    core_count: int = 0
    stufe2_count: int = 0
    extended_count: int = 0
    skipped_count: int = 0
    errors: List[str] = field(default_factory=list)


# Rueckwaertskompatibilitaet: CollectorResult ist ein Alias
CollectorResult = ScanResult


@dataclass
class TextBundle:
    """Extrahiertes Text-Bundle."""
    core_text: str = ""
    stufe2_text: str = ""
    extended_text: str = ""
    total_length: int = 0
    files_processed: int = 0


# ═══════════════════════════════════════════════════════════════
# Erkennungsmuster (aus document_collector.py uebernommen)
# ═══════════════════════════════════════════════════════════════

DOC_PATTERNS = {
    # CORE Dokumente
    "protokoll": [
        r"protokoll",
        r"einzel.*protokoll",
        r"gruppen.*protokoll",
        r"doku(?:mentation)?",
        r"verlauf",
        r"sitzung",
    ],
    "aktendeckblatt": [
        r"aktendeckblatt",
        r"aktendatenblatt",
        r"anmeld(?:ung|ebogen|eunterlagen)",
        r"stammdaten",
    ],
    "hilfeplan": [
        r"hilfeplan",
        r"kostenzusage",
        r"bewilligung",
        r"bescheid",
    ],
    "proautismus_bericht": [
        r"entwicklungsbericht",
        r"foerderbericht",
        r"icf.?bericht",
        r"bericht.*proautismus",
    ],

    # STUFE2 Dokumente
    "arztbericht": [
        r"arztbericht",
        r"bericht.*dr\.?",
        r"entlass(?:ungs)?bericht",
        r"diagnose",
        r"befund",
        r"gutachten",
    ],
    "schulbericht": [
        r"schulbericht",
        r"zeugnis",
        r"lernentwicklung",
        r"foerderplan.*schule",
        r"nachteilsausgleich",
    ],
    "mail": [
        r"\.msg$",
        r"\.eml$",
    ],
}

FOLDER_HINTS = {
    "core": [
        "dokumentation", "protokolle", "anmeldeunterlagen",
        "hilfeplan", "amt", "intern"
    ],
    "stufe2": [
        "familie", "umfeld", "extern", "arzt", "schule", "mail"
    ],
    "extended": [
        "archiv", "alt", "historie", "backup"
    ],
    "skip": [
        "output", "_prepare", "_ready", "_archive", "temp", "_bundle"
    ],
}

# Unterstuetzte Dateitypen (Scan)
SUPPORTED_EXTENSIONS = {".docx", ".doc", ".txt", ".md", ".pdf", ".xlsx", ".xls", ".msg", ".eml"}

# Dateitypen die der Bundle-Lesemodus unterstuetzt (FileAccessHook-Kompatibilitaet)
BUNDLE_READ_EXTENSIONS = {".txt", ".md", ".docx", ".pdf", ".xlsx", ".xls"}


# ═══════════════════════════════════════════════════════════════
# DocumentPipeline — Haupt-Klasse
# ═══════════════════════════════════════════════════════════════

class DocumentPipeline:
    """
    Vereinter Dokumenten-Verarbeitungs-Service.

    Kombiniert die gemeinsame Logik aus:
      - DocumentCollector (document_collector.py)
      - FileAccessHook._extract_text (file_access_hook.py)

    Die Klasse bietet drei Interface-Schichten:
      1. Scan-Interface: scan_folder() -> ScanResult
      2. Extraktion-Interface: extract_bundle() -> TextBundle
      3. Bundle-Lese-Interface: read_bundle_dir() -> str (fuer FileAccessHook-Compat.)

    Sicherheitshinweis: Wenn anonym_profile uebergeben wird, werden alle
    echten Namen/Daten im Ausgabe-Text ersetzt. Die Originaldateien werden
    niemals veraendert.
    """

    def __init__(
        self,
        berichtszeitraum_monate: int = 12,
        stichtag: Optional[datetime] = None
    ):
        """
        Args:
            berichtszeitraum_monate: Zeitraum fuer aktuelle Dokumente (Default: 12)
            stichtag: Referenzdatum (Default: heute)
        """
        self.berichtszeitraum_monate = berichtszeitraum_monate
        self.stichtag = stichtag or datetime.now()
        self.zeitraum_start = self.stichtag - timedelta(days=berichtszeitraum_monate * 30)

        # Compile regex patterns
        self._compiled_patterns: Dict[str, List[re.Pattern]] = {}
        for doc_type, patterns in DOC_PATTERNS.items():
            self._compiled_patterns[doc_type] = [
                re.compile(p, re.IGNORECASE) for p in patterns
            ]

    # ─────────────────────────────────────────────────────────────
    # Interface 1: Ordner scannen und kategorisieren
    # ─────────────────────────────────────────────────────────────

    def scan_folder(self, folder: str, recursive: bool = True) -> ScanResult:
        """
        Scannt einen Ordner und kategorisiert alle Dokumente.

        Args:
            folder: Pfad zum Klienten-Ordner
            recursive: Unterordner einbeziehen (Default: True)

        Returns:
            ScanResult mit kategorisierten Dokumenten
        """
        result = ScanResult()
        folder_path = Path(folder)

        if not folder_path.exists():
            result.errors.append(f"Ordner nicht gefunden: {folder}")
            return result

        # Dateien sammeln
        if recursive:
            files = list(folder_path.rglob("*"))
        else:
            files = list(folder_path.iterdir())

        for filepath in sorted(files):
            if not filepath.is_file():
                continue
            if filepath.name.startswith("."):
                continue
            if filepath.suffix.lower() not in SUPPORTED_EXTENSIONS:
                continue

            doc_info = self._analyze_document(filepath, folder_path)
            result.documents.append(doc_info)

            if doc_info.category == DocumentCategory.CORE:
                result.core_count += 1
            elif doc_info.category == DocumentCategory.STUFE2:
                result.stufe2_count += 1
            elif doc_info.category == DocumentCategory.EXTENDED:
                result.extended_count += 1
            else:
                result.skipped_count += 1

        # Sortierung: Kategorie-Prioritaet, dann Typ-Prioritaet,
        # dann Datum absteigend (neueste Dokumente zuerst)
        ten_years_ago = self.stichtag - timedelta(days=10 * 365)

        def _sort_key(d: DocumentInfo):
            cat_prio = CATEGORY_PRIORITY.get(d.category, 9)
            type_prio = DOC_TYPE_PRIORITY.get(d.doc_type, 8)
            # Arztberichte aelter als 10 Jahre -> niedrigste Prioritaet
            if d.doc_type == "arztbericht" and d.date_hint and d.date_hint < ten_years_ago:
                type_prio = 99
            # Datum absteigend: neueste zuerst
            # Dokumente ohne Datum bekommen mittlere Prioritaet
            if d.date_hint:
                date_val = -d.date_hint.timestamp()
            else:
                date_val = 0  # Mitte
            return (cat_prio, type_prio, date_val)

        result.documents.sort(key=_sort_key)

        return result

    def _analyze_document(self, filepath: Path, base_folder: Path) -> DocumentInfo:
        """Analysiert ein einzelnes Dokument und bestimmt Kategorie."""
        filename = filepath.name
        suffix = filepath.suffix.lower()

        try:
            rel_path = filepath.relative_to(base_folder)
            folder_parts = [p.lower() for p in rel_path.parts[:-1]]
        except ValueError:
            folder_parts = []

        doc_info = DocumentInfo(
            path=filepath,
            filename=filename,
            suffix=suffix,
            category=DocumentCategory.EXTENDED,  # Default
            doc_type="unknown",
            size_bytes=filepath.stat().st_size if filepath.exists() else 0
        )

        doc_info.date_hint = self._extract_date(filename)

        # Ordner-basierte Skip-Erkennung
        for skip_folder in FOLDER_HINTS["skip"]:
            if any(skip_folder in part for part in folder_parts):
                doc_info.category = DocumentCategory.SKIP
                return doc_info

        doc_type = self._detect_doc_type(filename, suffix)
        doc_info.doc_type = doc_type

        # Dateien im Root-Ordner sind CORE (ausser _-Dateien)
        is_root = len(folder_parts) == 0
        if is_root and not filename.startswith("_"):
            doc_info.category = DocumentCategory.CORE
            if doc_type == "unknown":
                # Root .docx/.doc ohne Pattern-Match = Verlaufsprotokoll (Klientenname.docx)
                if suffix in [".docx", ".doc"]:
                    doc_info.doc_type = "protokoll"
                else:
                    doc_info.doc_type = "root_dokument"
            return doc_info

        # Kategorie basierend auf Typ und Datum
        if doc_type in ["protokoll", "aktendeckblatt"]:
            if self._is_in_berichtszeitraum(doc_info.date_hint):
                doc_info.category = DocumentCategory.CORE
            else:
                doc_info.category = DocumentCategory.EXTENDED

        elif doc_type == "hilfeplan":
            doc_info.category = DocumentCategory.CORE

        elif doc_type == "proautismus_bericht":
            doc_info.category = DocumentCategory.CORE

        elif doc_type in ["arztbericht", "schulbericht"]:
            if self._is_recent(doc_info.date_hint, years=2):
                doc_info.category = DocumentCategory.STUFE2
            else:
                doc_info.category = DocumentCategory.EXTENDED

        elif doc_type == "mail":
            if self._is_in_berichtszeitraum(doc_info.date_hint):
                doc_info.category = DocumentCategory.STUFE2
            else:
                doc_info.category = DocumentCategory.EXTENDED

        # Ordner-Hints koennen Kategorie ueberschreiben
        for part in folder_parts:
            if any(hint in part for hint in FOLDER_HINTS["core"]):
                if doc_info.category != DocumentCategory.SKIP:
                    if self._is_in_berichtszeitraum(doc_info.date_hint):
                        doc_info.category = DocumentCategory.CORE
            elif any(hint in part for hint in FOLDER_HINTS["extended"]):
                if doc_info.category not in [DocumentCategory.SKIP, DocumentCategory.CORE]:
                    doc_info.category = DocumentCategory.EXTENDED

        return doc_info

    def _detect_doc_type(self, filename: str, suffix: str) -> str:
        """Erkennt den Dokumenttyp anhand des Dateinamens."""
        if suffix in [".msg", ".eml"]:
            return "mail"

        for doc_type, patterns in self._compiled_patterns.items():
            for pattern in patterns:
                if pattern.search(filename):
                    return doc_type

        return "unknown"

    def _extract_date(self, filename: str) -> Optional[datetime]:
        """Versucht ein Datum aus dem Dateinamen zu extrahieren."""
        patterns = [
            r"\b(\d{4})-(\d{2})-(\d{2})\b",     # 2024-01-15
            r"\b(\d{2})\.(\d{2})\.(\d{4})\b",   # 15.01.2024
            r"\b(\d{4})-(\d{2})\b",             # 2024-01
            r"\b(\d{2})-(\d{4})\b",             # 01-2024
            r"_(\d{4})(?:[_\.\s]|$)",           # _2024_
        ]

        for pattern in patterns:
            match = re.search(pattern, filename)
            if match:
                groups = match.groups()
                try:
                    if len(groups) == 3:
                        if len(groups[0]) == 4:
                            return datetime(int(groups[0]), int(groups[1]), int(groups[2]))
                        else:
                            return datetime(int(groups[2]), int(groups[1]), int(groups[0]))
                    elif len(groups) == 2:
                        if len(groups[0]) == 4:
                            return datetime(int(groups[0]), int(groups[1]), 1)
                        else:
                            return datetime(int(groups[1]), int(groups[0]), 1)
                    elif len(groups) == 1:
                        return datetime(int(groups[0]), 6, 15)
                except (ValueError, IndexError):
                    pass

        return None

    def _is_in_berichtszeitraum(self, date: Optional[datetime]) -> bool:
        """Prueft ob ein Datum im aktuellen Berichtszeitraum liegt."""
        if date is None:
            return True
        return date >= self.zeitraum_start

    def _is_recent(self, date: Optional[datetime], years: int = 2) -> bool:
        """Prueft ob ein Datum innerhalb der letzten X Jahre liegt."""
        if date is None:
            return True
        cutoff = self.stichtag - timedelta(days=years * 365)
        return date >= cutoff

    # ─────────────────────────────────────────────────────────────
    # Interface 2: Text extrahieren und Bundle bauen
    # ─────────────────────────────────────────────────────────────

    def extract_bundle(
        self,
        documents: List[DocumentInfo],
        include_extended: bool = False,
        anonym_profile: Optional["AnonymProfile"] = None
    ) -> TextBundle:
        """
        Extrahiert Text aus den kategorisierten Dokumenten.

        Args:
            documents: Liste von DocumentInfo (aus scan_folder)
            include_extended: Auch EXTENDED Dokumente einbeziehen
            anonym_profile: Optional - AnonymProfile fuer sofortige Anonymisierung

        Returns:
            TextBundle mit extrahiertem (und optional anonymisiertem) Text
        """
        bundle = TextBundle()

        # Anonymisierungs-Mappings vorbereiten
        if anonym_profile:
            all_replacements = {}
            for category in anonym_profile.mappings.values():
                all_replacements.update(category)
            sorted_replacements = sorted(
                all_replacements.items(),
                key=lambda x: len(x[0]),
                reverse=True
            )
        else:
            sorted_replacements = []

        def anonymize_text(text: str) -> str:
            if not sorted_replacements:
                return text
            for old, new in sorted_replacements:
                text = text.replace(old, new)
            return text

        # Prioritaets-Gruppen innerhalb der Kategorien
        high_prio_parts = []    # Aktuelle Protokolle, Hilfeplan, Bewilligung, Aktendeckblatt
        medium_prio_parts = []  # proAutismus-Berichte, Arztberichte (<10J), Schulberichte
        low_prio_parts = []     # Arztberichte aelter als 10 Jahre
        stufe2_parts = []
        extended_parts = []

        ten_years_ago = datetime.now() - timedelta(days=10 * 365)

        for doc in documents:
            if doc.category == DocumentCategory.SKIP:
                continue
            if doc.category == DocumentCategory.EXTENDED and not include_extended:
                continue

            text = self._extract_text_from_file(doc.path)
            if not text or text.startswith("[FEHLER"):
                continue

            text = anonymize_text(text)

            doc.text_extracted = True
            doc.text_length = len(text)
            bundle.files_processed += 1

            anon_filename = anonymize_text(doc.filename)
            header = f"--- {doc.doc_type.upper()}: {anon_filename} ---"
            content = f"{header}\n{text}"

            if doc.category == DocumentCategory.CORE:
                # Innerhalb CORE nach Typ-Prioritaet aufteilen
                if doc.doc_type in ("protokoll", "aktendeckblatt", "hilfeplan"):
                    high_prio_parts.append(content)
                elif doc.doc_type == "arztbericht" and doc.date_hint and doc.date_hint < ten_years_ago:
                    low_prio_parts.append(content)
                else:
                    medium_prio_parts.append(content)
            elif doc.category == DocumentCategory.STUFE2:
                if doc.doc_type == "arztbericht" and doc.date_hint and doc.date_hint < ten_years_ago:
                    low_prio_parts.append(content)
                else:
                    stufe2_parts.append(content)
            else:
                extended_parts.append(content)

        # Bundle mit Prioritaets-Sections zusammenbauen
        core_sections = []
        if high_prio_parts:
            core_sections.append(
                "=== HOHE PRIORITAET: Aktuelle Protokolle, Dokumentation, "
                "Hilfeplan, Bewilligung, Aktendeckblatt ==="
            )
            core_sections.extend(high_prio_parts)
        if medium_prio_parts:
            core_sections.append(
                "\n=== MITTLERE PRIORITAET: Foerderberichte, "
                "Arztberichte, Netzwerkarbeit ==="
            )
            core_sections.extend(medium_prio_parts)
        if low_prio_parts:
            core_sections.append(
                "\n=== NIEDRIGE PRIORITAET: Aeltere Berichte (>10 Jahre) ==="
            )
            core_sections.extend(low_prio_parts)

        bundle.core_text = "\n\n".join(core_sections)
        bundle.stufe2_text = "\n\n".join(stufe2_parts)
        bundle.extended_text = "\n\n".join(extended_parts)
        bundle.total_length = len(bundle.core_text) + len(bundle.stufe2_text)

        return bundle

    def save_bundle(
        self,
        bundle: TextBundle,
        output_path: str,
        combined: bool = True
    ) -> Path:
        """
        Speichert ein Bundle als Textdatei.

        Args:
            bundle: Das TextBundle
            output_path: Zielpfad
            combined: CORE + STUFE2 kombiniert speichern

        Returns:
            Pfad zur gespeicherten Datei
        """
        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)

        if combined:
            content = "=== CORE DOKUMENTE ===\n\n"
            content += bundle.core_text
            content += "\n\n=== STUFE 2 DOKUMENTE ===\n\n"
            content += bundle.stufe2_text
            path.write_text(content, encoding="utf-8")
        else:
            core_path = path.with_suffix(".core.txt")
            core_path.write_text(bundle.core_text, encoding="utf-8")
            stufe2_path = path.with_suffix(".stufe2.txt")
            stufe2_path.write_text(bundle.stufe2_text, encoding="utf-8")

        return path

    # ─────────────────────────────────────────────────────────────
    # Interface 3: Bundle-Verzeichnis lesen (FileAccessHook-Compat.)
    # ─────────────────────────────────────────────────────────────

    def read_bundle_dir(
        self,
        bundle_path: Path,
        filename_filter: Optional[str] = None,
        include_extended: bool = False
    ) -> str:
        """
        Liest Inhalte aus einem anonymisierten Bundle-Verzeichnis.

        Ersetzt FileAccessHook.get_safe_content() und _get_bundle_content().
        Das Bundle-Verzeichnis enthaelt bereits anonymisierte Dateien.

        Args:
            bundle_path: Pfad zum Bundle-Ordner (z.B. bundles/K_XXX/)
            filename_filter: Optionaler Dateiname-Filter
            include_extended: Auch extended/ einbeziehen

        Returns:
            Zusammengefuehrter Textinhalt aus dem Bundle
        """
        folders = [bundle_path / "core"]
        if include_extended and (bundle_path / "extended").exists():
            folders.append(bundle_path / "extended")

        all_content = []

        for folder in folders:
            if not folder.exists():
                continue

            for filepath in sorted(folder.rglob("*")):
                if not filepath.is_file():
                    continue
                if filepath.suffix.lower() not in BUNDLE_READ_EXTENSIONS:
                    continue
                if filepath.name.startswith(".") or filepath.name.startswith("_"):
                    continue
                if "output" in filepath.parts:
                    continue

                if filename_filter and filename_filter.lower() not in filepath.name.lower():
                    continue

                text = self._extract_text_from_file(filepath)
                if text.strip():
                    rel_path = filepath.relative_to(bundle_path)
                    all_content.append(f"--- Quelle: {rel_path} ---\n{text}")

        if not all_content:
            return "[KEINE DATEN] Keine passenden Dokumente im Bundle gefunden."

        return "\n\n".join(all_content)

    def list_bundle_files(self, bundle_path: Path) -> List[dict]:
        """
        Listet alle verfuegbaren Dateien in einem Bundle auf.

        Ersetzt FileAccessHook.list_available_files().

        Args:
            bundle_path: Pfad zum Bundle-Ordner

        Returns:
            Liste von Datei-Infos
        """
        if not bundle_path.exists():
            return []

        files = []
        for filepath in sorted(bundle_path.rglob("*")):
            if not filepath.is_file():
                continue
            if filepath.name.startswith("."):
                continue
            if "output" in filepath.parts:
                continue

            rel_path = filepath.relative_to(bundle_path)
            files.append({
                "name": filepath.name,
                "path": str(rel_path),
                "size": filepath.stat().st_size,
                "type": filepath.suffix.lower()
            })

        return files

    # ─────────────────────────────────────────────────────────────
    # Anonymisierungs-Hilfsfunktionen
    # ─────────────────────────────────────────────────────────────

    def apply_anonymization(
        self,
        bundle: TextBundle,
        profile: "AnonymProfile"
    ) -> TextBundle:
        """
        Anonymisiert ein fertiges Text-Bundle nachtraeglich.

        Kann eingesetzt werden wenn extract_bundle() ohne anonym_profile
        aufgerufen wurde und die Anonymisierung erst danach erfolgen soll.

        Args:
            bundle: Extrahiertes TextBundle
            profile: Anonymisierungsprofil

        Returns:
            Neues TextBundle mit anonymisiertem Text
        """
        all_replacements = {}
        for category in profile.mappings.values():
            all_replacements.update(category)

        sorted_replacements = sorted(
            all_replacements.items(),
            key=lambda x: len(x[0]),
            reverse=True
        )

        def anonymize_text(text: str) -> str:
            for old, new in sorted_replacements:
                text = text.replace(old, new)
            return text

        return TextBundle(
            core_text=anonymize_text(bundle.core_text),
            stufe2_text=anonymize_text(bundle.stufe2_text),
            extended_text=anonymize_text(bundle.extended_text),
            total_length=bundle.total_length,
            files_processed=bundle.files_processed
        )

    # ─────────────────────────────────────────────────────────────
    # Text-Extraktion (konsolidierter Code aus beiden Quell-Dateien)
    # ─────────────────────────────────────────────────────────────

    def _extract_text_from_file(self, filepath: Path) -> str:
        """
        Extrahiert Text aus einer Datei.

        Konsolidiert die identische Logik aus:
          - DocumentCollector._extract_text()
          - FileAccessHook._extract_text()

        Zusaetzlich gegenueber FileAccessHook: DOC, MSG, EML Unterstuetzung.
        """
        suffix = filepath.suffix.lower()
        filepath_str = str(filepath)

        try:
            if suffix == ".docx":
                return self._extract_docx(filepath_str)
            elif suffix == ".doc":
                return self._extract_doc(filepath_str)
            elif suffix in [".txt", ".md"]:
                return self._extract_txt(filepath_str)
            elif suffix == ".pdf":
                return self._extract_pdf(filepath_str)
            elif suffix in [".xlsx", ".xls"]:
                return self._extract_excel(filepath_str)
            elif suffix == ".msg":
                return self._extract_msg(filepath_str)
            elif suffix == ".eml":
                return self._extract_eml(filepath_str)
            else:
                return f"[Nicht unterstuetztes Format: {suffix}]"
        except Exception as e:
            return f"[FEHLER bei {filepath.name}: {e}]"

    def _extract_docx(self, filepath: str) -> str:
        """Extrahiert Text aus Word-Dokument (.docx)."""
        try:
            from docx import Document
            doc = Document(filepath)
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        except ImportError:
            return "[python-docx nicht installiert]"
        except Exception as e:
            return f"[DOCX-Fehler: {e}]"

    def _extract_doc(self, filepath: str) -> str:
        """Extrahiert Text aus altem Word-Format (.doc) via antiword oder LibreOffice."""
        import subprocess
        import shutil

        if shutil.which("antiword"):
            try:
                result = subprocess.run(
                    ["antiword", filepath],
                    capture_output=True, text=True, timeout=30
                )
                if result.returncode == 0 and result.stdout.strip():
                    return result.stdout.strip()
            except Exception:
                pass

        if shutil.which("soffice"):
            try:
                import tempfile
                with tempfile.TemporaryDirectory() as tmpdir:
                    result = subprocess.run(
                        ["soffice", "--headless", "--convert-to", "txt:Text",
                         "--outdir", tmpdir, filepath],
                        capture_output=True, timeout=60
                    )
                    if result.returncode == 0:
                        txt_file = Path(tmpdir) / (Path(filepath).stem + ".txt")
                        if txt_file.exists():
                            return txt_file.read_text(encoding="utf-8", errors="replace")
            except Exception:
                pass

        try:
            import textract
            text = textract.process(filepath).decode("utf-8", errors="replace")
            return text.strip()
        except ImportError:
            pass
        except Exception:
            pass

        return "[.doc-Extraktion fehlgeschlagen - antiword/LibreOffice/textract nicht verfuegbar]"

    def _extract_txt(self, filepath: str) -> str:
        """Liest Textdatei."""
        path = Path(filepath)
        try:
            return path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return path.read_text(encoding="latin-1")

    def _extract_pdf(self, filepath: str) -> str:
        """Extrahiert Text aus PDF, mit OCR-Fallback fuer Bild-PDFs."""
        text_parts = []
        needs_ocr = False

        # Primaer: pypdf (MIT-Lizenz)
        try:
            from pypdf import PdfReader
            reader = PdfReader(filepath)
            for page in reader.pages:
                page_text = (page.extract_text() or "").strip()
                if page_text:
                    text_parts.append(page_text)
                else:
                    needs_ocr = True
            if text_parts:
                return "\n".join(text_parts)
        except ImportError:
            pass
        except Exception:
            pass

        # Fallback: pdfplumber (MIT-Lizenz)
        try:
            import pdfplumber
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    page_text = (page.extract_text() or "").strip()
                    if page_text:
                        text_parts.append(page_text)
                    else:
                        needs_ocr = True
            if text_parts:
                return "\n".join(text_parts)
        except ImportError:
            pass
        except Exception:
            pass

        # Optional: PyMuPDF (AGPL)
        try:
            import fitz
            doc = fitz.open(filepath)
            fitz_parts = []
            fitz_needs_ocr = False
            for page in doc:
                page_text = page.get_text().strip()
                if page_text:
                    fitz_parts.append(page_text)
                else:
                    fitz_needs_ocr = True
            doc.close()
            if fitz_parts:
                return "\n".join(fitz_parts)
        except ImportError:
            pass
        except Exception:
            pass

        if needs_ocr:
            return self._extract_pdf_ocr(filepath)

        return "[PDF-Extraktion fehlgeschlagen: pypdf, pdfplumber und PyMuPDF nicht verfuegbar]"

    def _extract_pdf_ocr(self, filepath: str) -> str:
        """OCR fuer Bild-PDFs via tools.ocr.engine.OCREngine."""
        try:
            try:
                from tools.ocr.engine import OCREngine
            except ImportError:
                # Fallback: tools-Verzeichnis manuell zum sys.path hinzufuegen
                tools_dir = Path(__file__).resolve().parent.parent.parent.parent / "tools"
                if tools_dir.exists() and str(tools_dir.parent) not in sys.path:
                    sys.path.insert(0, str(tools_dir.parent))
                from tools.ocr.engine import OCREngine

            engine = OCREngine()
            if not engine.available:
                return "[OCR nicht verfuegbar - Tesseract nicht installiert]"

            text = engine.extract_text(filepath)
            return text if text else "[OCR: Kein Text erkannt]"
        except ImportError as e:
            return f"[OCR Import-Fehler: {e}]"
        except Exception as e:
            return f"[OCR-Fehler: {e}]"

    def _extract_excel(self, filepath: str) -> str:
        """Extrahiert Text aus Excel."""
        try:
            import openpyxl
            wb = openpyxl.load_workbook(filepath, data_only=True)
            parts = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                parts.append(f"[Tabelle: {sheet_name}]")
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(c) if c else "" for c in row]
                    if any(cells):
                        parts.append(" | ".join(cells))
            wb.close()
            return "\n".join(parts)
        except ImportError:
            return "[openpyxl nicht installiert]"

    def _extract_msg(self, filepath: str) -> str:
        """Extrahiert Text aus Outlook .msg (inkl. Anhaenge)."""
        try:
            import extract_msg
            msg = extract_msg.Message(filepath)
            parts = []
            if msg.date:
                parts.append(f"Datum: {msg.date}")
            if msg.sender:
                parts.append(f"Von: {msg.sender}")
            if msg.to:
                parts.append(f"An: {msg.to}")
            if msg.subject:
                parts.append(f"Betreff: {msg.subject}")
            parts.append("")
            if msg.body:
                parts.append(msg.body)

            if hasattr(msg, 'attachments') and msg.attachments:
                for attach in msg.attachments:
                    try:
                        attach_name = attach.longFilename or attach.shortFilename or "unbekannt"
                        suffix = Path(attach_name).suffix.lower()

                        if suffix == ".pdf":
                            import tempfile
                            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                                tmp.write(attach.data)
                                tmp_path = tmp.name
                            parts.append(f"\n[Anhang: {attach_name}]")
                            parts.append(self._extract_pdf(tmp_path))
                            Path(tmp_path).unlink(missing_ok=True)

                        elif suffix == ".docx":
                            import tempfile
                            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                                tmp.write(attach.data)
                                tmp_path = tmp.name
                            parts.append(f"\n[Anhang: {attach_name}]")
                            parts.append(self._extract_docx(tmp_path))
                            Path(tmp_path).unlink(missing_ok=True)

                        elif suffix in [".txt", ".md"]:
                            parts.append(f"\n[Anhang: {attach_name}]")
                            try:
                                text = attach.data.decode('utf-8')
                            except Exception:
                                text = attach.data.decode('latin-1')
                            parts.append(text)

                    except Exception as e:
                        parts.append(f"\n[Anhang {attach_name}: Fehler - {e}]")

            msg.close()
            return "\n".join(parts)
        except ImportError:
            return "[extract-msg nicht installiert]"

    def _extract_eml(self, filepath: str) -> str:
        """Extrahiert Text aus .eml."""
        import email
        from email import policy
        from email.parser import BytesParser

        with open(filepath, 'rb') as f:
            msg = BytesParser(policy=policy.default).parse(f)

        parts = []
        if msg['Date']:
            parts.append(f"Datum: {msg['Date']}")
        if msg['From']:
            parts.append(f"Von: {msg['From']}")
        if msg['To']:
            parts.append(f"An: {msg['To']}")
        if msg['Subject']:
            parts.append(f"Betreff: {msg['Subject']}")
        parts.append("")

        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type() == 'text/plain':
                    charset = part.get_content_charset() or 'utf-8'
                    try:
                        body = part.get_payload(decode=True).decode(charset)
                        parts.append(body)
                        break
                    except Exception:
                        pass
        else:
            charset = msg.get_content_charset() or 'utf-8'
            try:
                body = msg.get_payload(decode=True).decode(charset)
                parts.append(body)
            except Exception:
                parts.append(str(msg.get_payload()))

        return "\n".join(parts)


# ═══════════════════════════════════════════════════════════════
# Convenience-Funktionen (ersetzen Top-Level-Funktionen aus
# document_collector.py und file_access_hook.py)
# ═══════════════════════════════════════════════════════════════

def scan_and_extract(
    folder_path: str,
    geburtsdatum: str = "01.01.2010",
    output_path: Optional[str] = None,
    include_extended: bool = False
) -> Tuple[TextBundle, Optional["AnonymProfile"]]:
    """
    Kompletter Workflow: Scannen + Extrahieren + optional Anonymisieren.

    Wenn 'geburtsdatum' angegeben ist, wird automatisch ein Anonymisierungsprofil
    aus dem Ordnernamen erstellt (gleicher Mechanismus wie create_profile_from_folder).

    Args:
        folder_path: Pfad zum Klienten-Ordner
        geburtsdatum: Geburtsdatum des Klienten (fuer Anonymisierung)
        output_path: Optional - Pfad zum Speichern des Bundles
        include_extended: EXTENDED-Dokumente einbeziehen

    Returns:
        (TextBundle, AnonymProfile oder None)
    """
    profile = None

    try:
        from .anonymizer_service import DocumentAnonymizer
        folder = Path(folder_path)
        folder_name = folder.name
        if ", " in folder_name:
            parts = folder_name.split(", ")
            name = f"{parts[1]} {parts[0]}"
        else:
            name = folder_name

        anonymizer = DocumentAnonymizer()
        profile = anonymizer.create_profile(
            real_name=name.strip(),
            geburtsdatum=geburtsdatum,
            whitelist=[]
        )
    except (ImportError, Exception):
        pass

    pipeline = DocumentPipeline()
    result = pipeline.scan_folder(folder_path)
    bundle = pipeline.extract_bundle(
        result.documents,
        include_extended=include_extended,
        anonym_profile=profile
    )

    if output_path:
        pipeline.save_bundle(bundle, output_path)

    return bundle, profile
