#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# SPDX-License-Identifier: MIT
"""
Copyright (c) 2026 Lukas Geiger

Permission is hereby granted, free of charge, to any person obtaining a copy
of this software and associated documentation files (the "Software"), to deal
in the Software without restriction, including without limitation the rights
to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
copies of the Software, and to permit persons to whom the Software is
furnished to do so, subject to the following conditions:

The above copyright notice and this permission notice shall be included in all
copies or substantial portions of the Software.

THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
SOFTWARE.
"""

"""
Word Template Service
=====================

Allgemeiner Service fuer Word-Template-Befuellung via python-docx.
Unterstuetzt Platzhalter-Ersetzung, Tabellenzeilen-Filterung,
Checkbox-Steuerung und Satzbausteine.

Abhaengigkeit: python-docx (pip install python-docx)

Version: 1.1.0 (Tabellen-Support für Förderziele)
Erstellt: 2026-01-27
Aktualisiert: 2026-01-31
"""

import copy
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Callable

from docx import Document
from docx.table import Table, _Row
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn


# ═══════════════════════════════════════════════════════════════
# Hilfsfunktionen
# ═══════════════════════════════════════════════════════════════

def _get_full_text(paragraph: Paragraph) -> str:
    """Gibt den vollstaendigen Text eines Paragraphen zurueck (ueber alle Runs)."""
    return "".join(run.text for run in paragraph.runs)


def _replace_in_paragraph(paragraph: Paragraph, old: str, new: str) -> bool:
    """
    Ersetzt Text in einem Paragraphen unter Beibehaltung der Formatierung.
    Behandelt den Fall, dass der Suchtext ueber mehrere Runs verteilt ist.
    """
    full_text = _get_full_text(paragraph)
    if old not in full_text:
        return False

    # Einfacher Fall: Text in einem einzigen Run
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True

    # Komplexer Fall: Text ueber mehrere Runs verteilt
    # Runs zusammenfuegen, ersetzen, dann zurueckschreiben
    new_text = full_text.replace(old, new)
    for i, run in enumerate(paragraph.runs):
        if i == 0:
            run.text = new_text
        else:
            run.text = ""
    return True


def _replace_in_cell(cell, old: str, new: str) -> bool:
    """
    Ersetzt Text in einer Tabellenzelle (alle Paragraphen).
    Beruecksichtigt auch Inhalte in Content Controls (SDT).
    """
    replaced = False

    # Normale Paragraphen
    for paragraph in cell.paragraphs:
        if _replace_in_paragraph(paragraph, old, new):
            replaced = True

    # Auch Paragraphen innerhalb von Content Controls (SDT)
    # Diese werden von cell.paragraphs nicht erfasst
    for sdt in cell._element.iter(qn("w:sdt")):
        sdt_content = sdt.find(qn("w:sdtContent"))
        if sdt_content is not None:
            for p_elem in sdt_content.findall(qn("w:p")):
                # Text zusammenbauen und pruefen
                full_text = "".join(t.text or "" for t in p_elem.iter(qn("w:t")))
                if old in full_text:
                    # Ersetze in allen w:t Elementen
                    for t_elem in p_elem.iter(qn("w:t")):
                        if t_elem.text and old in t_elem.text:
                            t_elem.text = t_elem.text.replace(old, new)
                            replaced = True

    return replaced


def _get_cell_text(cell) -> str:
    """Gibt den gesamten Text einer Tabellenzelle zurueck."""
    return "\n".join(p.text for p in cell.paragraphs).strip()


# ═══════════════════════════════════════════════════════════════
# Hauptklasse
# ═══════════════════════════════════════════════════════════════

class WordTemplateService:
    """
    Allgemeiner Service fuer Word-Template-Operationen.

    Verwendung:
        svc = WordTemplateService()
        doc = svc.load_template("vorlage.docx")
        svc.replace_placeholders(doc, {"{{NAME}}": "Max Mustermann"})
        svc.filter_table_rows(doc.tables[2], keep_column=0, keep_values=["D350", "D710"])
        svc.set_checkbox(doc, "Einzelfoerderung", checked=True)
        svc.save(doc, "bericht.docx")
    """

    def load_template(self, template_path: str) -> Document:
        """Laedt eine Word-Vorlage."""
        path = Path(template_path)
        if not path.exists():
            raise FileNotFoundError(f"Vorlage nicht gefunden: {template_path}")
        return Document(str(path))

    def replace_placeholders(self, doc: Document, data: Dict[str, str]) -> int:
        """
        Ersetzt Platzhalter im gesamten Dokument (Paragraphen + Tabellen).

        Args:
            doc: Word-Dokument
            data: Mapping von Platzhalter -> Ersetzungstext
                  z.B. {"{{NAME}}": "Max Mustermann", "Vorname und Nachname": "Max M."}

        Returns:
            Anzahl der durchgefuehrten Ersetzungen
        """
        count = 0

        # Paragraphen (alle, auch in Textboxen)
        # doc.paragraphs findet nicht alles (z.B. Textboxen), daher XML-Iteration
        body = doc._element.body
        for p_elem in body.iter(qn("w:p")):
            paragraph = Paragraph(p_elem, doc)
            for old, new in data.items():
                if _replace_in_paragraph(paragraph, old, new):
                    count += 1

        # Direkte Suche in allen w:t Elementen (erfasst auch SDT-Inhalte)
        body = doc._element.body
        for t_elem in body.iter(qn("w:t")):
            if t_elem.text:
                for old, new in data.items():
                    if old in t_elem.text:
                        t_elem.text = t_elem.text.replace(old, new)
                        count += 1

        # Tabellen (inkl. verschachtelte)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for old, new in data.items():
                        if _replace_in_cell(cell, old, new):
                            count += 1

        # Header und Footer
        for section in doc.sections:
            for header_para in section.header.paragraphs:
                for old, new in data.items():
                    if _replace_in_paragraph(header_para, old, new):
                        count += 1
            for footer_para in section.footer.paragraphs:
                for old, new in data.items():
                    if _replace_in_paragraph(footer_para, old, new):
                        count += 1

        return count

    def filter_table_rows(
        self,
        table: Table,
        keep_column: int,
        keep_values: List[str],
        header_rows: int = 1,
        keep_section_headers: bool = True
    ) -> int:
        """
        Entfernt Zeilen aus einer Tabelle, deren Wert in einer bestimmten
        Spalte NICHT in der keep_values-Liste ist.

        Args:
            table: Die zu filternde Tabelle
            keep_column: Spaltenindex (0-basiert) fuer den Filterwert
            keep_values: Liste der Werte, deren Zeilen behalten werden
            header_rows: Anzahl der Header-Zeilen (werden nie geloescht)
            keep_section_headers: Behaelt Zeilen ohne ICF-Code (Kapitel-Header)

        Returns:
            Anzahl geloeschter Zeilen
        """
        rows_to_remove = []
        keep_values_set = set(v.strip().upper() for v in keep_values)

        for i, row in enumerate(table.rows):
            # Header-Zeilen nie loeschen
            if i < header_rows:
                continue

            cell_text = _get_cell_text(row.cells[keep_column]).strip().upper()

            # Leere Zelle in der Filterspalte = Kapitel-Header
            if not cell_text and keep_section_headers:
                continue

            if cell_text not in keep_values_set:
                rows_to_remove.append(row)

        # Zeilen von hinten loeschen (Index-Stabilität)
        for row in reversed(rows_to_remove):
            tbl = table._tbl
            tbl.remove(row._tr)

        return len(rows_to_remove)

    def set_checkbox(self, doc: Document, label_contains: str, checked: bool) -> int:
        """
        Setzt Checkboxen basierend auf dem umgebenden Text.
        Unterstuetzt sowohl einfache Unicode-Checkboxen (☐/☒)
        als auch Word Content Control Checkboxen (SDT/w14:checkbox).

        Args:
            label_contains: Text neben der Checkbox
            checked: True = ☒, False = ☐

        Returns:
            Anzahl geaenderter Checkboxen
        """
        count = 0
        check_char = "\u2612"    # ☒
        uncheck_char = "\u2610"  # ☐

        NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        NS_W14 = "http://schemas.microsoft.com/office/word/2010/wordml"

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = _get_cell_text(cell)
                    if label_contains not in cell_text:
                        continue

                    # Methode 1: Word Content Control (SDT) Checkboxen
                    for sdt in cell._element.iter(qn("w:sdt")):
                        sdt_pr = sdt.find(qn("w:sdtPr"))
                        if sdt_pr is None:
                            continue
                        checkbox_elem = sdt_pr.find(f"{{{NS_W14}}}checkbox")
                        if checkbox_elem is None:
                            continue

                        # w14:checked Wert setzen
                        checked_elem = checkbox_elem.find(f"{{{NS_W14}}}checked")
                        if checked_elem is not None:
                            checked_elem.set(f"{{{NS_W14}}}val", "1" if checked else "0")

                        # Run-Text aktualisieren (sichtbares Symbol)
                        sdt_content = sdt.find(qn("w:sdtContent"))
                        if sdt_content is not None:
                            for t_elem in sdt_content.iter(qn("w:t")):
                                if t_elem.text in (check_char, uncheck_char):
                                    t_elem.text = check_char if checked else uncheck_char

                        count += 1

                    # Methode 2: Einfache Unicode-Checkboxen (Fallback)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if checked and uncheck_char in run.text:
                                run.text = run.text.replace(uncheck_char, check_char)
                                count += 1
                            elif not checked and check_char in run.text:
                                run.text = run.text.replace(check_char, uncheck_char)
                                count += 1

        return count

    def set_cell_text(
        self,
        table: Table,
        row_index: int,
        col_index: int,
        text: str,
        preserve_formatting: bool = True
    ) -> bool:
        """
        Setzt den Text einer bestimmten Tabellenzelle.

        Args:
            table: Tabelle
            row_index: Zeilenindex (0-basiert)
            col_index: Spaltenindex (0-basiert)
            text: Neuer Text
            preserve_formatting: Formatierung des ersten Runs beibehalten
        """
        if row_index >= len(table.rows) or col_index >= len(table.rows[row_index].cells):
            return False

        cell = table.rows[row_index].cells[col_index]

        if preserve_formatting and cell.paragraphs and cell.paragraphs[0].runs:
            # Formatierung des ersten Runs beibehalten
            first_run = cell.paragraphs[0].runs[0]
            first_run.text = text
            # Weitere Runs leeren
            for run in cell.paragraphs[0].runs[1:]:
                run.text = ""
            # Weitere Paragraphen leeren
            for para in cell.paragraphs[1:]:
                for run in para.runs:
                    run.text = ""
        else:
            cell.text = text

        return True

    def find_table_row_by_text(
        self,
        table: Table,
        search_text: str,
        column: int = 0
    ) -> Optional[int]:
        """
        Findet eine Tabellenzeile anhand eines Textes in einer bestimmten Spalte.

        Returns:
            Zeilenindex oder None
        """
        for i, row in enumerate(table.rows):
            if column < len(row.cells):
                cell_text = _get_cell_text(row.cells[column]).strip()
                if search_text.strip().upper() == cell_text.upper():
                    return i
        return None

    def activate_textblock(
        self,
        doc: Document,
        marker: str,
        replacement_text: str
    ) -> bool:
        """
        Ersetzt einen Textblock-Marker (z.B. [MOBILITÄT]) durch individualisierten Text.

        Args:
            marker: Der Marker-Text (z.B. "[MOBILITÄT]")
            replacement_text: Der neue Text

        Returns:
            True wenn Marker gefunden und ersetzt
        """
        for paragraph in doc.paragraphs:
            full_text = _get_full_text(paragraph)
            if marker in full_text:
                _replace_in_paragraph(paragraph, marker, "")
                # Den gesamten Paragraphen-Text ersetzen
                if paragraph.runs:
                    paragraph.runs[0].text = replacement_text
                    for run in paragraph.runs[1:]:
                        run.text = ""
                return True

        # Auch in Tabellenzellen suchen
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        full_text = _get_full_text(paragraph)
                        if marker in full_text:
                            if paragraph.runs:
                                paragraph.runs[0].text = replacement_text
                                for run in paragraph.runs[1:]:
                                    run.text = ""
                            return True
        return False

    def remove_textblock(
        self,
        doc: Document,
        marker: str,
        remove_paragraph: bool = True
    ) -> bool:
        """
        Entfernt einen Textblock-Marker und optional den gesamten Paragraphen.

        Sucht nach dem Marker und entfernt den gesamten Paragraphen oder
        nur den Marker-Text, abhaengig von remove_paragraph.
        """
        # In Paragraphen suchen
        for paragraph in doc.paragraphs:
            full_text = _get_full_text(paragraph)
            if marker in full_text:
                if remove_paragraph:
                    # Paragraphen aus dem XML-Baum entfernen
                    p_element = paragraph._element
                    p_element.getparent().remove(p_element)
                else:
                    _replace_in_paragraph(paragraph, marker, "")
                return True

        # In Tabellenzellen suchen
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        full_text = _get_full_text(paragraph)
                        if marker in full_text:
                            if remove_paragraph:
                                p_element = paragraph._element
                                p_element.getparent().remove(p_element)
                            else:
                                _replace_in_paragraph(paragraph, marker, "")
                            return True
        return False

    def remove_table_row(self, table: Table, row_index: int) -> bool:
        """Entfernt eine bestimmte Zeile aus einer Tabelle."""
        if row_index < 0 or row_index >= len(table.rows):
            return False
        tbl = table._tbl
        tr = table.rows[row_index]._tr
        tbl.remove(tr)
        return True

    def get_table_data(self, table: Table) -> List[List[str]]:
        """Gibt den Inhalt einer Tabelle als 2D-Liste zurueck."""
        data = []
        for row in table.rows:
            row_data = [_get_cell_text(cell) for cell in row.cells]
            data.append(row_data)
        return data

    def fill_icf_placeholders_and_cleanup(
        self,
        table: Table,
        foerderziele: List[Dict],
        header_rows: int = 1
    ) -> Tuple[int, int]:
        """
        Ersetzt ICF-Zeilen-Platzhalter und loescht Zeilen ohne Daten.

        Template-Format pro Zeile:
          ICF-Code | Kapitel | {CODE-Ziel} | {CODE-Ist} | {CODE-E} | {CODE-G}

        Beispiel:
          D350 | Konversation | {D350-Ziel} | {D350-Ist} | {D350-E} | {D350-G}

        Logik:
        1. Fuer jedes Foerderziel: Ersetze {CODE-Ziel}, {CODE-Ist}, {CODE-E}, {CODE-G}
        2. Loesche alle Zeilen die noch ungefuellte {xxx} Platzhalter haben

        Args:
            table: Die ICF-Tabelle
            foerderziele: Liste von Foerderzielen
                          [{"icf_code": "D350", "zielformulierung": "...",
                            "ist_stand": "...", "erreicht": 2, "grund_nichterreichung": null}]
            header_rows: Anzahl Header-Zeilen (werden nie geloescht)

        Returns:
            (ersetzungen, geloeschte_zeilen)
        """
        import re
        placeholder_pattern = re.compile(r'\{[^}]+\}')

        # Mapping: ICF-Code -> Daten
        ziel_map = {}
        for ziel in foerderziele:
            code = ziel.get("icf_code", "").upper()
            if code:
                erreicht_map = {1: "(1)", 2: "(2)", 3: "(3)"}
                grund_map = {1: "(1)", 2: "(2)", 3: "(3)", 4: "(4)"}
                ziel_map[code] = {
                    "Ziel": ziel.get("zielformulierung", ""),
                    "Ist": ziel.get("ist_stand", ""),
                    "E": erreicht_map.get(ziel.get("erreicht"), ""),
                    "G": grund_map.get(ziel.get("grund_nichterreichung"), "")
                }

        replacements = 0
        rows_to_delete = []

        # Durch alle Zeilen gehen
        for row_idx, row in enumerate(table.rows):
            if row_idx < header_rows:
                continue

            # ICF-Code aus erster Spalte lesen
            if not row.cells:
                continue
            icf_code = _get_cell_text(row.cells[0]).strip().upper()

            # Platzhalter in dieser Zeile ersetzen
            row_has_unfilled = False
            for cell in row.cells:
                for para in cell.paragraphs:
                    full_text = _get_full_text(para)

                    # Suche nach Platzhaltern wie {D350-Ziel}, {D350-Ist}, etc.
                    matches = placeholder_pattern.findall(full_text)
                    for match in matches:
                        # Extrahiere Code und Feldname: {D350-Ziel} -> D350, Ziel
                        inner = match[1:-1]  # ohne { }
                        if "-" in inner:
                            ph_code, ph_field = inner.split("-", 1)
                            ph_code = ph_code.upper()

                            if ph_code in ziel_map and ph_field in ziel_map[ph_code]:
                                replacement = ziel_map[ph_code][ph_field]
                                # Ersetze den Platzhalter - auch mit leerem String
                                # (z.B. wenn grund_nichterreichung nicht gesetzt ist)
                                _replace_in_paragraph(para, match, replacement)
                                replacements += 1
                            else:
                                # Kein Mapping fuer diesen Code -> Zeile loeschen
                                row_has_unfilled = True
                        else:
                            row_has_unfilled = True

            # Zeile zum Loeschen markieren wenn ungefuellte Platzhalter
            if row_has_unfilled:
                rows_to_delete.append(row)

        # Zeilen von hinten loeschen
        for row in reversed(rows_to_delete):
            table._tbl.remove(row._tr)

        return replacements, len(rows_to_delete)

    def fill_table_rows(
        self,
        table: Table,
        data_rows: List[Dict[str, str]],
        column_mapping: Dict[str, int],
        template_row_index: int = 1,
        clear_template: bool = True
    ) -> int:
        """
        Fuellt eine Tabelle mit dynamischen Daten aus einer Liste von Dictionaries.

        Ideal fuer Foerderziele-Tabellen mit variabler Anzahl von Zeilen.

        Args:
            table: Die zu befuellende Tabelle
            data_rows: Liste von Dictionaries mit den Daten pro Zeile
                       z.B. [{"icf_code": "D350", "ist_stand": "...", "erreicht": "2"}]
            column_mapping: Mapping von Dictionary-Key -> Spaltenindex
                            z.B. {"icf_code": 0, "ist_stand": 2, "erreicht": 3}
            template_row_index: Index der Vorlagezeile (wird als Formatvorlage genutzt)
            clear_template: Vorlagezeile nach dem Kopieren loeschen

        Returns:
            Anzahl eingefuegter Zeilen
        """
        if not data_rows:
            return 0

        if template_row_index >= len(table.rows):
            return 0

        template_row = table.rows[template_row_index]
        tbl = table._tbl
        template_tr = template_row._tr

        # Neue Zeilen nach der Vorlagezeile einfuegen
        inserted = 0
        for data in data_rows:
            # Vorlagezeile kopieren (Deep Copy des XML)
            new_tr = copy.deepcopy(template_tr)
            template_tr.addnext(new_tr)

            # Daten in die neue Zeile schreiben
            # Wir muessen die Zeile neu aus der Tabelle holen
            new_row_index = template_row_index + inserted + 1
            if new_row_index < len(table.rows):
                new_row = table.rows[new_row_index]
                for key, col_idx in column_mapping.items():
                    if key in data and col_idx < len(new_row.cells):
                        cell = new_row.cells[col_idx]
                        # Text setzen, Formatierung beibehalten
                        if cell.paragraphs and cell.paragraphs[0].runs:
                            cell.paragraphs[0].runs[0].text = str(data[key])
                            for run in cell.paragraphs[0].runs[1:]:
                                run.text = ""
                        else:
                            cell.text = str(data[key])
            inserted += 1

        # Vorlagezeile entfernen wenn gewuenscht
        if clear_template:
            tbl.remove(template_tr)

        return inserted

    def fill_foerderziele_table(
        self,
        doc: Document,
        foerderziele: List[Dict],
        table_index: int = 0
    ) -> int:
        """
        Spezialisierte Methode fuer Foerderziele-Tabellen.

        Sucht eine Tabelle mit typischen Foerderziel-Spalten und fuellt sie.

        Args:
            doc: Word-Dokument
            foerderziele: Liste von Foerderzielen aus dem JSON
                          [{"icf_code": "D350", "zielformulierung": "...",
                            "ist_stand": "...", "erreicht": 2}]
            table_index: Index der Tabelle (falls mehrere)

        Returns:
            Anzahl eingefuegter Zeilen
        """
        if table_index >= len(doc.tables):
            return 0

        table = doc.tables[table_index]

        # Header analysieren um Spalten zu identifizieren
        if not table.rows:
            return 0

        header_texts = [_get_cell_text(c).lower() for c in table.rows[0].cells]

        # Typische Spalten-Keywords fuer Foerderziele
        column_mapping = {}
        for i, header in enumerate(header_texts):
            if "icf" in header or "code" in header:
                column_mapping["icf_code"] = i
            elif "ziel" in header and "formulierung" in header:
                column_mapping["zielformulierung"] = i
            elif "ist" in header or "stand" in header or "beschreibung" in header:
                column_mapping["ist_stand"] = i
            elif "erreicht" in header or "bewertung" in header:
                column_mapping["erreicht"] = i
            elif "grund" in header:
                column_mapping["grund_nichterreichung"] = i

        if not column_mapping:
            # Fallback: Einfache Reihenfolge annehmen
            column_mapping = {
                "icf_code": 0,
                "zielformulierung": 1,
                "ist_stand": 2,
                "erreicht": 3
            }

        # Daten vorbereiten
        data_rows = []
        for ziel in foerderziele:
            row_data = {}
            for key in column_mapping.keys():
                if key in ziel:
                    val = ziel[key]
                    # Erreicht: Zahl zu Text
                    if key == "erreicht":
                        erreicht_map = {1: "nicht erreicht", 2: "teilweise erreicht", 3: "erreicht"}
                        val = erreicht_map.get(val, str(val))
                    row_data[key] = str(val) if val else ""
            data_rows.append(row_data)

        return self.fill_table_rows(table, data_rows, column_mapping)

    def save(self, doc: Document, output_path: str) -> Path:
        """Speichert das Dokument."""
        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(path))
        return path


# ═══════════════════════════════════════════════════════════════
# CLI Interface
# ═══════════════════════════════════════════════════════════════

def main():
    """Minimales CLI fuer Tests."""
    import sys

    if len(sys.argv) < 2:
        print("Usage: python word_template_service.py <template.docx> [--info]")
        print("  --info    Zeigt Dokumentstruktur (Tabellen, Platzhalter)")
        return

    template = sys.argv[1]
    svc = WordTemplateService()

    try:
        doc = svc.load_template(template)
    except FileNotFoundError as e:
        print(f"FEHLER: {e}")
        return

    if "--info" in sys.argv:
        print(f"Dokument: {template}")
        print(f"Paragraphen: {len(doc.paragraphs)}")
        print(f"Tabellen: {len(doc.tables)}")
        for i, table in enumerate(doc.tables):
            print(f"\n  Tabelle {i + 1}: {len(table.rows)} Zeilen x {len(table.columns)} Spalten")
            if table.rows:
                header = [_get_cell_text(c)[:40] for c in table.rows[0].cells]
                print(f"    Header: {header}")
    else:
        print(f"Geladen: {template}")
        print(f"  {len(doc.paragraphs)} Paragraphen, {len(doc.tables)} Tabellen")


if __name__ == "__main__":
    main()
