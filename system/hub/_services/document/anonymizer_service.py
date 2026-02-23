#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# SPDX-License-Identifier: MIT
"""
Copyright (c) 2026 Lukas Geiger

Permission is hereby granted, free of charge, to any person obtaining a copy
of this software and associated documentation files (the "Software"), to deal
in the Software without restriction, including without limitation the rights
to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
copies of the Software, and to permit persons to whom the Software is
furnished to do so, subject to the following conditions:

The above copyright notice and this permission notice shall be included in all
copies or substantial portions of the Software.

THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
SOFTWARE.
"""

"""
Document Anonymizer Service
============================

Pseudonymisierung und De-Anonymisierung von Dokumenten.
Ersetzt personenbezogene Daten durch konsistente Tarnnamen.
Schluessel wird AES-256 verschluesselt gespeichert.

Nutzt bestehende DokuZentrum-Komponenten:
  - RedactionDetector (core/redaction/detector.py)
  - OCREngine (core/ocr/engine.py)
  - PDFProcessor (core/pdf/processor.py)

Abhaengigkeiten:
  pip install python-docx cryptography openpyxl

Version: 1.2.0 (Excel-Support + Dateinamen-Anonymisierung)
Erstellt: 2026-01-27
"""

import json
import os
import re
import secrets
import string
from dataclasses import dataclass, field
from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict, List, Optional, Tuple

# AES-Verschluesselung
try:
    from cryptography.fernet import Fernet
    from cryptography.hazmat.primitives import hashes
    from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
    import base64
    CRYPTO_AVAILABLE = True
except ImportError:
    CRYPTO_AVAILABLE = False

# Word-Dokumente
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# PDF-Verarbeitung (PyMuPDF)
try:
    import fitz
    FITZ_AVAILABLE = True
except ImportError:
    FITZ_AVAILABLE = False

# PDF-Verschluesselung (pikepdf)
try:
    import pikepdf
    PIKEPDF_AVAILABLE = True
except ImportError:
    PIKEPDF_AVAILABLE = False

# Excel-Verarbeitung (openpyxl)
try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False


# ═══════════════════════════════════════════════════════════════
# Datenklassen
# ═══════════════════════════════════════════════════════════════

@dataclass
class AnonymProfile:
    """Anonymisierungsprofil fuer einen Klienten."""
    client_id: str                          # z.B. "K_0042"
    tarnname: str                           # z.B. "Felix Bergmann"
    fake_geburtsdatum: str                  # z.B. "22.07.2016" (gleiches Alter)
    mappings: Dict[str, Dict[str, str]] = field(default_factory=dict)
    # mappings = {
    #   "names": {"Max Mustermann": "Felix Bergmann", "Dr. Meyer": "Dr. Lindner"},
    #   "dates": {"15.03.2016": "22.07.2016"},
    #   "addresses": {"Musterstr. 5": "Waldweg 12"},
    #   "misc": {"07761/123456": "07741/987654"}
    # }
    created: str = ""
    version: int = 1


@dataclass
class AnonymResult:
    """Ergebnis einer Anonymisierung."""
    processed_files: int = 0
    anonymized_files: int = 0
    skipped_files: int = 0
    errors: List[str] = field(default_factory=list)
    replacements_total: int = 0


@dataclass
class ProgressInfo:
    """Fortschrittsinfo fuer GUI."""
    total_files: int = 0
    processed_files: int = 0
    current_file: str = ""
    status: str = "idle"  # idle, scanning, anonymizing, done, error

    @property
    def percent(self) -> float:
        if self.total_files == 0:
            return 0.0
        return (self.processed_files / self.total_files) * 100.0


# ═══════════════════════════════════════════════════════════════
# Tarnnamen-Generator
# ═══════════════════════════════════════════════════════════════

# Deutsche Phantasienamen fuer konsistente Pseudonymisierung
# Getrennt nach Geschlecht fuer konsistente Pronomen
TARN_VORNAMEN_M = [
    "Felix", "Jonas", "Paul", "Leon", "Finn", "Lukas", "Noah", "Elias",
    "Ben", "Maximilian", "Tim", "David", "Moritz", "Julian", "Niklas", "Erik"
]

TARN_VORNAMEN_W = [
    "Lena", "Marie", "Sophie", "Emma", "Hannah", "Mia", "Clara", "Lina",
    "Lea", "Anna", "Laura", "Sarah", "Julia", "Lisa", "Emily", "Nina"
]

# Kombinierte Liste fuer Rueckwaertskompatibilitaet
TARN_VORNAMEN = TARN_VORNAMEN_M + TARN_VORNAMEN_W

# Haeufige deutsche Vornamen fuer Gender-Erkennung
DEUTSCHE_VORNAMEN_M = {
    "jaden", "max", "paul", "leon", "jonas", "felix", "lukas", "elias", "ben", "noah",
    "tim", "david", "jan", "finn", "niklas", "moritz", "julian", "erik", "tom", "luca",
    "michael", "thomas", "peter", "hans", "klaus", "christian", "andreas", "stefan",
    "markus", "daniel", "tobias", "sebastian", "florian", "matthias", "alexander",
    "johannes", "philipp", "simon", "marco", "oliver", "martin", "frank", "wolfgang",
    "karl", "josef", "heinrich", "werner", "günter", "jürgen", "horst", "dieter",
    "harald", "helmut", "manfred", "bernhard", "gerhard", "rainer", "rolf", "walter",
    # Zusätzliche Namen für bessere Erkennung
    "skaven", "kevin", "justin", "jason", "brandon", "tyler", "ryan", "kyle",
    "pascal", "dennis", "sven", "lars", "nils", "björn", "jens", "uwe", "kai"
}

DEUTSCHE_VORNAMEN_W = {
    "marie", "sophie", "emma", "lena", "hannah", "mia", "clara", "lina", "lea", "anna",
    "laura", "sarah", "julia", "lisa", "emily", "nina", "amelie", "leonie", "johanna",
    "maria", "sabine", "petra", "susanne", "monika", "karin", "ursula", "renate",
    "brigitte", "helga", "ingrid", "gisela", "erika", "christa", "hildegard", "gertrud",
    "elisabeth", "christine", "andrea", "claudia", "martina", "nicole", "katrin",
    "birgit", "silke", "heike", "anja", "melanie", "stefanie", "sandra", "jennifer"
}

TARN_NACHNAMEN = [
    "Bergmann", "Fischer", "Lindner", "Sommer", "Richter", "Vogel",
    "Baumann", "Krause", "Werner", "Hartmann", "Lehmann", "Brandt",
    "Keller", "Bauer", "Schuster", "Hofmann", "Albrecht", "Steiner"
]

TARN_STRASSEN = [
    "Waldweg", "Birkenallee", "Sonnenstr.", "Gartenweg", "Lindenstr.",
    "Bergstr.", "Rosenweg", "Eichenstr.", "Parkstr.", "Wiesenweg"
]

TARN_STAEDTE = [
    "79800 Tiengen", "79761 Waldshut", "79725 Laufenburg",
    "79780 Stühlingen", "79807 Lottstetten", "79737 Herrischried"
]


def _detect_gender(vorname: str) -> str:
    """
    Erkennt das Geschlecht anhand des Vornamens.

    Returns:
        'm' fuer maennlich, 'w' fuer weiblich, 'u' fuer unbekannt
    """
    vorname_lower = vorname.lower().strip()

    if vorname_lower in DEUTSCHE_VORNAMEN_M:
        return 'm'
    if vorname_lower in DEUTSCHE_VORNAMEN_W:
        return 'w'

    # Heuristiken fuer unbekannte Namen
    # Viele weibliche Namen enden auf -a, -e, -ie, -ine
    if vorname_lower.endswith(('a', 'ie', 'ine', 'ette', 'ella', 'ina')):
        return 'w'
    # Viele maennliche Namen enden auf Konsonanten oder -o, -us, -er
    if vorname_lower.endswith(('us', 'er', 'o', 'ian', 'en')):
        return 'm'

    return 'u'


def _generate_tarnname(used_names: set = None, gender: str = None, original_vorname: str = None) -> str:
    """
    Generiert einen zufaelligen Tarnnamen.

    Args:
        used_names: Set bereits verwendeter Namen (Kollisionsvermeidung)
        gender: 'm' fuer maennlich, 'w' fuer weiblich, None fuer auto-detect
        original_vorname: Echter Vorname fuer Gender-Erkennung

    Returns:
        Tarnname im Format "Vorname Nachname"
    """
    if used_names is None:
        used_names = set()

    # Gender automatisch erkennen wenn nicht angegeben
    if gender is None and original_vorname:
        gender = _detect_gender(original_vorname)

    # Passende Vornamenliste waehlen
    if gender == 'm':
        vornamen_pool = TARN_VORNAMEN_M
    elif gender == 'w':
        vornamen_pool = TARN_VORNAMEN_W
    else:
        vornamen_pool = TARN_VORNAMEN  # Fallback: alle

    for _ in range(100):
        name = f"{secrets.choice(vornamen_pool)} {secrets.choice(TARN_NACHNAMEN)}"
        if name not in used_names:
            return name
    return f"Person_{secrets.token_hex(4)}"


def _shift_date(date_str: str, days_offset: int) -> str:
    """
    Verschiebt ein Datum um eine feste Anzahl Tage.
    Akzeptiert dd.mm.yyyy Format.
    """
    try:
        dt = datetime.strptime(date_str.strip(), "%d.%m.%Y")
        shifted = dt + timedelta(days=days_offset)
        return shifted.strftime("%d.%m.%Y")
    except ValueError:
        return date_str


def _generate_fake_date_same_age(real_date: str) -> Tuple[str, int]:
    """
    Generiert ein falsches Geburtsdatum mit demselben Alter.
    Returns: (fake_date, days_offset)
    """
    offset = secrets.randbelow(300) - 150  # -150 bis +149 Tage
    fake = _shift_date(real_date, offset)
    return fake, offset


def _generate_fake_phone() -> str:
    """Generiert eine falsche Telefonnummer."""
    prefix = secrets.choice(["07741", "07742", "07743", "07744"])
    number = "".join([str(secrets.randbelow(10)) for _ in range(6)])
    return f"{prefix}/{number}"


def _generate_fake_email(original: str = None) -> str:
    """Generiert eine falsche E-Mail-Adresse."""
    vorname = secrets.choice(TARN_VORNAMEN).lower()
    nachname = secrets.choice(TARN_NACHNAMEN).lower()
    domain = secrets.choice(["beispiel.de", "muster.org", "test-mail.de", "privat.net"])
    return f"{vorname}.{nachname}@{domain}"


def _generate_fake_address() -> str:
    """Generiert eine falsche Adresse."""
    strasse = secrets.choice(TARN_STRASSEN)
    nr = secrets.randbelow(50) + 1
    stadt = secrets.choice(TARN_STAEDTE)
    return f"{strasse} {nr}, {stadt}"


# ═══════════════════════════════════════════════════════════════
# RegEx-Patterns fuer automatische Erkennung
# ═══════════════════════════════════════════════════════════════

# Telefonnummern (deutsch): 07761/123456, +49 7761 123456, 0049-7761-123456, 0761 12345678
PHONE_PATTERN = re.compile(
    r'(?:'
    r'(?:\+49|0049)[\s\-/]?\d{2,4}[\s\-/]?\d{4,8}|'  # +49 oder 0049
    r'0\d{2,4}[\s\-/]?\d{4,8}'                        # 0xxx/xxxxxxx
    r')',
    re.IGNORECASE
)

# E-Mail-Adressen
EMAIL_PATTERN = re.compile(
    r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}',
    re.IGNORECASE
)

# Straßenadressen: "Musterstr. 5", "Hauptstraße 123a", "Am Waldweg 7 b"
STREET_PATTERN = re.compile(
    r'(?:'
    r'(?:[A-ZÄÖÜ][a-zäöüß]+(?:str\.|straße|weg|gasse|platz|allee|ring|damm|ufer|berg|tal|hof|feld|wiese|grund|rain|steig|pfad))'  # Straßenname
    r'\s*'
    r'\d{1,4}\s?[a-zA-Z]?'  # Hausnummer + opt. Zusatz
    r')',
    re.IGNORECASE
)


# ═══════════════════════════════════════════════════════════════
# Lokaler Schluessel-Speicher (NICHT in OneDrive)
# ═══════════════════════════════════════════════════════════════

def get_local_keys_dir() -> Path:
    """
    Gibt den lokalen Schluessel-Ordner zurueck (ausserhalb OneDrive).

    Schluessel-Dateien (.schluessel.enc) duerfen NICHT ueber Cloud-Dienste
    synchronisiert werden. Dieser Ordner liegt unter %LOCALAPPDATA%\\BACH\\keys\\
    und wird bei Bedarf automatisch erstellt.

    Returns:
        Path zum lokalen keys-Ordner
    """
    if os.name == "nt":
        base = Path(os.environ.get("LOCALAPPDATA", os.path.expanduser("~")))
    else:
        base = Path.home() / ".local" / "share"
    keys_dir = base / "BACH" / "keys"
    keys_dir.mkdir(parents=True, exist_ok=True)
    return keys_dir


def get_key_path(client_id: str) -> Path:
    """
    Gibt den sicheren Pfad fuer eine Schluessel-Datei zurueck.

    Args:
        client_id: Klienten-ID (z.B. 'K_0042')

    Returns:
        Path zur .schluessel.enc Datei im lokalen Speicher
    """
    return get_local_keys_dir() / f"{client_id}.schluessel.enc"


def _clear_hidden_attribute(filepath: Path):
    """Entfernt das Windows Hidden-Attribut (noetig vor Ueberschreiben)."""
    if os.name == "nt" and filepath.exists():
        try:
            import ctypes
            attrs = ctypes.windll.kernel32.GetFileAttributesW(str(filepath))
            if attrs != -1 and (attrs & 0x02):
                ctypes.windll.kernel32.SetFileAttributesW(str(filepath), attrs & ~0x02)
        except Exception:
            pass


def _set_hidden_attribute(filepath: Path):
    """Setzt das Windows Hidden-Attribut auf eine Datei (nur Windows)."""
    if os.name == "nt":
        try:
            import ctypes
            FILE_ATTRIBUTE_HIDDEN = 0x02
            ctypes.windll.kernel32.SetFileAttributesW(str(filepath), FILE_ATTRIBUTE_HIDDEN)
        except Exception:
            pass  # Nicht kritisch


# ═══════════════════════════════════════════════════════════════
# AES-Verschluesselung fuer Schluessel
# ═══════════════════════════════════════════════════════════════

def _derive_key(password: str, salt: bytes) -> bytes:
    """Leitet AES-Schluessel aus Passwort ab (PBKDF2)."""
    if not CRYPTO_AVAILABLE:
        raise RuntimeError("cryptography-Paket nicht installiert: pip install cryptography")

    kdf = PBKDF2HMAC(
        algorithm=hashes.SHA256(),
        length=32,
        salt=salt,
        iterations=480000,
    )
    return base64.urlsafe_b64encode(kdf.derive(password.encode("utf-8")))


def encrypt_key_file(profile: AnonymProfile, output_path: str, password: str) -> Path:
    """
    Speichert das AnonymProfile AES-256 verschluesselt.

    Args:
        profile: Das Anonymisierungsprofil mit allen Mappings
        output_path: Pfad fuer die verschluesselte Datei (.schluessel.enc)
        password: Passwort fuer die Verschluesselung
    """
    if not CRYPTO_AVAILABLE:
        raise RuntimeError("cryptography-Paket nicht installiert: pip install cryptography")

    # Profil als JSON serialisieren
    data = {
        "version": profile.version,
        "client_id": profile.client_id,
        "tarnname": profile.tarnname,
        "fake_geburtsdatum": profile.fake_geburtsdatum,
        "mappings": profile.mappings,
        "created": profile.created or datetime.now().isoformat()
    }
    plaintext = json.dumps(data, ensure_ascii=False, indent=2).encode("utf-8")

    # Verschluesseln
    salt = os.urandom(16)
    key = _derive_key(password, salt)
    fernet = Fernet(key)
    encrypted = fernet.encrypt(plaintext)

    # Speichern: Salt (16 Bytes fest) + verschluesselte Daten
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    _clear_hidden_attribute(path)  # Falls Datei existiert und Hidden ist
    with open(path, "wb") as f:
        f.write(salt)  # genau 16 Bytes
        f.write(encrypted)

    # Hidden-Attribut setzen (Schutz vor versehentlichem Zugriff)
    _set_hidden_attribute(path)

    return path


def decrypt_key_file(key_path: str, password: str) -> AnonymProfile:
    """
    Laedt und entschluesselt ein AnonymProfile.

    Args:
        key_path: Pfad zur .schluessel.enc Datei
        password: Passwort fuer die Entschluesselung
    """
    if not CRYPTO_AVAILABLE:
        raise RuntimeError("cryptography-Paket nicht installiert: pip install cryptography")

    path = Path(key_path)
    with open(path, "rb") as f:
        content = f.read()

    # Salt (erste 16 Bytes) und verschluesselte Daten trennen
    salt = content[:16]
    encrypted = content[16:]

    # Entschluesseln
    key = _derive_key(password, salt)
    fernet = Fernet(key)
    plaintext = fernet.decrypt(encrypted)

    # JSON parsen
    data = json.loads(plaintext.decode("utf-8"))

    return AnonymProfile(
        client_id=data["client_id"],
        tarnname=data["tarnname"],
        fake_geburtsdatum=data["fake_geburtsdatum"],
        mappings=data.get("mappings", {}),
        created=data.get("created", ""),
        version=data.get("version", 1)
    )


# ═══════════════════════════════════════════════════════════════
# Anonymisierer
# ═══════════════════════════════════════════════════════════════

class DocumentAnonymizer:
    """
    Anonymisiert Dokumente durch konsistente Pseudonymisierung.

    Workflow:
        1. create_profile() - Profil mit Tarnnamen erstellen
        2. anonymize_folder() - Alle Dokumente im Ordner anonymisieren
        3. Schluessel wird AES-verschluesselt gespeichert
    """

    def __init__(self):
        self._progress = ProgressInfo()
        self._used_names: set = set()
        self.global_whitelist = self._load_global_whitelist()

    def _load_global_whitelist(self) -> dict:
        """Laedt die globale Whitelist aus data/anonymizer_whitelist.json."""
        whitelist_file = Path(__file__).parent.parent.parent.parent / "data" / "anonymizer_whitelist.json"
        if whitelist_file.exists():
            try:
                return json.loads(whitelist_file.read_text(encoding="utf-8"))
            except Exception as e:
                print(f"[WARN] Anonymizer Whitelist konnte nicht geladen werden: {e}")
        return {"titles": [], "names": [], "organizations": []}

    @property
    def progress(self) -> ProgressInfo:
        return self._progress

    def scan_text_for_sensitive_data(self, text: str) -> Dict[str, List[str]]:
        """
        Scannt Text nach sensiblen Daten (Telefon, E-Mail, Adressen).

        Args:
            text: Der zu scannende Text

        Returns:
            Dict mit Listen: {"phones": [...], "emails": [...], "addresses": [...]}
        """
        found = {
            "phones": [],
            "emails": [],
            "addresses": []
        }

        # Telefonnummern finden
        phones = PHONE_PATTERN.findall(text)
        for phone in phones:
            cleaned = phone.strip()
            if cleaned and cleaned not in found["phones"]:
                found["phones"].append(cleaned)

        # E-Mails finden
        emails = EMAIL_PATTERN.findall(text)
        for email in emails:
            cleaned = email.strip().lower()
            if cleaned and cleaned not in found["emails"]:
                found["emails"].append(cleaned)

        # Adressen finden
        addresses = STREET_PATTERN.findall(text)
        for addr in addresses:
            cleaned = addr.strip()
            if cleaned and cleaned not in found["addresses"]:
                found["addresses"].append(cleaned)

        return found

    def create_profile(
        self,
        real_name: str,
        geburtsdatum: str,
        weitere_namen: Optional[List[str]] = None,
        weitere_daten: Optional[Dict[str, str]] = None,
        whitelist: Optional[List[str]] = None,
        scanned_data: Optional[Dict[str, List[str]]] = None
    ) -> AnonymProfile:
        """
        Erstellt ein Anonymisierungsprofil.

        Args:
            real_name: Echter Name des Klienten (z.B. "Max Mustermann")
            geburtsdatum: Echtes Geburtsdatum (dd.mm.yyyy)
            weitere_namen: Weitere zu ersetzende Namen (Aerzte, Angehoerige)
            weitere_daten: Weitere Daten {"adresse": "Musterstr. 5", "telefon": "07761/123"}
            whitelist: Namen die NICHT anonymisiert werden (z.B. Sachbearbeiter vom Amt)
            scanned_data: Automatisch erkannte Daten aus scan_text_for_sensitive_data()
                          {"phones": [...], "emails": [...], "addresses": [...]}
        """
        # Whitelist zusammenstellen (Global + Parameter)
        whitelist_set = set(whitelist) if whitelist else set()
        whitelist_set.update(self.global_whitelist.get("names", []))
        whitelist_set.update(self.global_whitelist.get("organizations", []))

        # Titel fuer automatische Erkennung
        titles = self.global_whitelist.get("titles", [])
        # Client-ID generieren
        client_id = f"K_{secrets.token_hex(3).upper()}"

        # Vorname extrahieren für Gender-Erkennung
        real_parts = real_name.strip().split()
        original_vorname = real_parts[0] if real_parts else ""

        # Tarnname mit Geschlechts-Matching generieren
        tarnname = _generate_tarnname(
            used_names=self._used_names,
            original_vorname=original_vorname
        )
        self._used_names.add(tarnname)

        # Falsches Geburtsdatum (gleiches Alter)
        fake_geb, date_offset = _generate_fake_date_same_age(geburtsdatum)

        # Mappings aufbauen
        mappings: Dict[str, Dict[str, str]] = {
            "names": {},
            "dates": {},
            "addresses": {},
            "phones": {},
            "emails": {},
            "misc": {}
        }

        # Hauptname
        real_parts = real_name.strip().split()
        tarn_parts = tarnname.strip().split()

        mappings["names"][real_name] = tarnname
        # Auch Vorname und Nachname einzeln
        if len(real_parts) >= 2 and len(tarn_parts) >= 2:
            mappings["names"][real_parts[0]] = tarn_parts[0]   # Vorname
            mappings["names"][real_parts[-1]] = tarn_parts[-1]  # Nachname

        # Weitere Namen (nur wenn nicht auf Whitelist und keine Amtsperson)
        if weitere_namen:
            for name in weitere_namen:
                # Check ob explizit whitelisted
                if name in whitelist_set:
                    continue

                # Check ob Amtsperson durch Titel (z.B. "Dr. Meyer" -> "Dr." ist Title)
                is_amtsperson = any(title in name for title in titles)
                if is_amtsperson:
                    continue

                fake = _generate_tarnname(self._used_names)
                self._used_names.add(fake)
                mappings["names"][name] = fake

        # Geburtsdatum
        mappings["dates"][geburtsdatum] = fake_geb

        # Weitere Daten (explizit uebergeben)
        if weitere_daten:
            if "adresse" in weitere_daten:
                mappings["addresses"][weitere_daten["adresse"]] = _generate_fake_address()
            if "telefon" in weitere_daten:
                mappings["phones"][weitere_daten["telefon"]] = _generate_fake_phone()
            if "email" in weitere_daten:
                mappings["emails"][weitere_daten["email"]] = _generate_fake_email()
            # Restliche als misc
            for key, val in weitere_daten.items():
                if key not in ("adresse", "telefon", "email") and val:
                    mappings["misc"][val] = f"[{key.upper()}_ANON]"

        # Automatisch gescannte Daten (Telefon, E-Mail, Adressen)
        if scanned_data:
            # Telefonnummern
            for phone in scanned_data.get("phones", []):
                if phone not in mappings["phones"]:
                    mappings["phones"][phone] = _generate_fake_phone()

            # E-Mail-Adressen
            for email in scanned_data.get("emails", []):
                if email not in mappings["emails"]:
                    mappings["emails"][email] = _generate_fake_email(email)

            # Straßenadressen
            for addr in scanned_data.get("addresses", []):
                if addr not in mappings["addresses"]:
                    mappings["addresses"][addr] = _generate_fake_address()

        return AnonymProfile(
            client_id=client_id,
            tarnname=tarnname,
            fake_geburtsdatum=fake_geb,
            mappings=mappings,
            created=datetime.now().isoformat()
        )

    def extract_text_from_file(self, filepath: str) -> str:
        """
        Extrahiert Text aus einer Datei zum Scannen.

        Args:
            filepath: Pfad zur Datei

        Returns:
            Extrahierter Text
        """
        path = Path(filepath)
        suffix = path.suffix.lower()
        text = ""

        try:
            if suffix == ".docx" and DOCX_AVAILABLE:
                doc = Document(str(path))
                paragraphs = [p.text for p in doc.paragraphs]
                # Auch Tabellen
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            paragraphs.append(cell.text)
                text = "\n".join(paragraphs)

            elif suffix in (".txt", ".md"):
                try:
                    text = path.read_text(encoding="utf-8")
                except UnicodeDecodeError:
                    text = path.read_text(encoding="latin-1")

            elif suffix == ".pdf" and FITZ_AVAILABLE:
                doc = fitz.open(str(path))
                for page in doc:
                    text += page.get_text()
                doc.close()

            elif suffix in (".xlsx", ".xls") and EXCEL_AVAILABLE:
                wb = openpyxl.load_workbook(str(path), data_only=True)
                for sheet in wb.worksheets:
                    for row in sheet.iter_rows():
                        for cell in row:
                            if cell.value:
                                text += str(cell.value) + " "
                wb.close()
        except Exception as e:
            print(f"[WARN] Text-Extraktion fehlgeschlagen fuer {path.name}: {e}")

        return text

    def scan_folder_for_sensitive_data(self, folder: str) -> Dict[str, List[str]]:
        """
        Scannt alle Dateien in einem Ordner nach sensiblen Daten.

        Args:
            folder: Pfad zum Ordner

        Returns:
            Aggregierte gefundene Daten {"phones": [...], "emails": [...], "addresses": [...]}
        """
        combined = {
            "phones": [],
            "emails": [],
            "addresses": []
        }

        src = Path(folder)
        supported = {".docx", ".txt", ".md", ".pdf", ".xlsx", ".xls"}

        for filepath in src.rglob("*"):
            if filepath.is_file() and filepath.suffix.lower() in supported:
                text = self.extract_text_from_file(str(filepath))
                found = self.scan_text_for_sensitive_data(text)

                for key in combined:
                    for item in found.get(key, []):
                        if item not in combined[key]:
                            combined[key].append(item)

        return combined

    def anonymize_file(self, filepath: str, profile: AnonymProfile) -> Tuple[bool, int]:
        """
        Anonymisiert eine einzelne Datei.

        Returns:
            (success, replacement_count)
        """
        path = Path(filepath)
        suffix = path.suffix.lower()

        if suffix == ".docx":
            return self._anonymize_docx(path, profile)
        elif suffix == ".txt" or suffix == ".md":
            return self._anonymize_text(path, profile)
        elif suffix == ".pdf":
            return self._anonymize_pdf(path, profile)
        elif suffix in (".xlsx", ".xls"):
            return self._anonymize_excel(path, profile)
        else:
            return False, 0

    def anonymize_folder(
        self,
        folder: str,
        profile: AnonymProfile,
        password: str,
        output_folder: Optional[str] = None
    ) -> AnonymResult:
        """
        Anonymisiert alle Dokumente in einem Ordner.

        Args:
            folder: Quellordner
            profile: Anonymisierungsprofil
            password: Passwort fuer den Schluessel
            output_folder: Zielordner (default: klienten/<client_id>/)
        """
        result = AnonymResult()
        src = Path(folder)

        if output_folder:
            dest = Path(output_folder)
        else:
            try:
                from bach_paths import KLIENTEN_DIR
                dest = KLIENTEN_DIR / profile.client_id
            except ImportError:
                dest = src.parent.parent / "klienten" / profile.client_id

        dest.mkdir(parents=True, exist_ok=True)
        (dest / "output").mkdir(exist_ok=True)

        # Dateien zaehlen
        files = [f for f in src.rglob("*") if f.is_file() and not f.name.startswith(".")]
        self._progress = ProgressInfo(total_files=len(files), status="anonymizing")

        # Dateien verarbeiten
        for filepath in files:
            self._progress.current_file = filepath.name
            self._progress.processed_files += 1
            result.processed_files += 1

            try:
                # Relative Struktur beibehalten
                rel = filepath.relative_to(src)
                dest_file = dest / rel
                dest_file.parent.mkdir(parents=True, exist_ok=True)

                # Datei kopieren und anonymisieren
                import shutil
                shutil.copy2(filepath, dest_file)

                success, count = self.anonymize_file(str(dest_file), profile)
                if success:
                    result.anonymized_files += 1
                    result.replacements_total += count
                else:
                    result.skipped_files += 1

                # Dateinamen anonymisieren (falls Name enthalten)
                dest_file = self._anonymize_filename(dest_file, profile)

            except Exception as e:
                result.errors.append(f"{filepath.name}: {e}")

        # Schluessel im lokalen Speicher speichern (NICHT in OneDrive)
        key_path = get_key_path(profile.client_id)
        encrypt_key_file(profile, str(key_path), password)

        # Profil-Info (ohne sensible Daten) speichern
        profil_info = {
            "client_id": profile.client_id,
            "tarnname": profile.tarnname,
            "fake_geburtsdatum": profile.fake_geburtsdatum,
            "created": profile.created,
            "files_anonymized": result.anonymized_files,
            "key_location": str(key_path),
            "key_info": "Schluessel liegt LOKAL (nicht in OneDrive)"
        }
        profil_path = dest / ".profil.json"
        profil_path.write_text(json.dumps(profil_info, indent=2, ensure_ascii=False), encoding="utf-8")

        self._progress.status = "done"
        return result

    def _anonymize_docx(self, path: Path, profile: AnonymProfile) -> Tuple[bool, int]:
        """Anonymisiert ein Word-Dokument."""
        if not DOCX_AVAILABLE:
            return False, 0

        doc = Document(str(path))
        count = 0

        # Alle Mappings anwenden
        all_replacements = {}
        for category in profile.mappings.values():
            all_replacements.update(category)

        # Nach Laenge sortieren (laengste zuerst, verhindert Teilersetzungen)
        sorted_replacements = sorted(all_replacements.items(), key=lambda x: len(x[0]), reverse=True)

        def replace_in_paragraphs(paragraphs):
            nonlocal count
            for paragraph in paragraphs:
                full_text = "".join(run.text for run in paragraph.runs)
                
                # Check ob ueberhaupt was zu tun ist
                if not any(old in full_text for old, _ in sorted_replacements):
                    continue

                for old, new in sorted_replacements:
                    if old in full_text:
                        # Versuch 1: Einzel-Run Ersetzung
                        replaced_in_runs = False
                        for run in paragraph.runs:
                            if old in run.text:
                                run.text = run.text.replace(old, new)
                                count += 1
                                replaced_in_runs = True
                        
                        # Versuch 2: Falls gesplittet (destruktiver Fallback fuer Privacy)
                        if not replaced_in_runs:
                            p_text = paragraph.text
                            new_p_text = p_text.replace(old, new)
                            if new_p_text != p_text:
                                # Alle Runs leeren
                                for run in paragraph.runs:
                                    run.text = ""
                                # Text in den ersten Run schreiben
                                if paragraph.runs:
                                    paragraph.runs[0].text = new_p_text
                                else:
                                    paragraph.add_run(new_p_text)
                                count += 1
                                full_text = new_p_text # Update fuer naechste Ersetzung

        # Paragraphen
        replace_in_paragraphs(doc.paragraphs)

        # Tabellen
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_paragraphs(cell.paragraphs)

        # Header/Footer
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    replace_in_paragraphs(header.paragraphs)
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    replace_in_paragraphs(footer.paragraphs)

        doc.save(str(path))
        return True, count

    def _anonymize_text(self, path: Path, profile: AnonymProfile) -> Tuple[bool, int]:
        """Anonymisiert eine Textdatei."""
        try:
            text = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            text = path.read_text(encoding="latin-1")

        count = 0
        all_replacements = {}
        for category in profile.mappings.values():
            all_replacements.update(category)

        sorted_replacements = sorted(all_replacements.items(), key=lambda x: len(x[0]), reverse=True)

        for old, new in sorted_replacements:
            occurrences = text.count(old)
            if occurrences > 0:
                text = text.replace(old, new)
                count += occurrences

        path.write_text(text, encoding="utf-8")
        return True, count

    def _anonymize_excel(self, path: Path, profile: AnonymProfile) -> Tuple[bool, int]:
        """
        Anonymisiert eine Excel-Datei (.xlsx, .xls).

        Ersetzt sensible Begriffe in allen Zellen aller Tabellenblätter.
        """
        if not EXCEL_AVAILABLE:
            return False, 0

        try:
            wb = openpyxl.load_workbook(str(path))
        except Exception:
            return False, 0

        count = 0
        all_replacements = {}
        for category in profile.mappings.values():
            all_replacements.update(category)

        # Nach Laenge sortieren (laengste zuerst, verhindert Teilersetzungen)
        sorted_replacements = sorted(all_replacements.items(), key=lambda x: len(x[0]), reverse=True)

        # Alle Tabellenblätter durchgehen
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        original = cell.value
                        new_value = original
                        for old, new in sorted_replacements:
                            if old in new_value:
                                new_value = new_value.replace(old, new)
                                count += 1
                        if new_value != original:
                            cell.value = new_value

        wb.save(str(path))
        wb.close()
        return True, count

    def _anonymize_filename(self, filepath: Path, profile: AnonymProfile) -> Path:
        """
        Anonymisiert einen Dateinamen, falls er sensible Begriffe enthält.

        Returns:
            Neuer Pfad (umbenannt) oder urspruenglicher Pfad (unveraendert)
        """
        filename = filepath.stem
        suffix = filepath.suffix

        all_replacements = {}
        for category in profile.mappings.values():
            all_replacements.update(category)

        sorted_replacements = sorted(all_replacements.items(), key=lambda x: len(x[0]), reverse=True)

        new_filename = filename
        changed = False
        for old, new in sorted_replacements:
            if old in new_filename:
                new_filename = new_filename.replace(old, new)
                changed = True

        if changed:
            new_path = filepath.parent / f"{new_filename}{suffix}"
            filepath.rename(new_path)
            return new_path
        return filepath

    def _anonymize_pdf(self, path: Path, profile: AnonymProfile,
                       encrypt_password: Optional[str] = None) -> Tuple[bool, int]:
        """
        PDF-Anonymisierung via PyMuPDF (fitz) Schwärzung + pikepdf AES-256 Verschluesselung.

        Hybrid-Ansatz aus:
          - DokuZentrum RedactionDetector (Erkennungslogik)
          - PDFSchwaerzer Pro (Redact+Encrypt Pipeline)

        Pipeline:
          1. PDF oeffnen (fitz)
          2. Fuer jede Seite: Sensitive Begriffe suchen und schwärzen
          3. Geschwärztes PDF speichern
          4. Optional: AES-256 verschluesseln (pikepdf, R=6)

        Args:
            path: PDF-Datei
            profile: Anonymisierungsprofil mit Mappings
            encrypt_password: Optionales Passwort fuer PDF-Verschluesselung
        """
        if not FITZ_AVAILABLE:
            return False, 0

        # Alle zu schwärzenden Begriffe sammeln (Original-Namen, nicht Tarnnamen)
        sensitive_words = []
        for category in profile.mappings.values():
            for original_text in category.keys():
                if len(original_text) >= 2:
                    sensitive_words.append(original_text)

        # Nach Laenge sortieren (laengste zuerst)
        sensitive_words.sort(key=len, reverse=True)

        if not sensitive_words:
            return True, 0

        count = 0
        import shutil
        temp_path = path.with_suffix(".tmp.pdf")

        try:
            doc = fitz.open(str(path))

            for page in doc:
                for word in sensitive_words:
                    hits = page.search_for(word)
                    for rect in hits:
                        page.add_redact_annot(rect, fill=(0, 0, 0))
                        count += 1
                page.apply_redactions()

            # Immer in Temp-Datei speichern (fitz verbietet non-incremental save zum Original)
            doc.save(str(temp_path))
            doc.close()

            # Optional: AES-256 Verschluesselung (pikepdf, R=6)
            if encrypt_password and PIKEPDF_AVAILABLE:
                try:
                    pdf = pikepdf.open(str(temp_path))
                    enc = pikepdf.Encryption(
                        owner=encrypt_password,
                        user=encrypt_password,
                        R=6
                    )
                    pdf.save(str(path), encryption=enc)
                    pdf.close()
                    temp_path.unlink()
                except Exception:
                    # Fallback: Unverschluesselt
                    if temp_path.exists():
                        shutil.move(str(temp_path), str(path))
            else:
                if temp_path.exists():
                    shutil.move(str(temp_path), str(path))

            return True, count

        except Exception:
            if temp_path.exists():
                temp_path.unlink()
            return False, 0


# ═══════════════════════════════════════════════════════════════
# De-Anonymisierer
# ═══════════════════════════════════════════════════════════════

class DocumentDeanonymizer:
    """
    Stellt anonymisierte Dokumente wieder her.
    Ersetzt Tarnnamen durch echte Namen basierend auf dem Schluessel.
    """

    def deanonymize_file(
        self,
        filepath: str,
        profile: AnonymProfile
    ) -> Tuple[bool, int]:
        """
        De-anonymisiert eine einzelne Datei (umgekehrte Mappings).
        """
        # Umgekehrte Mappings erstellen
        reverse_profile = AnonymProfile(
            client_id=profile.client_id,
            tarnname="",
            fake_geburtsdatum="",
            mappings={}
        )

        for category, mapping in profile.mappings.items():
            reverse_profile.mappings[category] = {v: k for k, v in mapping.items()}

        # Anonymizer mit umgekehrten Mappings nutzen
        anon = DocumentAnonymizer()
        return anon.anonymize_file(filepath, reverse_profile)

    def deanonymize_folder(
        self,
        folder: str,
        schluessel_path: str,
        password: str,
        output_folder: str,
        client_id: str = None
    ) -> AnonymResult:
        """
        De-anonymisiert alle Dokumente in einem Ordner.

        Args:
            folder: Anonymisierter Ordner (z.B. klienten/K_0042/)
            schluessel_path: Pfad zur .schluessel.enc Datei (oder None wenn client_id gegeben)
            password: Passwort fuer den Schluessel
            output_folder: Zielordner (z.B. _ready_for_export/Max_Mustermann/)
            client_id: Klienten-ID — wenn angegeben, wird der lokale Schluessel genutzt
        """
        # Schluessel laden
        if client_id and not schluessel_path:
            schluessel_path = str(get_key_path(client_id))
        profile = decrypt_key_file(schluessel_path, password)

        result = AnonymResult()
        src = Path(folder)
        dest = Path(output_folder)
        dest.mkdir(parents=True, exist_ok=True)

        import shutil

        # Dateien kopieren und de-anonymisieren
        deanon_suffixes = {".docx", ".txt", ".md"}
        copy_only_suffixes = {".pdf"}
        all_suffixes = deanon_suffixes | copy_only_suffixes

        files = [f for f in src.rglob("*")
                 if f.is_file()
                 and not f.name.startswith(".")
                 and f.suffix.lower() in all_suffixes]

        for filepath in files:
            try:
                rel = filepath.relative_to(src)
                dest_file = dest / rel
                dest_file.parent.mkdir(parents=True, exist_ok=True)
                shutil.copy2(filepath, dest_file)
                result.processed_files += 1

                # PDFs sind geschwärzt — keine De-Anonymisierung moeglich
                if filepath.suffix.lower() in copy_only_suffixes:
                    continue

                success, count = self.deanonymize_file(str(dest_file), profile)
                if success:
                    result.anonymized_files += 1
                    result.replacements_total += count

            except Exception as e:
                result.errors.append(f"{filepath.name}: {e}")

        return result


# ═══════════════════════════════════════════════════════════════
# CLI Interface
# ═══════════════════════════════════════════════════════════════

def main():
    """CLI Einstiegspunkt."""
    import sys

    if len(sys.argv) < 2:
        print("Document Anonymizer Service v1.0.0")
        print()
        print("Usage:")
        print("  python anonymizer_service.py create-profile <name> <geburtsdatum>")
        print("  python anonymizer_service.py anonymize <ordner> <passwort>")
        print("  python anonymizer_service.py deanonymize <ordner> <schluessel> <passwort> <output>")
        print("  python anonymizer_service.py test")
        return

    cmd = sys.argv[1]

    if cmd == "test":
        print("[TEST] Erstelle Testprofil...")
        anon = DocumentAnonymizer()
        profile = anon.create_profile(
            real_name="Max Mustermann",
            geburtsdatum="15.03.2016",
            weitere_namen=["Dr. Meyer", "Frau Schmidt"],
            weitere_daten={"adresse": "Musterstr. 5, 79713 Bad Säckingen", "telefon": "07761/123456"}
        )
        print(f"  Client-ID: {profile.client_id}")
        print(f"  Tarnname: {profile.tarnname}")
        print(f"  Fake-Geb.: {profile.fake_geburtsdatum}")
        print(f"  Mappings:")
        for cat, mapping in profile.mappings.items():
            print(f"    {cat}:")
            for k, v in mapping.items():
                print(f"      {k} -> {v}")

        # Verschluesselungstest
        if CRYPTO_AVAILABLE:
            test_path = Path("__test_schluessel.enc")
            encrypt_key_file(profile, str(test_path), "testpasswort123")
            print(f"\n  Schluessel gespeichert: {test_path}")

            loaded = decrypt_key_file(str(test_path), "testpasswort123")
            print(f"  Entschluesselt: {loaded.tarnname} (ID: {loaded.client_id})")
            assert loaded.tarnname == profile.tarnname
            assert loaded.mappings == profile.mappings
            print("  [OK] Verschluesselungstest bestanden")

            test_path.unlink()
        else:
            print("\n  [WARN] cryptography nicht installiert, Verschluesselungstest uebersprungen")

        print("\n[OK] Test abgeschlossen")

    else:
        print(f"Unbekannter Befehl: {cmd}")
        print("Nutze 'test' fuer einen Selbsttest")


if __name__ == "__main__":
    main()
