#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# SPDX-License-Identifier: MIT
"""
Copyright (c) 2026 Lukas Geiger

Permission is hereby granted, free of charge, to any person obtaining a copy
of this software and associated documentation files (the "Software"), to deal
in the Software without restriction, including without limitation the rights
to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
copies of the Software, and to permit persons to whom the Software is
furnished to do so, subject to the following conditions:

The above copyright notice and this permission notice shall be included in all
copies or substantial portions of the Software.

THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
SOFTWARE.
"""

"""
Document Collector Service
===========================

Sammelt und kategorisiert Dokumente nach dem CORE/STUFE2/EXTENDED Schema.
Extrahiert Text und erstellt anonymisierte Bundles fuer LLM-Verarbeitung.

Workflow:
  1. scan_folder() - Dokumente scannen und kategorisieren
  2. extract_bundle() - Text extrahieren (CORE + STUFE2)
  3. anonymize_bundle() - Text anonymisieren
  4. Das anonymisierte Bundle geht ans LLM

Die Original-PDFs werden NIE ans LLM geschickt - nur der extrahierte,
anonymisierte Text. Daher ist keine PDF-Schwaerzung noetig.

Kategorien:
  CORE: Protokolle, Aktendeckblatt, aktueller Hilfeplan, letzter Foerderstelle-Bericht
  STUFE2: Mails (Berichtszeitraum), aktuellster Arztbericht, aktuellster Schulbericht
  EXTENDED: Aeltere Dokumente (separate Datei, nur auf Anfrage)

Version: 1.0.1 (Bugfix: Datums-Regex & OCR-Pfad)
Erstellt: 2026-01-31
Aktualisiert: 2026-02-02
"""

import re
import sys
from dataclasses import dataclass, field
from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict, List, Optional, Tuple, TYPE_CHECKING
from enum import Enum

if TYPE_CHECKING:
    from .anonymizer_service import AnonymProfile


class DocumentCategory(Enum):
    """Dokument-Kategorien nach Prioritaet."""
    CORE = "core"           # Immer an LLM senden
    STUFE2 = "stufe2"       # Nach CORE senden
    EXTENDED = "extended"   # Nur auf Anfrage
    SKIP = "skip"           # Nicht verarbeiten


@dataclass
class DocumentInfo:
    """Informationen zu einem Dokument."""
    path: Path
    filename: str
    suffix: str
    category: DocumentCategory
    doc_type: str           # z.B. "protokoll", "hilfeplan", "arztbericht"
    date_hint: Optional[datetime] = None  # Erkanntes Datum aus Dateinamen
    size_bytes: int = 0
    text_extracted: bool = False
    text_length: int = 0


@dataclass
class CollectorResult:
    """Ergebnis der Dokumenten-Sammlung."""
    documents: List[DocumentInfo] = field(default_factory=list)
    core_count: int = 0
    stufe2_count: int = 0
    extended_count: int = 0
    skipped_count: int = 0
    errors: List[str] = field(default_factory=list)


@dataclass
class TextBundle:
    """Extrahiertes Text-Bundle."""
    core_text: str = ""
    stufe2_text: str = ""
    extended_text: str = ""
    total_length: int = 0
    files_processed: int = 0


# ═══════════════════════════════════════════════════════════════
# Erkennungsmuster fuer Dokumenttypen
# ═══════════════════════════════════════════════════════════════

# Muster fuer Dateinamen (case-insensitive)
DOC_PATTERNS = {
    # CORE Dokumente
    "protokoll": [
        r"protokoll",
        r"einzel.*protokoll",
        r"gruppen.*protokoll",
        r"doku(?:mentation)?",
        r"verlauf",
        r"sitzung",
    ],
    "aktendeckblatt": [
        r"aktendeckblatt",
        r"aktendatenblatt",
        r"anmeld(?:ung|ebogen|eunterlagen)",
        r"stammdaten",
    ],
    "hilfeplan": [
        r"hilfeplan",
        r"kostenzusage",
        r"bewilligung",
        r"bescheid",
    ],
    "proautismus_bericht": [
        r"entwicklungsbericht",
        r"foerderbericht",
        r"icf.?bericht",
        r"bericht.*proautismus",
    ],

    # STUFE2 Dokumente
    "arztbericht": [
        r"arztbericht",
        r"bericht.*dr\.?",
        r"entlass(?:ungs)?bericht",
        r"diagnose",
        r"befund",
        r"gutachten",
    ],
    "schulbericht": [
        r"schulbericht",
        r"zeugnis",
        r"lernentwicklung",
        r"foerderplan.*schule",
        r"nachteilsausgleich",
    ],
    "mail": [
        r"\.msg$",
        r"\.eml$",
    ],
}

# Ordnernamen die auf Kategorie hinweisen
FOLDER_HINTS = {
    "core": [
        "dokumentation", "protokolle", "anmeldeunterlagen",
        "hilfeplan", "amt", "intern"
    ],
    "stufe2": [
        "familie", "umfeld", "extern", "arzt", "schule", "mail"
    ],
    "extended": [
        "archiv", "alt", "historie", "backup"
    ],
    "skip": [
        "output", "_prepare", "_ready", "_archive", "temp", "_bundle"
    ],
}


class DocumentCollector:
    """
    Sammelt und kategorisiert Dokumente fuer die Berichtserstellung.

    Usage:
        collector = DocumentCollector(berichtszeitraum_monate=12)
        result = collector.scan_folder("/path/to/client/folder")
        bundle = collector.extract_bundle(result.documents)
    """

    def __init__(
        self,
        berichtszeitraum_monate: int = 12,
        stichtag: Optional[datetime] = None
    ):
        """
        Args:
            berichtszeitraum_monate: Zeitraum fuer aktuelle Dokumente (Default: 12)
            stichtag: Referenzdatum (Default: heute)
        """
        self.berichtszeitraum_monate = berichtszeitraum_monate
        self.stichtag = stichtag or datetime.now()
        self.zeitraum_start = self.stichtag - timedelta(days=berichtszeitraum_monate * 30)

        # Compile regex patterns
        self._compiled_patterns: Dict[str, List[re.Pattern]] = {}
        for doc_type, patterns in DOC_PATTERNS.items():
            self._compiled_patterns[doc_type] = [
                re.compile(p, re.IGNORECASE) for p in patterns
            ]

    def scan_folder(self, folder: str, recursive: bool = True) -> CollectorResult:
        """
        Scannt einen Ordner und kategorisiert alle Dokumente.

        Args:
            folder: Pfad zum Klienten-Ordner
            recursive: Unterordner einbeziehen (Default: True)

        Returns:
            CollectorResult mit kategorisierten Dokumenten
        """
        result = CollectorResult()
        folder_path = Path(folder)

        if not folder_path.exists():
            result.errors.append(f"Ordner nicht gefunden: {folder}")
            return result

        # Unterstuetzte Dateitypen
        supported = {".docx", ".doc", ".txt", ".md", ".pdf", ".xlsx", ".xls", ".msg", ".eml"}

        # Dateien sammeln
        if recursive:
            files = list(folder_path.rglob("*"))
        else:
            files = list(folder_path.iterdir())

        for filepath in sorted(files):
            if not filepath.is_file():
                continue
            if filepath.name.startswith("."):
                continue
            if filepath.suffix.lower() not in supported:
                continue

            # Dokument analysieren
            doc_info = self._analyze_document(filepath, folder_path)
            result.documents.append(doc_info)

            # Zaehler aktualisieren
            if doc_info.category == DocumentCategory.CORE:
                result.core_count += 1
            elif doc_info.category == DocumentCategory.STUFE2:
                result.stufe2_count += 1
            elif doc_info.category == DocumentCategory.EXTENDED:
                result.extended_count += 1
            else:
                result.skipped_count += 1

        # Nach Kategorie und dann Datum sortieren
        result.documents.sort(key=lambda d: (
            d.category.value,
            d.date_hint or datetime.min
        ))

        return result

    def _analyze_document(self, filepath: Path, base_folder: Path) -> DocumentInfo:
        """Analysiert ein einzelnes Dokument und bestimmt Kategorie."""
        filename = filepath.name
        suffix = filepath.suffix.lower()
        
        try:
            rel_path = filepath.relative_to(base_folder)
            folder_parts = [p.lower() for p in rel_path.parts[:-1]]
        except ValueError:
            # Fallback falls relative_to fehlschlaegt (z.B. bei Symlinks)
            folder_parts = []

        # Basis-Info
        doc_info = DocumentInfo(
            path=filepath,
            filename=filename,
            suffix=suffix,
            category=DocumentCategory.EXTENDED,  # Default
            doc_type="unknown",
            size_bytes=filepath.stat().st_size if filepath.exists() else 0
        )

        # Datum aus Dateinamen extrahieren
        doc_info.date_hint = self._extract_date(filename)

        # Ordner-basierte Skip-Erkennung
        for skip_folder in FOLDER_HINTS["skip"]:
            if any(skip_folder in part for part in folder_parts):
                doc_info.category = DocumentCategory.SKIP
                return doc_info

        # Dokumenttyp erkennen
        doc_type = self._detect_doc_type(filename, suffix)
        doc_info.doc_type = doc_type

        # Regel: Dateien im Root-Ordner sind IMMER wichtig -> CORE
        # Ausnahme: Dateien die mit _ beginnen (temp/generated files)
        is_root = len(folder_parts) == 0
        if is_root and not filename.startswith("_"):
            doc_info.category = DocumentCategory.CORE
            if doc_type == "unknown":
                doc_info.doc_type = "root_dokument"
            return doc_info

        # Kategorie basierend auf Typ und Datum
        if doc_type in ["protokoll", "aktendeckblatt"]:
            # Protokolle im Berichtszeitraum sind CORE
            if self._is_in_berichtszeitraum(doc_info.date_hint):
                doc_info.category = DocumentCategory.CORE
            else:
                doc_info.category = DocumentCategory.EXTENDED

        elif doc_type == "hilfeplan":
            # Aktuellster Hilfeplan ist CORE
            doc_info.category = DocumentCategory.CORE

        elif doc_type == "proautismus_bericht":
            # Letzter eigener Bericht ist CORE (fuer Kontinuitaet)
            doc_info.category = DocumentCategory.CORE

        elif doc_type in ["arztbericht", "schulbericht"]:
            # Aktuellster ist STUFE2, aeltere EXTENDED
            if self._is_recent(doc_info.date_hint, years=2):
                doc_info.category = DocumentCategory.STUFE2
            else:
                doc_info.category = DocumentCategory.EXTENDED

        elif doc_type == "mail":
            # Mails im Berichtszeitraum sind STUFE2
            if self._is_in_berichtszeitraum(doc_info.date_hint):
                doc_info.category = DocumentCategory.STUFE2
            else:
                doc_info.category = DocumentCategory.EXTENDED

        # Ordner-Hints koennen Kategorie ueberschreiben
        for part in folder_parts:
            if any(hint in part for hint in FOLDER_HINTS["core"]):
                if doc_info.category != DocumentCategory.SKIP:
                    # Nur hochstufen wenn im Berichtszeitraum
                    if self._is_in_berichtszeitraum(doc_info.date_hint):
                        doc_info.category = DocumentCategory.CORE
            elif any(hint in part for hint in FOLDER_HINTS["extended"]):
                if doc_info.category not in [DocumentCategory.SKIP, DocumentCategory.CORE]:
                    doc_info.category = DocumentCategory.EXTENDED

        return doc_info

    def _detect_doc_type(self, filename: str, suffix: str) -> str:
        """Erkennt den Dokumenttyp anhand des Dateinamens."""
        # Mail-Dateien direkt erkennen
        if suffix in [".msg", ".eml"]:
            return "mail"

        # Pattern-Matching
        for doc_type, patterns in self._compiled_patterns.items():
            for pattern in patterns:
                if pattern.search(filename):
                    return doc_type

        return "unknown"

    def _extract_date(self, filename: str) -> Optional[datetime]:
        """Versucht ein Datum aus dem Dateinamen zu extrahieren."""
        # Muster geordnet von spezifisch nach generisch
        # \b erzwingt Wortgrenzen, um 2024-01-15 nicht als 2024-01 zu erkennen
        patterns = [
            r"\b(\d{4})-(\d{2})-(\d{2})\b",     # 2024-01-15 (Priority 1)
            r"\b(\d{2})\.(\d{2})\.(\d{4})\b",   # 15.01.2024
            r"\b(\d{4})-(\d{2})\b",             # 2024-01
            r"\b(\d{2})-(\d{4})\b",             # 01-2024
            r"_(\d{4})(?:[_\.\s]|$)",           # _2024_ (Jahr allein)
        ]

        for pattern in patterns:
            match = re.search(pattern, filename)
            if match:
                groups = match.groups()
                try:
                    if len(groups) == 3:
                        if len(groups[0]) == 4:  # YYYY-MM-DD
                            return datetime(int(groups[0]), int(groups[1]), int(groups[2]))
                        else:  # DD.MM.YYYY
                            return datetime(int(groups[2]), int(groups[1]), int(groups[0]))
                    elif len(groups) == 2:
                        if len(groups[0]) == 4:  # YYYY-MM
                            return datetime(int(groups[0]), int(groups[1]), 1)
                        else:  # MM-YYYY
                            return datetime(int(groups[1]), int(groups[0]), 1)
                    elif len(groups) == 1:  # Just year
                        return datetime(int(groups[0]), 6, 15)  # Mitte des Jahres
                except (ValueError, IndexError):
                    pass

        return None

    def _is_in_berichtszeitraum(self, date: Optional[datetime]) -> bool:
        """Prueft ob ein Datum im aktuellen Berichtszeitraum liegt."""
        if date is None:
            return True  # Im Zweifel einbeziehen
        return date >= self.zeitraum_start

    def _is_recent(self, date: Optional[datetime], years: int = 2) -> bool:
        """Prueft ob ein Datum innerhalb der letzten X Jahre liegt."""
        if date is None:
            return True  # Im Zweifel einbeziehen
        cutoff = self.stichtag - timedelta(days=years * 365)
        return date >= cutoff

    def extract_bundle(
        self,
        documents: List[DocumentInfo],
        include_extended: bool = False,
        anonym_profile: Optional["AnonymProfile"] = None
    ) -> TextBundle:
        """
        Extrahiert Text aus den kategorisierten Dokumenten.

        Args:
            documents: Liste von DocumentInfo (aus scan_folder)
            include_extended: Auch EXTENDED Dokumente einbeziehen
            anonym_profile: Optional - AnonymProfile fuer sofortige Anonymisierung

        Returns:
            TextBundle mit extrahiertem (und optional anonymisiertem) Text
        """
        bundle = TextBundle()

        # Anonymisierungs-Mappings vorbereiten
        if anonym_profile:
            all_replacements = {}
            for category in anonym_profile.mappings.values():
                all_replacements.update(category)
            sorted_replacements = sorted(
                all_replacements.items(),
                key=lambda x: len(x[0]),
                reverse=True
            )
        else:
            sorted_replacements = []

        def anonymize_text(text: str) -> str:
            """Wendet Anonymisierung auf Text an."""
            if not sorted_replacements:
                return text
            for old, new in sorted_replacements:
                text = text.replace(old, new)
            return text

        def anonymize_filename(filename: str) -> str:
            """Anonymisiert Dateinamen fuer Header."""
            if not sorted_replacements:
                return filename
            for old, new in sorted_replacements:
                filename = filename.replace(old, new)
            return filename

        core_parts = []
        stufe2_parts = []
        extended_parts = []

        for doc in documents:
            if doc.category == DocumentCategory.SKIP:
                continue
            if doc.category == DocumentCategory.EXTENDED and not include_extended:
                continue

            # Text extrahieren
            text = self._extract_text(doc)
            if not text or text.startswith("[FEHLER"):
                continue

            # Text anonymisieren
            text = anonymize_text(text)

            doc.text_extracted = True
            doc.text_length = len(text)
            bundle.files_processed += 1

            # Header mit anonymisiertem Dateinamen
            anon_filename = anonymize_filename(doc.filename)
            header = f"--- {doc.doc_type.upper()}: {anon_filename} ---"
            content = f"{header}\n{text}"

            if doc.category == DocumentCategory.CORE:
                core_parts.append(content)
            elif doc.category == DocumentCategory.STUFE2:
                stufe2_parts.append(content)
            else:
                extended_parts.append(content)

        bundle.core_text = "\n\n".join(core_parts)
        bundle.stufe2_text = "\n\n".join(stufe2_parts)
        bundle.extended_text = "\n\n".join(extended_parts)
        bundle.total_length = len(bundle.core_text) + len(bundle.stufe2_text)

        return bundle

    def _extract_text(self, doc: DocumentInfo) -> str:
        """Extrahiert Text aus einem Dokument."""
        filepath = str(doc.path)
        suffix = doc.suffix

        try:
            if suffix == ".docx":
                return self._extract_docx(filepath)
            elif suffix == ".doc":
                return self._extract_doc(filepath)
            elif suffix in [".txt", ".md"]:
                return self._extract_txt(filepath)
            elif suffix == ".pdf":
                return self._extract_pdf(filepath)
            elif suffix in [".xlsx", ".xls"]:
                return self._extract_excel(filepath)
            elif suffix == ".msg":
                return self._extract_msg(filepath)
            elif suffix == ".eml":
                return self._extract_eml(filepath)
            else:
                return f"[Nicht unterstuetztes Format: {suffix}]"
        except Exception as e:
            return f"[FEHLER bei {doc.filename}: {e}]"

    def _extract_docx(self, filepath: str) -> str:
        """Extrahiert Text aus Word-Dokument (.docx)."""
        try:
            from docx import Document
            doc = Document(filepath)
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        except ImportError:
            return "[python-docx nicht installiert]"
        except Exception as e:
            # Manchmal ist eine .doc Datei als .docx benanned
            return f"[DOCX-Fehler: {e}]"

    def _extract_doc(self, filepath: str) -> str:
        """Extrahiert Text aus altem Word-Format (.doc) via antiword oder textract."""
        import subprocess
        import shutil

        # Methode 1: antiword (Linux/Windows mit antiword installiert)
        if shutil.which("antiword"):
            try:
                result = subprocess.run(
                    ["antiword", filepath],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                if result.returncode == 0 and result.stdout.strip():
                    return result.stdout.strip()
            except Exception:
                pass

        # Methode 2: LibreOffice Konvertierung zu txt
        if shutil.which("soffice"):
            try:
                import tempfile
                with tempfile.TemporaryDirectory() as tmpdir:
                    result = subprocess.run(
                        ["soffice", "--headless", "--convert-to", "txt:Text",
                         "--outdir", tmpdir, filepath],
                        capture_output=True,
                        timeout=60
                    )
                    if result.returncode == 0:
                        txt_file = Path(tmpdir) / (Path(filepath).stem + ".txt")
                        if txt_file.exists():
                            return txt_file.read_text(encoding="utf-8", errors="replace")
            except Exception:
                pass

        # Methode 3: textract (falls installiert)
        try:
            import textract
            text = textract.process(filepath).decode("utf-8", errors="replace")
            return text.strip()
        except ImportError:
            pass
        except Exception:
            pass

        return "[.doc-Extraktion fehlgeschlagen - antiword/LibreOffice/textract nicht verfuegbar]"

    def _extract_txt(self, filepath: str) -> str:
        """Liest Textdatei."""
        path = Path(filepath)
        try:
            return path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return path.read_text(encoding="latin-1")

    def _extract_pdf(self, filepath: str) -> str:
        """Extrahiert Text aus PDF, mit OCR-Fallback fuer Bild-PDFs."""
        text_parts = []
        needs_ocr = False

        # Primaer: pypdf (MIT-Lizenz)
        try:
            from pypdf import PdfReader
            reader = PdfReader(filepath)
            for page in reader.pages:
                page_text = (page.extract_text() or "").strip()
                if page_text:
                    text_parts.append(page_text)
                else:
                    needs_ocr = True

            if text_parts:
                if needs_ocr and not text_parts:
                    return self._extract_pdf_ocr(filepath)
                return "\n".join(text_parts)
        except ImportError:
            pass
        except Exception:
            pass

        # Fallback: pdfplumber (MIT-Lizenz) -- besser bei Tabellen
        try:
            import pdfplumber
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    page_text = (page.extract_text() or "").strip()
                    if page_text:
                        text_parts.append(page_text)
                    else:
                        needs_ocr = True
            if text_parts:
                return "\n".join(text_parts)
        except ImportError:
            pass
        except Exception:
            pass

        # Optional: PyMuPDF (AGPL) -- nur wenn verfuegbar
        try:
            import fitz  # PyMuPDF optional
            doc = fitz.open(filepath)
            fitz_parts = []
            fitz_needs_ocr = False
            for page in doc:
                page_text = page.get_text().strip()
                if page_text:
                    fitz_parts.append(page_text)
                else:
                    fitz_needs_ocr = True
            doc.close()
            if fitz_parts:
                if fitz_needs_ocr and not fitz_parts:
                    return self._extract_pdf_ocr(filepath)
                return "\n".join(fitz_parts)
        except ImportError:
            pass
        except Exception:
            pass

        # Wenn kein Text gefunden, OCR versuchen
        if needs_ocr:
            return self._extract_pdf_ocr(filepath)

        return "[PDF-Extraktion fehlgeschlagen: pypdf, pdfplumber und PyMuPDF nicht verfuegbar]"

    def _extract_pdf_ocr(self, filepath: str) -> str:
        """OCR fuer Bild-PDFs via c_ocr_engine."""
        try:
            # Versuche den existierenden OCR-Engine zu nutzen
            # Nutze bach_paths.py als Single Source of Truth fuer Pfade

            # 1. Versuche Import, falls im Pfad
            try:
                from c_ocr_engine import ocr_pdf, is_ocr_available
            except ImportError:
                # 2. Nutze bach_paths.py fuer den Tools-Pfad (Single Source of Truth)
                tools_dir = None
                try:
                    # bach_paths sollte bereits im sys.path sein (via hub/)
                    from bach_paths import get_path
                    tools_dir = get_path("tools")
                except ImportError:
                    # Fallback: Manuell suchen wenn bach_paths nicht verfuegbar
                    current = Path(__file__).resolve()
                    for _ in range(6):
                        current = current.parent
                        candidates = [
                            current / "skills" / "tools",
                            current / "system" / "skills" / "tools",
                        ]
                        for candidate in candidates:
                            if candidate.exists() and (candidate / "c_ocr_engine.py").exists():
                                tools_dir = candidate
                                break
                        if tools_dir:
                            break

                if tools_dir and str(tools_dir) not in sys.path:
                    sys.path.insert(0, str(tools_dir))

                from c_ocr_engine import ocr_pdf, is_ocr_available

            if not is_ocr_available():
                return "[OCR nicht verfuegbar - Tesseract nicht installiert]"

            text = ocr_pdf(filepath)
            return text if text else "[OCR: Kein Text erkannt]"
        except ImportError as e:
            return f"[OCR Import-Fehler: {e}]"
        except Exception as e:
            return f"[OCR-Fehler: {e}]"

    def _extract_excel(self, filepath: str) -> str:
        """Extrahiert Text aus Excel."""
        try:
            import openpyxl
            wb = openpyxl.load_workbook(filepath, data_only=True)
            parts = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                parts.append(f"[Tabelle: {sheet_name}]")
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(c) if c else "" for c in row]
                    if any(cells):
                        parts.append(" | ".join(cells))
            wb.close()
            return "\n".join(parts)
        except ImportError:
            return "[openpyxl nicht installiert]"

    def _extract_msg(self, filepath: str) -> str:
        """Extrahiert Text aus Outlook .msg (inkl. Anhaenge)."""
        try:
            import extract_msg
            msg = extract_msg.Message(filepath)
            parts = []
            if msg.date:
                parts.append(f"Datum: {msg.date}")
            if msg.sender:
                parts.append(f"Von: {msg.sender}")
            if msg.to:
                parts.append(f"An: {msg.to}")
            if msg.subject:
                parts.append(f"Betreff: {msg.subject}")
            parts.append("")
            if msg.body:
                parts.append(msg.body)

            # Anhaenge extrahieren (PDF, DOCX, TXT)
            if hasattr(msg, 'attachments') and msg.attachments:
                for attach in msg.attachments:
                    try:
                        attach_name = attach.longFilename or attach.shortFilename or "unbekannt"
                        suffix = Path(attach_name).suffix.lower()

                        if suffix == ".pdf":
                            # PDF-Anhang temporaer speichern und extrahieren
                            import tempfile
                            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                                tmp.write(attach.data)
                                tmp_path = tmp.name
                            parts.append(f"\n[Anhang: {attach_name}]")
                            pdf_text = self._extract_pdf(tmp_path)
                            parts.append(pdf_text)
                            Path(tmp_path).unlink(missing_ok=True)

                        elif suffix == ".docx":
                            import tempfile
                            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                                tmp.write(attach.data)
                                tmp_path = tmp.name
                            parts.append(f"\n[Anhang: {attach_name}]")
                            docx_text = self._extract_docx(tmp_path)
                            parts.append(docx_text)
                            Path(tmp_path).unlink(missing_ok=True)

                        elif suffix in [".txt", ".md"]:
                            parts.append(f"\n[Anhang: {attach_name}]")
                            try:
                                text = attach.data.decode('utf-8')
                            except:
                                text = attach.data.decode('latin-1')
                            parts.append(text)

                    except Exception as e:
                        parts.append(f"\n[Anhang {attach_name}: Fehler - {e}]")

            msg.close()
            return "\n".join(parts)
        except ImportError:
            return "[extract-msg nicht installiert]"

    def _extract_eml(self, filepath: str) -> str:
        """Extrahiert Text aus .eml."""
        import email
        from email import policy
        from email.parser import BytesParser

        with open(filepath, 'rb') as f:
            msg = BytesParser(policy=policy.default).parse(f)

        parts = []
        if msg['Date']:
            parts.append(f"Datum: {msg['Date']}")
        if msg['From']:
            parts.append(f"Von: {msg['From']}")
        if msg['To']:
            parts.append(f"An: {msg['To']}")
        if msg['Subject']:
            parts.append(f"Betreff: {msg['Subject']}")
        parts.append("")

        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type() == 'text/plain':
                    charset = part.get_content_charset() or 'utf-8'
                    try:
                        body = part.get_payload(decode=True).decode(charset)
                        parts.append(body)
                        break
                    except:
                        pass
        else:
            charset = msg.get_content_charset() or 'utf-8'
            try:
                body = msg.get_payload(decode=True).decode(charset)
                parts.append(body)
            except:
                parts.append(str(msg.get_payload()))

        return "\n".join(parts)

    def create_anonymized_bundle(
        self,
        bundle: TextBundle,
        profile: "AnonymProfile"
    ) -> TextBundle:
        """
        Anonymisiert ein Text-Bundle.

        Args:
            bundle: Extrahiertes TextBundle
            profile: Anonymisierungsprofil

        Returns:
            Neues TextBundle mit anonymisiertem Text
        """
        # Alle Mappings sammeln
        all_replacements = {}
        for category in profile.mappings.values():
            all_replacements.update(category)

        # Nach Laenge sortieren (laengste zuerst)
        sorted_replacements = sorted(
            all_replacements.items(),
            key=lambda x: len(x[0]),
            reverse=True
        )

        def anonymize_text(text: str) -> str:
            for old, new in sorted_replacements:
                text = text.replace(old, new)
            return text

        return TextBundle(
            core_text=anonymize_text(bundle.core_text),
            stufe2_text=anonymize_text(bundle.stufe2_text),
            extended_text=anonymize_text(bundle.extended_text),
            total_length=bundle.total_length,
            files_processed=bundle.files_processed
        )

    def save_bundle(
        self,
        bundle: TextBundle,
        output_path: str,
        combined: bool = True
    ) -> Path:
        """
        Speichert ein Bundle als Textdatei.

        Args:
            bundle: Das TextBundle
            output_path: Zielpfad
            combined: CORE + STUFE2 kombiniert speichern

        Returns:
            Pfad zur gespeicherten Datei
        """
        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)

        if combined:
            content = "=== CORE DOKUMENTE ===\n\n"
            content += bundle.core_text
            content += "\n\n=== STUFE 2 DOKUMENTE ===\n\n"
            content += bundle.stufe2_text

            path.write_text(content, encoding="utf-8")
        else:
            # Separate Dateien
            core_path = path.with_suffix(".core.txt")
            core_path.write_text(bundle.core_text, encoding="utf-8")

            stufe2_path = path.with_suffix(".stufe2.txt")
            stufe2_path.write_text(bundle.stufe2_text, encoding="utf-8")

        return path


# ═══════════════════════════════════════════════════════════════
# Hilfs-Funktionen fuer Anonymisierung
# ═══════════════════════════════════════════════════════════════

def create_profile_from_folder(folder_path: str, geburtsdatum: str = "01.01.2010") -> "AnonymProfile":
    """
    Erstellt ein AnonymProfile basierend auf dem Ordnernamen.

    Erwartet Ordnernamen im Format: "Nachname, Vorname" oder "Vorname Nachname"

    Args:
        folder_path: Pfad zum Klienten-Ordner
        geburtsdatum: Geburtsdatum falls bekannt (dd.mm.yyyy)

    Returns:
        AnonymProfile mit generierten Tarnnamen
    """
    from .anonymizer_service import DocumentAnonymizer

    folder = Path(folder_path)
    folder_name = folder.name

    # Name aus Ordnernamen extrahieren
    # Format: "Nachname, Vorname" -> "Vorname Nachname"
    if ", " in folder_name:
        parts = folder_name.split(", ")
        name = f"{parts[1]} {parts[0]}"
    else:
        name = folder_name

    # Profil erstellen
    anonymizer = DocumentAnonymizer()
    profile = anonymizer.create_profile(
        real_name=name.strip(),
        geburtsdatum=geburtsdatum,
        whitelist=[]  # User-specific: add names to preserve here
    )

    return profile


def extract_with_anonymization(
    folder_path: str,
    geburtsdatum: str = "01.01.2010",
    output_path: Optional[str] = None
) -> Tuple[TextBundle, "AnonymProfile"]:
    """
    Kompletter Workflow: Scannen + Extrahieren + Anonymisieren.

    Args:
        folder_path: Pfad zum Klienten-Ordner
        geburtsdatum: Geburtsdatum des Klienten
        output_path: Optional - Pfad zum Speichern des Bundles

    Returns:
        (TextBundle, AnonymProfile)
    """
    # Profil erstellen
    profile = create_profile_from_folder(folder_path, geburtsdatum)

    # Collector
    collector = DocumentCollector()
    result = collector.scan_folder(folder_path)

    # Bundle mit Anonymisierung extrahieren
    bundle = collector.extract_bundle(
        result.documents,
        include_extended=False,
        anonym_profile=profile
    )

    # Optional speichern
    if output_path:
        collector.save_bundle(bundle, output_path)

    return bundle, profile


# ═══════════════════════════════════════════════════════════════
# CLI Interface
# ═══════════════════════════════════════════════════════════════

def main():
    """CLI Einstiegspunkt."""
    import sys

    if len(sys.argv) < 2:
        print("Document Collector Service v1.0.1")
        print()
        print("Usage:")
        print("  python document_collector.py scan <ordner>")
        print("  python document_collector.py extract <ordner> -o <output.txt>")
        print("  python document_collector.py test")
        return

    cmd = sys.argv[1]

    if cmd == "scan":
        if len(sys.argv) < 3:
            print("Fehler: Ordner angeben")
            return

        folder = sys.argv[2]
        collector = DocumentCollector()
        result = collector.scan_folder(folder)

        print(f"\nGefundene Dokumente: {len(result.documents)}")
        print(f"  CORE: {result.core_count}")
        print(f"  STUFE2: {result.stufe2_count}")
        print(f"  EXTENDED: {result.extended_count}")
        print(f"  Uebersprungen: {result.skipped_count}")

        print("\n--- CORE ---")
        for doc in result.documents:
            if doc.category == DocumentCategory.CORE:
                print(f"  [{doc.doc_type}] {doc.filename}")

        print("\n--- STUFE2 ---")
        for doc in result.documents:
            if doc.category == DocumentCategory.STUFE2:
                print(f"  [{doc.doc_type}] {doc.filename}")

        if result.errors:
            print("\nFehler:")
            for err in result.errors:
                print(f"  {err}")

    elif cmd == "extract":
        if len(sys.argv) < 3:
            print("Fehler: Ordner angeben")
            return

        folder = sys.argv[2]
        output = None
        if "-o" in sys.argv:
            idx = sys.argv.index("-o")
            if idx + 1 < len(sys.argv):
                output = sys.argv[idx + 1]

        collector = DocumentCollector()
        result = collector.scan_folder(folder)
        bundle = collector.extract_bundle(result.documents)

        print(f"Extrahiert: {bundle.files_processed} Dateien")
        print(f"Textlaenge: {bundle.total_length} Zeichen")

        if output:
            collector.save_bundle(bundle, output)
            print(f"Gespeichert: {output}")
        else:
            print("\n--- CORE TEXT (Auszug) ---")
            print(bundle.core_text[:2000] + "..." if len(bundle.core_text) > 2000 else bundle.core_text)

    elif cmd == "test":
        print("[TEST] Document Collector")
        collector = DocumentCollector(berichtszeitraum_monate=12)

        # Test Datumsextraktion
        test_files = [
            "Protokoll_2024-01-15.docx",
            "Hilfeplan 01.06.2023.pdf",
            "Bericht_2022.docx",
            "Arztbericht Dr. Mueller.pdf",
        ]

        print("\nDatumsextraktion:")
        for f in test_files:
            date = collector._extract_date(f)
            print(f"  {f} -> {date}")

        print("\nDokumenttyp-Erkennung:")
        for f in test_files:
            dtype = collector._detect_doc_type(f, Path(f).suffix)
            print(f"  {f} -> {dtype}")

        print("\n[OK] Test abgeschlossen")

    else:
        print(f"Unbekannter Befehl: {cmd}")


if __name__ == "__main__":
    main()
