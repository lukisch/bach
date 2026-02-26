#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# SPDX-License-Identifier: MIT
"""
Copyright (c) 2026 Lukas Geiger

Permission is hereby granted, free of charge, to any person obtaining a copy
of this software and associated documentation files (the "Software"), to deal
in the Software without restriction, including without limitation the rights
to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
copies of the Software, and to permit persons to whom the Software is
furnished to do so, subject to the following conditions:

The above copyright notice and this permission notice shall be included in all
copies or substantial portions of the Software.

THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
SOFTWARE.
"""

"""
Report Workflow Service
========================

Vereinheitlichter Workflow fuer Foerderberichte:
  1. Dateien importieren
  2. Temporaeres Anonymisierungsprofil erstellen
  3. Dokumente anonymisieren und in bundles/ sortieren
  4. Prompt mit Wissensdatenbank-Kontext generieren
  5. Word-Dokument generieren und de-anonymisieren
  6. Cleanup - alle temporaeren Daten loeschen

KEIN permanentes Speichern von Profilen mehr!
Jeder Durchlauf startet frisch.

Version: 1.0.0
Erstellt: 2026-02-01
"""

import json
import shutil
import sys
import tempfile
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import secrets

# Lokale Imports - mit Fallback für CLI-Ausführung
try:
    from .anonymizer_service import (
        DocumentAnonymizer,
        DocumentDeanonymizer,
        AnonymProfile,
        _generate_tarnname,
        _detect_gender,
        _generate_fake_date_same_age,
        _generate_fake_phone,
        _generate_fake_email,
        _generate_fake_address,
        TARN_VORNAMEN,
        TARN_VORNAMEN_M,
        TARN_VORNAMEN_W,
        TARN_NACHNAMEN
    )
    from .document_collector import (
        DocumentCollector,
        DocumentCategory,
        TextBundle,
        CollectorResult
    )
    from .word_template_service import WordTemplateService
except ImportError:
    # CLI-Modus: Absoluter Import aus dem gleichen Ordner
    import sys
    _this_dir = Path(__file__).parent.resolve()
    if str(_this_dir) not in sys.path:
        sys.path.insert(0, str(_this_dir))
    # Auch BACH root fuer andere imports
    _bach_root = _this_dir.parent.parent.parent.parent
    if str(_bach_root / "system") not in sys.path:
        sys.path.insert(0, str(_bach_root / "system"))

    from anonymizer_service import (
        DocumentAnonymizer,
        DocumentDeanonymizer,
        AnonymProfile,
        _generate_tarnname,
        _detect_gender,
        _generate_fake_date_same_age,
        _generate_fake_phone,
        _generate_fake_email,
        _generate_fake_address,
        TARN_VORNAMEN,
        TARN_VORNAMEN_M,
        TARN_VORNAMEN_W,
        TARN_NACHNAMEN
    )
    from document_collector import (
        DocumentCollector,
        DocumentCategory,
        TextBundle,
        CollectorResult
    )
    from word_template_service import WordTemplateService


# ═══════════════════════════════════════════════════════════════
# Universal Import: bach_paths.py (Single Source of Truth)
# ═══════════════════════════════════════════════════════════════
# Dieses Pattern funktioniert von ueberall im System

_current = Path(__file__).resolve()
for _parent in [_current] + list(_current.parents):
    _hub = _parent / "system" / "hub"
    if _hub.exists():
        if str(_hub) not in sys.path:
            sys.path.insert(0, str(_hub))
        break

# Jetzt koennen wir bach_paths importieren
try:
    from bach_paths import get_path, BACH_ROOT
    _USE_BACH_PATHS = True
except ImportError:
    _USE_BACH_PATHS = False


# ═══════════════════════════════════════════════════════════════
# Konfiguration
# ═══════════════════════════════════════════════════════════════

# Basis-Pfade - aus bach_paths.py oder Fallback
if _USE_BACH_PATHS:
    BACH_BASE = BACH_ROOT
    BERICHTE_BASE = get_path("berichte")
    TEMPLATE_PATH = get_path("bericht_template")
    WISSENSDATENBANK = get_path("wissensdatenbank")
else:
    # Fallback: Dynamisch ermitteln
    _SERVICE_DIR = Path(__file__).resolve().parent
    BACH_BASE = _SERVICE_DIR.parent.parent.parent.parent
    BERICHTE_BASE = BACH_BASE / "user" / "documents" / "foerderplaner" / "Berichte"
    TEMPLATE_PATH = BACH_BASE / "system" / "skills" / "_templates" / "bericht_template_geiger_universal.docx"
    WISSENSDATENBANK = Path.home() / "OneDrive" / "Dokumente" / "_Wissensdatenbank"

ICF_REF_PATH = BACH_BASE / "system" / "skills" / "_experts" / "foerderplaner" / "Kontext" / "ICF-Checker Tabelle_Geiger.docx"

# Whitelist: Diese Namen werden NICHT anonymisiert
DEFAULT_WHITELIST = [
    "YOUR_NAME",  # User-specific: add names to preserve here
    "Landratsamt", "Jugendamt", "Sozialamt",
    "Händle", "Frau Händle"
]


# ═══════════════════════════════════════════════════════════════
# Datenklassen
# ═══════════════════════════════════════════════════════════════

@dataclass
class TempAnonymProfile:
    """
    Temporaeres Anonymisierungsprofil.

    Wird NICHT persistent gespeichert - existiert nur waehrend
    des Workflow-Durchlaufs im Speicher.
    """
    tarnname: str
    fake_geburtsdatum: str
    mappings: Dict[str, Dict[str, str]]
    original_name: str = ""
    _temp_dir: Optional[Path] = None
    created: str = field(default_factory=lambda: datetime.now().isoformat())

    def get_all_mappings(self) -> Dict[str, str]:
        """Flache Map aller Ersetzungen (Original -> Tarn)."""
        result = {}
        for category in self.mappings.values():
            result.update(category)
        return result

    def get_reverse_mappings(self) -> Dict[str, str]:
        """Umgekehrte Map (Tarn -> Original) fuer De-Anonymisierung."""
        result = {}
        for category in self.mappings.values():
            for original, tarn in category.items():
                result[tarn] = original
        return result

    def cleanup(self):
        """Loescht temporaere Daten."""
        if self._temp_dir and self._temp_dir.exists():
            shutil.rmtree(self._temp_dir, ignore_errors=True)
            self._temp_dir = None

    def to_anonym_profile(self) -> AnonymProfile:
        """Konvertiert zu AnonymProfile fuer bestehende Services."""
        return AnonymProfile(
            client_id=f"TEMP_{secrets.token_hex(4)}",
            tarnname=self.tarnname,
            fake_geburtsdatum=self.fake_geburtsdatum,
            mappings=self.mappings,
            created=self.created
        )


@dataclass
class WorkflowSession:
    """Session-Status fuer einen Workflow-Durchlauf."""
    session_id: str
    status: str = "initialized"  # initialized, importing, anonymizing, generating, completed, error
    profile: Optional[TempAnonymProfile] = None
    input_files: List[Path] = field(default_factory=list)
    bundle: Optional[TextBundle] = None
    output_path: Optional[Path] = None
    error_message: str = ""
    created: str = field(default_factory=lambda: datetime.now().isoformat())

    def to_dict(self) -> dict:
        """Serialisiert Session fuer API-Response."""
        return {
            "session_id": self.session_id,
            "status": self.status,
            "tarnname": self.profile.tarnname if self.profile else None,
            "input_files_count": len(self.input_files),
            "output_path": str(self.output_path) if self.output_path else None,
            "error": self.error_message or None,
            "created": self.created
        }


# ═══════════════════════════════════════════════════════════════
# Report Workflow Service
# ═══════════════════════════════════════════════════════════════

class ReportWorkflowService:
    """
    Vereinheitlichter Workflow fuer Foerderberichte.

    Workflow:
        1. start_session() - Neue Session starten
        2. import_files() - Dateien nach data/ kopieren
        3. create_temp_profile() - Temporaeres Anonymisierungsprofil
        4. anonymize_to_bundles() - Dateien anonymisieren
        5. generate_prompt() - Prompt mit Wissensdatenbank
        6. generate_report() - Word-Dokument erstellen
        7. cleanup() - Alles aufräumen

    Usage:
        service = ReportWorkflowService()
        session = service.start_session()
        service.import_files(session, [Path("dokument.docx")])
        service.create_temp_profile(session, "Max Mustermann", "15.03.2016")
        service.anonymize_to_bundles(session)
        prompt = service.generate_prompt(session)
        # ... LLM aufrufen ...
        service.generate_report(session, llm_response)
        service.cleanup(session)
    """

    def __init__(self, base_path: Optional[Path] = None):
        """
        Args:
            base_path: Basis-Ordner fuer Berichte (Default: user/foerderplaner/Berichte)
        """
        self.base_path = base_path or BERICHTE_BASE
        self._ensure_folder_structure()
        self._sessions: Dict[str, WorkflowSession] = {}
        self._used_tarnnames: set = set()

    def _get_icf_reference(self) -> str:
        """Gibt die detaillierte Liste aller ICF-Codes zurück."""
        icf_path = Path(ICF_REF_PATH)
        if icf_path.exists():
            try:
                # Nutze DocumentAnonymizer für Text-Extraktion aus Docx
                anonymizer = DocumentAnonymizer()
                text = anonymizer.extract_text_from_file(str(icf_path))
                if text.strip():
                    # Vorne wegwerfen falls wir wissen wo die Codes anfangen
                    return text.strip()
            except Exception as e:
                print(f"[WARN] konnte ICF-Referenz nicht aus DOCX lesen: {e}")

        # Fallback Liste (alt + erweitert)
        return """
d1 – Lernen und Wissensanwendung
- Spielverhalten: D1310 (Symbolspiel), D1314 (So-tun-als-ob)
- Theory of Mind: D1630 (Perspektivwechsel)
- Logische Zusammenhänge: D1631, D1632
- Aufmerksamkeit: D160 (Umlenken), D161 (Aufrechterhalten)
- Auswahl treffen: D177
- Lernen durch Wiederholung: D135

d2 – Allgemeine Aufgaben und Anforderungen
- D210 (Einzelaufgabe), D220 (Mehrfachaufgaben), D230 (Tagesroutine), D250 (Verhaltensregulation)

d3 – Kommunikation
- D331 (Lautieren), D310 (Verbalsprache verstehen), D330 (Sprechen)
- D315 (Nonverbal verstehen), D335 (Nonverbal senden)
- D3350 (Körpersprache), D3351 (Zeichen/Symbole), D3352 (Piktogramme)
- D350 (Konversation), D3500-D3504 (Unterhaltung führen)

d7 – Interpersonelle Interaktionen
- D7100-D7106 (Soziale Basiskompetenzen: Mitgefühl, Respekt, Nähe/Distanz)
- D7200-D7204 (Beziehungen führen: Gefühle regulieren, Regeln beachten)
- D730 (Zweckgebundene Interaktion), D7500 (Freundschaften), D760 (Familie)
"""

    def _ensure_folder_structure(self):
        """Erstellt die Ordnerstruktur falls nicht vorhanden."""
        (self.base_path / "data").mkdir(parents=True, exist_ok=True)
        (self.base_path / "bundles" / "core").mkdir(parents=True, exist_ok=True)
        (self.base_path / "bundles" / "extended").mkdir(parents=True, exist_ok=True)
        (self.base_path / "output").mkdir(parents=True, exist_ok=True)

    # ─────────────────────────────────────────────────────────────
    # Schritt 1: Session starten
    # ─────────────────────────────────────────────────────────────

    def start_session(self) -> WorkflowSession:
        """
        Startet eine neue Workflow-Session.

        Returns:
            WorkflowSession mit eindeutiger ID
        """
        session_id = f"WF_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{secrets.token_hex(4)}"
        session = WorkflowSession(session_id=session_id)
        self._sessions[session_id] = session
        return session

    def get_session(self, session_id: str) -> Optional[WorkflowSession]:
        """Holt eine bestehende Session."""
        return self._sessions.get(session_id)

    # ─────────────────────────────────────────────────────────────
    # Schritt 2: Dateien importieren
    # ─────────────────────────────────────────────────────────────

    def import_files(
        self,
        session: WorkflowSession,
        files: List[Path],
        clear_existing: bool = True
    ) -> int:
        """
        Importiert Dateien in den data/ Ordner.

        Args:
            session: Aktive Session
            files: Liste von Dateipfaden
            clear_existing: Bestehende Dateien in data/ loeschen

        Returns:
            Anzahl importierter Dateien
        """
        session.status = "importing"
        data_dir = self.base_path / "data"

        # HINWEIS: data/ wird manuell vom Benutzer verwaltet - kein auto-cleanup

        # Dateien kopieren
        imported = 0
        for filepath in files:
            if not filepath.exists():
                continue

            if filepath.is_file():
                dest = data_dir / filepath.name
                shutil.copy2(filepath, dest)
                session.input_files.append(dest)
                imported += 1
            elif filepath.is_dir():
                # Ordner rekursiv kopieren
                dest = data_dir / filepath.name
                shutil.copytree(filepath, dest, dirs_exist_ok=True)
                session.input_files.extend(list(dest.rglob("*")))
                imported += len(list(dest.rglob("*")))

        return imported

    def import_from_folder(
        self,
        session: WorkflowSession,
        folder: Path,
        clear_existing: bool = True
    ) -> int:
        """
        Importiert alle Dateien aus einem Ordner.

        Args:
            session: Aktive Session
            folder: Quellordner
            clear_existing: Bestehende Dateien loeschen

        Returns:
            Anzahl importierter Dateien
        """
        if not folder.exists() or not folder.is_dir():
            session.error_message = f"Ordner nicht gefunden: {folder}"
            session.status = "error"
            return 0

        files = list(folder.iterdir())
        return self.import_files(session, files, clear_existing)

    # ─────────────────────────────────────────────────────────────
    # Schritt 3: Temporaeres Anonymisierungsprofil
    # ─────────────────────────────────────────────────────────────

    def create_temp_profile(
        self,
        session: WorkflowSession,
        client_name: str,
        geburtsdatum: str = "01.01.2010",
        additional_terms: Optional[List[str]] = None,
        whitelist: Optional[List[str]] = None,
        scan_folder: Optional[Path] = None,
        parent_names: Optional[List[str]] = None,
        client_address: Optional[str] = None
    ) -> TempAnonymProfile:
        """
        Erstellt ein temporaeres Anonymisierungsprofil.

        Args:
            session: Aktive Session
            client_name: Echter Name des Klienten
            geburtsdatum: Geburtsdatum (dd.mm.yyyy)
            additional_terms: Weitere zu anonymisierende Begriffe
            whitelist: Begriffe die NICHT anonymisiert werden
            scan_folder: Ordner fuer automatische Erkennung von Tel/Mail/Adressen
            parent_names: Liste der Elternnamen (z.B. ["Maria Mustermann", "Hans Mustermann"])
            client_address: Adresse des Klienten (wird anonymisiert)

        Returns:
            TempAnonymProfile (nicht persistent!)
        """
        session.status = "anonymizing"

        # Whitelist
        wl = set(DEFAULT_WHITELIST)
        if whitelist:
            wl.update(whitelist)

        # Vorname extrahieren fuer Gender-Erkennung
        real_parts = client_name.strip().split()
        original_vorname = real_parts[0] if real_parts else ""

        # Gender erkennen
        detected_gender = _detect_gender(original_vorname)
        print(f"[INFO] Gender-Erkennung: {original_vorname} -> {detected_gender}")

        # Tarnname mit passendem Geschlecht generieren
        tarnname = _generate_tarnname(
            self._used_tarnnames,
            gender=detected_gender,
            original_vorname=original_vorname
        )
        self._used_tarnnames.add(tarnname)

        # Falsches Geburtsdatum
        fake_geb, _ = _generate_fake_date_same_age(geburtsdatum)

        # Mappings aufbauen
        mappings: Dict[str, Dict[str, str]] = {
            "names": {},
            "dates": {},
            "addresses": {},
            "phones": {},
            "emails": {},
            "misc": {}
        }

        # Hauptname
        tarn_parts = tarnname.strip().split()

        mappings["names"][client_name] = tarnname

        # Varianten fuer Namen (CAPS, etc.)
        name_variants = {}
        for real, fake in mappings["names"].items():
            name_variants[real.upper()] = fake.upper()
        mappings["names"].update(name_variants)

        # Vorname und Nachname einzeln
        if len(real_parts) >= 2 and len(tarn_parts) >= 2:
            vorname_real = real_parts[0]
            vorname_fake = tarn_parts[0]
            nachname_real = real_parts[-1]
            nachname_fake = tarn_parts[-1]
            
            mappings["names"][vorname_real] = vorname_fake
            mappings["names"][vorname_real.upper()] = vorname_fake.upper()
            
            mappings["names"][nachname_real] = nachname_fake
            mappings["names"][nachname_real.upper()] = nachname_fake.upper()

        # Bei "Nachname, Vorname" Format
        if ", " in client_name:
            parts = client_name.split(", ")
            if len(parts) == 2:
                # Nachname schon oben drin, aber hier nochmal explizit falls anders
                mappings["names"][parts[0]] = tarn_parts[-1]  # Nachname
                mappings["names"][parts[1]] = tarn_parts[0]   # Vorname
                mappings["names"][parts[0].upper()] = tarn_parts[-1].upper()
                mappings["names"][parts[1].upper()] = tarn_parts[0].upper()

        # Geburtsdatum
        mappings["dates"][geburtsdatum] = fake_geb

        # Elternnamen anonymisieren (mit Gender-Matching)
        if parent_names:
            for parent_name in parent_names:
                if parent_name and parent_name not in wl:
                    parent_parts = parent_name.strip().split()
                    parent_vorname = parent_parts[0] if parent_parts else ""
                    parent_gender = _detect_gender(parent_vorname)

                    # Eltern-Tarnname mit passendem Geschlecht
                    parent_tarnname = _generate_tarnname(
                        self._used_tarnnames,
                        gender=parent_gender,
                        original_vorname=parent_vorname
                    )
                    self._used_tarnnames.add(parent_tarnname)
                    parent_tarn_parts = parent_tarnname.strip().split()

                    # Vollstaendigen Namen
                    mappings["names"][parent_name] = parent_tarnname
                    mappings["names"][parent_name.upper()] = parent_tarnname.upper()

                    # Auch einzelne Teile (Vorname/Nachname)
                    if len(parent_parts) >= 2 and len(parent_tarn_parts) >= 2:
                        # Vorname
                        if parent_parts[0] not in mappings["names"]:
                            mappings["names"][parent_parts[0]] = parent_tarn_parts[0]
                            mappings["names"][parent_parts[0].upper()] = parent_tarn_parts[0].upper()
                        # Nachname (falls anders als Klient)
                        if parent_parts[-1] not in mappings["names"]:
                            mappings["names"][parent_parts[-1]] = parent_tarn_parts[-1]
                            mappings["names"][parent_parts[-1].upper()] = parent_tarn_parts[-1].upper()

                    print(f"[INFO] Elternname: {parent_name} -> {parent_tarnname}")

        # Klienten-Adresse explizit anonymisieren
        if client_address and client_address not in wl:
            fake_addr = _generate_fake_address()
            mappings["addresses"][client_address] = fake_addr
            print(f"[INFO] Klienten-Adresse: {client_address} -> {fake_addr}")

        # Zusaetzliche Begriffe
        if additional_terms:
            for term in additional_terms:
                if term not in wl:
                    mappings["misc"][term] = "[REDACTED]"

        # Automatische Erkennung von Telefon, E-Mail, Adressen
        if scan_folder and scan_folder.exists():
            print(f"[INFO] Scanne Ordner nach sensiblen Daten: {scan_folder}")
            anonymizer = DocumentAnonymizer()
            scanned = anonymizer.scan_folder_for_sensitive_data(str(scan_folder))

            # Telefonnummern
            for phone in scanned.get("phones", []):
                if phone not in wl:
                    fake_phone = _generate_fake_phone()
                    mappings["phones"][phone] = fake_phone
                    print(f"  [SCAN] Telefon: {phone} -> {fake_phone}")

            # E-Mail-Adressen
            for email in scanned.get("emails", []):
                if email not in wl:
                    fake_email = _generate_fake_email()
                    mappings["emails"][email] = fake_email
                    print(f"  [SCAN] E-Mail: {email} -> {fake_email}")

            # Straßenadressen
            for addr in scanned.get("addresses", []):
                if addr not in wl:
                    fake_addr = _generate_fake_address()
                    mappings["addresses"][addr] = fake_addr
                    print(f"  [SCAN] Adresse: {addr} -> {fake_addr}")

            print(f"[INFO] Scan abgeschlossen: {len(scanned.get('phones', []))} Tel, "
                  f"{len(scanned.get('emails', []))} Mail, {len(scanned.get('addresses', []))} Adressen")

        # Temporaeres Verzeichnis
        temp_dir = Path(tempfile.mkdtemp(prefix="bach_report_"))

        profile = TempAnonymProfile(
            tarnname=tarnname,
            fake_geburtsdatum=fake_geb,
            mappings=mappings,
            original_name=client_name,
            _temp_dir=temp_dir
        )

        session.profile = profile
        return profile

    def auto_detect_profile(
        self,
        session: WorkflowSession,
        folder_name: str
    ) -> TempAnonymProfile:
        """
        Erstellt Profil automatisch aus Ordnernamen.

        Erwartet Format: "Nachname, Vorname" oder "Vorname Nachname"

        Args:
            session: Aktive Session
            folder_name: Name des Klienten-Ordners

        Returns:
            TempAnonymProfile
        """
        # Name aus Ordnernamen extrahieren
        if ", " in folder_name:
            parts = folder_name.split(", ")
            name = f"{parts[1]} {parts[0]}"  # "Vorname Nachname"
        else:
            name = folder_name

        return self.create_temp_profile(session, name.strip())

    # ─────────────────────────────────────────────────────────────
    # Schritt 4: Anonymisieren und in Bundles sortieren
    # ─────────────────────────────────────────────────────────────

    def anonymize_to_bundles(
        self,
        session: WorkflowSession,
        include_extended: bool = False
    ) -> CollectorResult:
        """
        Anonymisiert Dokumente und sortiert sie in bundles/.

        Args:
            session: Aktive Session mit Profil
            include_extended: Auch EXTENDED Dokumente verarbeiten

        Returns:
            CollectorResult mit Scan-Ergebnissen
        """
        if not session.profile:
            session.error_message = "Kein Profil erstellt"
            session.status = "error"
            return CollectorResult()

        data_dir = self.base_path / "data"
        core_dir = self.base_path / "bundles" / "core"
        extended_dir = self.base_path / "bundles" / "extended"

        # Bundles-Ordner leeren - Robuste Variante (deep_cleanup)
        for d in [core_dir, extended_dir]:
            self._deep_cleanup(d)

        # Dokumente scannen
        collector = DocumentCollector()
        result = collector.scan_folder(str(data_dir))

        # Anonymizer
        anonymizer = DocumentAnonymizer()
        anonym_profile = session.profile.to_anonym_profile()

        # Dateien verarbeiten
        for doc in result.documents:
            if doc.category == DocumentCategory.SKIP:
                continue
            if doc.category == DocumentCategory.EXTENDED and not include_extended:
                continue

            # Zielordner bestimmen
            if doc.category == DocumentCategory.CORE:
                dest_dir = core_dir
            elif doc.category == DocumentCategory.STUFE2:
                dest_dir = core_dir  # STUFE2 geht auch in core
            else:
                dest_dir = extended_dir

            # Datei kopieren
            dest_file = dest_dir / doc.path.name
            shutil.copy2(doc.path, dest_file)

            # Anonymisieren
            anonymizer.anonymize_file(str(dest_file), anonym_profile)

            # Dateinamen anonymisieren
            self._anonymize_filename(dest_file, session.profile)

        return result

    def anonymize_to_bundles_from_folder(
        self,
        session: WorkflowSession,
        source_folder: Path,
        include_extended: bool = False
    ) -> CollectorResult:
        """
        Anonymisiert Dokumente aus einem spezifischen Ordner und sortiert sie in bundles/.

        Diese Variante scannt einen explizit angegebenen Ordner statt data/.
        Nützlich wenn die Daten bereits in data/ liegen und nicht kopiert werden sollen.

        Args:
            session: Aktive Session mit Profil
            source_folder: Ordner mit Quelldokumenten
            include_extended: Auch EXTENDED Dokumente verarbeiten

        Returns:
            CollectorResult mit Scan-Ergebnissen
        """
        if not session.profile:
            session.error_message = "Kein Profil erstellt"
            session.status = "error"
            return CollectorResult()

        core_dir = self.base_path / "bundles" / "core"
        extended_dir = self.base_path / "bundles" / "extended"

        # Bundles-Ordner leeren - Robuste Variante (deep_cleanup)
        for d in [core_dir, extended_dir]:
            self._deep_cleanup(d)

        # Dokumente aus dem angegebenen Ordner scannen
        collector = DocumentCollector()
        result = collector.scan_folder(str(source_folder))

        print(f"[DEBUG] Scan von {source_folder}: {result.core_count} CORE, {result.stufe2_count} STUFE2, {result.extended_count} EXTENDED")

        # Anonymizer
        anonymizer = DocumentAnonymizer()
        anonym_profile = session.profile.to_anonym_profile()

        # Dateien verarbeiten
        processed = 0
        for doc in result.documents:
            if doc.category == DocumentCategory.SKIP:
                continue
            if doc.category == DocumentCategory.EXTENDED and not include_extended:
                continue

            # Zielordner bestimmen
            if doc.category == DocumentCategory.CORE:
                dest_dir = core_dir
            elif doc.category == DocumentCategory.STUFE2:
                dest_dir = core_dir  # STUFE2 geht auch in core
            else:
                dest_dir = extended_dir

            # Datei kopieren
            dest_file = dest_dir / doc.path.name
            try:
                shutil.copy2(doc.path, dest_file)
                processed += 1
            except Exception as e:
                print(f"[WARN] Konnte {doc.path.name} nicht kopieren: {e}")
                continue

            # Anonymisieren
            try:
                anonymizer.anonymize_file(str(dest_file), anonym_profile)
            except Exception as e:
                print(f"[WARN] Konnte {dest_file.name} nicht anonymisieren: {e}")

            # Dateinamen anonymisieren
            self._anonymize_filename(dest_file, session.profile)

        print(f"[DEBUG] {processed} Dateien in bundles/ verarbeitet")
        return result

    def _anonymize_filename(self, filepath: Path, profile: TempAnonymProfile) -> Path:
        """Anonymisiert Dateinamen."""
        filename = filepath.stem
        suffix = filepath.suffix

        mappings = profile.get_all_mappings()
        sorted_mappings = sorted(mappings.items(), key=lambda x: len(x[0]), reverse=True)

        new_filename = filename
        changed = False

        for old, new in sorted_mappings:
            if old in new_filename:
                new_filename = new_filename.replace(old, new)
                changed = True

        if changed:
            new_path = filepath.parent / f"{new_filename}{suffix}"
            
            # Handle duplicate filenames (Windows safety)
            if new_path.exists() and new_path.resolve() != filepath.resolve():
                counter = 1
                while new_path.exists():
                    new_path = filepath.parent / f"{new_filename}_{counter}{suffix}"
                    counter += 1
            
            if new_path != filepath:
                try:
                    filepath.rename(new_path)
                    return new_path
                except FileExistsError:
                    # Final fallback if still collision (race condition)
                    import uuid
                    new_path = filepath.parent / f"{new_filename}_{uuid.uuid4().hex[:4]}{suffix}"
                    filepath.rename(new_path)
                    return new_path

        return filepath

    # ─────────────────────────────────────────────────────────────
    # Schritt 5: Prompt generieren
    # ─────────────────────────────────────────────────────────────

    def extract_bundle_text(self, session: WorkflowSession, from_data: bool = False) -> TextBundle:
        """
        Extrahiert Text aus Dokumenten.

        Args:
            session: Aktive Session
            from_data: True = aus data/ scannen, False = aus bundles/core

        Returns:
            TextBundle mit extrahiertem Text
        """
        if from_data:
            # Direkt aus data/ scannen (wenn keine Anonymisierung zu bundles/ erfolgt)
            scan_dir = self.base_path / "data"
        else:
            # Aus bundles/core (nach anonymize_to_bundles)
            scan_dir = self.base_path / "bundles" / "core"

        collector = DocumentCollector()
        result = collector.scan_folder(str(scan_dir))

        # Mit Anonymisierung falls Profil vorhanden
        anonym_profile = session.profile.to_anonym_profile() if session.profile else None
        bundle = collector.extract_bundle(
            result.documents,
            include_extended=True,
            anonym_profile=anonym_profile
        )

        session.bundle = bundle
        return bundle

    def generate_prompt(
        self,
        session: WorkflowSession,
        include_wissensdatenbank: bool = True,
        berichtszeitraum: str = "01.01.2025 - 31.12.2025",
        custom_instructions: str = ""
    ) -> str:
        """
        Generiert den LLM-Prompt mit Kontext.

        Args:
            session: Aktive Session
            include_wissensdatenbank: Wissensdatenbank einbeziehen
            berichtszeitraum: Zeitraum fuer den Bericht
            custom_instructions: Zusaetzliche Anweisungen

        Returns:
            Fertiger Prompt fuer LLM
        """
        session.status = "generating"

        # Bundle extrahieren falls noch nicht geschehen
        # Prüfe ob bundles/core Dateien hat, sonst aus data/ lesen
        core_dir = self.base_path / "bundles" / "core"
        has_bundles = core_dir.exists() and any(core_dir.iterdir())

        if not session.bundle:
            self.extract_bundle_text(session, from_data=not has_bundles)

        bundle = session.bundle
        tarnname = session.profile.tarnname if session.profile else "Klient"

        # Wissensdatenbank-Kontext
        kb_context = ""
        if include_wissensdatenbank and WISSENSDATENBANK.exists():
            kb_context = self._search_wissensdatenbank([
                "ICF", "Autismus", "Foerderplanung", "Methoden"
            ])

        # Prompt zusammenbauen
        prompt = f"""=== FÖRDERBERICHT ERSTELLEN ===

Klient (anonymisiert): {tarnname}
Berichtszeitraum: {berichtszeitraum}

=== DOKUMENTEN-BUNDLE (CORE + STUFE2) ===

{bundle.core_text}

{bundle.stufe2_text}

"""

        if kb_context:
            prompt += f"""=== FACHLICHER KONTEXT (Wissensdatenbank) ===

{kb_context}

"""

        prompt += """=== AUFGABE ===

Erstelle einen strukturierten JSON-Output für einen ICF-basierten Förderbericht.

WICHTIGE REGELN:
1. Antworte NUR mit validem JSON - KEIN Markdown, KEINE Erklärungen
2. Verwende den anonymisierten Namen wie angegeben
3. Formuliere im fachlichen Berichtsstil (3. Person, sachlich)
4. Beziehe dich auf konkrete Beobachtungen aus den Protokollen
5. ICF-Codes: Verwende AUSSCHLIESSLICH die unten gelisteten ICF-Codes (d1-d9).
   - D880 existiert NICHT. Verwende D1310 oder D1314 für Spielverhalten.

STAMMDATEN-EXTRAKTION:
- landkreis: Aus Kostenzusage/Hilfeplan ablesen - "Loerrach" oder "Waldshut"
- zeichen_amt: Aktenzeichen des Amts (z.B. "51.2.1-XXX")
- foerderungsbeginn: Wann begann die Foerderung bei Foerderstelle? (aus Aktendeckblatt)
- kostenzusage_ende: Datum Bewilligungsende (dd.mm.yyyy).
  * WICHTIG: Suche nach Phrasen wie "Kostenzusage bis...", "befristet bis...", "Ende der Förderung am..."
  * Wenn Text wie "bis Ende Februar 2026", konvertiere zu "28.02.2026".
  * Wenn Text wie "bis Ende September 2025", konvertiere zu "30.09.2025".
  * Das Datum im Aktendeckblatt ist oft das Erstellungsdatum, NICHT das Ende - prüfe die Kostenzusage!
- diagnosen: Fuer JEDE Diagnose erfassen:
  * icd_code: z.B. "F84.5"
  * bezeichnung: z.B. "Asperger-Syndrom"
  * diagnostiker: Wer hat diagnostiziert? (z.B. "Dr. Ritter-Gekeler", "Kreiskrankenhaus Loerrach")
  * datum: Wann? (Jahr oder volles Datum)
  * quelle: Aus welchem Dokument stammt die Info? (z.B. "Bericht vom 12.03.2019")

VERFÜGBARE ICF-BEREICHE:
{self._get_icf_reference()}

JSON-SCHEMA:
{
  "stammdaten": {
    "name": "Anonymisierter Name",
    "geburtsdatum": "TT.MM.JJJJ",
    "berichtszeitraum": "TT.MM.JJJJ - TT.MM.JJJJ",
    "stundenumfang": "75 Minuten wöchentlich",
    "landkreis": "Lörrach oder Waldshut",
    "kostenzusage_ende": "TT.MM.JJJJ",
    "weiterbewilligung_ab": "TT.MM.JJJJ",
    "foerderungsbeginn": "TT.MM.JJJJO",
    "diagnosen": [{"icd_code": "F84.5", "bezeichnung": "...", "diagnostiker": "...", "datum": "..."}],
    "therapieangebot": {"einzelfoerderung": true, "gruppenfoerderung": false},
    "netzwerkarbeit": {"elternarbeit": true, "schule": false}
  },
  "foerderziele": [
    {
      "icf_code": "D350",
      "zielformulierung": "Konversationsfähigkeiten verbessern",
      "ist_stand": "Konkrete Beobachtung aus Protokollen (2-3 Sätze)",
      "erreicht": 2,
      "grund_nichterreichung": null
    }
  ],
  "bereiche": {
    "selbstversorgung": {"aktiv": true, "text": "Individualisierter Text zum Bereich..."},
    "gesellschaft_freizeit": {"aktiv": false, "text": ""}
  },
  "umweltfaktoren": {
    "kommunikation": {"aktiv": true, "text": "..."},
    "wahrnehmung": {"aktiv": false, "text": ""}
  },
  "besondere_faehigkeiten": "Stärken des Klienten",
  "neue_ziele": [
    {"icf_code": "D7503", "zielformulierung": "...", "ist_stand": "..."}
  ],
  "empfehlung": {
    "verlaengerung": true,
    "begruendung": ["therapieziele_nicht_erreicht"],
    "umfang": "75 Minuten wöchentlich",
    "sonstige_empfehlung": ""
  },
  "abschluss": {
    "aktuelle_entwicklungen": "Zusammenfassender Absatz...",
    "bedingungsmodell": "Wenn X, dann Y - Prognose..."
  }
}

ZIELERREICHUNG:
- 1 = nicht erreicht
- 2 = teilweise erreicht
- 3 = vollständig erreicht

BEGRÜNDUNGEN (für empfehlung.begruendung Array):
- "therapieziele_nicht_erreicht"
- "uebergang_lebensbereich"
- "krise_erhoehter_bedarf"
- "ziele_erreicht" (bei Beendigung)

JSON:
"""

        if custom_instructions:
            prompt += f"""=== ZUSÄTZLICHE ANWEISUNGEN ===

{custom_instructions}

"""

        return prompt

    def _search_wissensdatenbank(self, topics: List[str], max_chars: int = 3000) -> str:
        """Durchsucht die Wissensdatenbank nach relevanten Inhalten."""
        results = []
        total_chars = 0

        for topic in topics:
            # Direkte Topic-Ordner
            topic_dir = WISSENSDATENBANK / topic
            if topic_dir.exists():
                for f in topic_dir.glob("*.md"):
                    if total_chars >= max_chars:
                        break
                    try:
                        content = f.read_text(encoding="utf-8")[:1000]
                        results.append(f"[{topic}/{f.name}]\n{content}")
                        total_chars += len(content)
                    except:
                        pass

            # Suche in allen .md Dateien
            for f in WISSENSDATENBANK.rglob("*.md"):
                if total_chars >= max_chars:
                    break
                try:
                    content = f.read_text(encoding="utf-8")
                    if topic.lower() in content.lower():
                        excerpt = content[:500]
                        results.append(f"[{f.relative_to(WISSENSDATENBANK)}]\n{excerpt}")
                        total_chars += len(excerpt)
                except:
                    pass

        return "\n\n".join(results[:10])  # Max 10 Eintraege

    # ─────────────────────────────────────────────────────────────
    # Schritt 6: Bericht generieren
    # ─────────────────────────────────────────────────────────────

    def generate_report(
        self,
        session: WorkflowSession,
        llm_response: str,
        auto_deanonymize: bool = True,
        report_data: Optional[Dict] = None
    ) -> Path:
        """
        Generiert das finale Word-Dokument.

        Versucht zuerst JSON zu parsen und den Generator zu nutzen.
        Fallback: Direktes Template-Fuellen mit Markdown-Text.

        Args:
            session: Aktive Session
            llm_response: Antwort vom LLM (JSON oder Markdown-Text)
            auto_deanonymize: Automatisch de-anonymisieren
            report_data: Optional - strukturierte Daten fuer Platzhalter

        Returns:
            Pfad zum fertigen Dokument
        """
        original_name = session.profile.original_name if session.profile else "Klient"

        # Versuche JSON zu parsen
        json_data = self._extract_json_from_response(llm_response)

        if json_data:
            # JSON gefunden -> nutze den Generator
            return self._generate_from_json(session, json_data, auto_deanonymize)

        # Fallback: Alter Markdown-basierter Ansatz
        return self._generate_from_markdown(session, llm_response, auto_deanonymize, report_data)

    def _extract_json_from_response(self, llm_response: str) -> Optional[Dict]:
        """
        Extrahiert JSON aus der LLM-Response.

        Behandelt verschiedene Formate:
        - Reines JSON
        - JSON in Markdown-Codeblock (```json ... ```)
        - JSON nach Text
        """
        import re

        text = llm_response.strip()

        # Fall 1: Markdown-Codeblock
        json_block_match = re.search(r'```(?:json)?\s*\n?([\s\S]*?)\n?```', text)
        if json_block_match:
            try:
                return json.loads(json_block_match.group(1).strip())
            except json.JSONDecodeError:
                pass

        # Fall 2: Reines JSON (beginnt mit {)
        if text.startswith('{'):
            try:
                return json.loads(text)
            except json.JSONDecodeError:
                pass

        # Fall 3: JSON irgendwo im Text
        json_match = re.search(r'\{[\s\S]*\}', text)
        if json_match:
            try:
                return json.loads(json_match.group(0))
            except json.JSONDecodeError:
                pass

        return None

    def _generate_from_json(
        self,
        session: WorkflowSession,
        json_data: Dict,
        auto_deanonymize: bool
    ) -> Path:
        """Generiert Bericht aus JSON-Daten mit dem Generator."""
        try:
            # Generator importieren
            import sys
            sys.path.insert(0, str(BACH_BASE / "system"))
            from skills._experts.report_generator.generator import fill_template

            # De-Anonymisierung der JSON-Daten
            if auto_deanonymize and session.profile:
                json_data = self._deanonymize_json(json_data, session.profile)
                original_name = session.profile.original_name
            else:
                original_name = session.profile.tarnname if session.profile else "Klient"

            # Output-Pfad
            safe_name = "".join(c for c in original_name if c.isalnum() or c in " -_").strip()
            timestamp = datetime.now().strftime("%Y%m%d")
            output_name = f"Foerderbericht_{safe_name}_{timestamp}.docx"
            output_path = self.base_path / "output" / output_name

            # Generator nutzen
            result = fill_template(
                template_path=str(TEMPLATE_PATH),
                report_data=json_data,
                output_path=str(output_path)
            )

            if result.success:
                session.output_path = output_path
                session.status = "completed"
                return output_path
            else:
                session.error_message = f"Generator-Fehler: {'; '.join(result.errors)}"
                session.status = "error"
                return None

        except Exception as e:
            session.error_message = f"JSON-Generierung fehlgeschlagen: {e}"
            session.status = "error"
            # Fallback zu Markdown
            return self._generate_from_markdown(
                session,
                json.dumps(json_data, ensure_ascii=False, indent=2),
                auto_deanonymize,
                None
            )

    def _deanonymize_json(self, data: Dict, profile: TempAnonymProfile) -> Dict:
        """De-anonymisiert alle String-Werte in einem JSON-Dict rekursiv."""
        reverse_mappings = profile.get_reverse_mappings()
        sorted_mappings = sorted(reverse_mappings.items(), key=lambda x: len(x[0]), reverse=True)

        def deanon_value(value):
            if isinstance(value, str):
                result = value
                for tarn, original in sorted_mappings:
                    result = result.replace(tarn, original)
                return result
            elif isinstance(value, dict):
                return {k: deanon_value(v) for k, v in value.items()}
            elif isinstance(value, list):
                return [deanon_value(item) for item in value]
            else:
                return value

        return deanon_value(data)

    def _generate_from_markdown(
        self,
        session: WorkflowSession,
        llm_response: str,
        auto_deanonymize: bool,
        report_data: Optional[Dict]
    ) -> Path:
        """Fallback: Generiert Bericht aus Markdown-Text."""
        try:
            from docx import Document
        except ImportError:
            session.error_message = "python-docx nicht installiert"
            session.status = "error"
            return None

        # Text de-anonymisieren
        if auto_deanonymize and session.profile:
            final_text = self._deanonymize_text(llm_response, session.profile)
            original_name = session.profile.original_name
        else:
            final_text = llm_response
            original_name = session.profile.tarnname if session.profile else "Klient"

        # Basis-Platzhalter aus Session/Profil
        placeholders = self._build_placeholders(session, final_text, report_data)

        # Template laden
        if TEMPLATE_PATH.exists():
            doc = Document(str(TEMPLATE_PATH))
            # Platzhalter in Paragraphen ersetzen
            self._replace_placeholders_in_doc(doc, placeholders)

            # ICF-Tabelle bereinigen wenn Foerderziele vorhanden
            if report_data and len(doc.tables) >= 3:
                foerderziele = report_data.get("foerderziele", []) or report_data.get("förderziele", [])
                if foerderziele:
                    try:
                        svc = WordTemplateService()
                        icf_table = doc.tables[2]  # ICF-Tabelle ist Tabelle 3 (0-basiert: Index 2)
                        svc.fill_icf_placeholders_and_cleanup(icf_table, foerderziele, header_rows=1)
                    except Exception as e:
                        # Logging wäre gut, aber Fehler nicht kritisch
                        pass
        else:
            # Fallback: Neues Dokument mit rohem Text
            doc = Document()
            doc.add_heading("Förderbericht", 0)
            for para_text in final_text.split("\n\n"):
                if para_text.strip():
                    doc.add_paragraph(para_text.strip())

        # Speichern
        safe_name = "".join(c for c in original_name if c.isalnum() or c in " -_").strip()
        timestamp = datetime.now().strftime("%Y%m%d")
        output_name = f"Foerderbericht_{safe_name}_{timestamp}.docx"
        output_path = self.base_path / "output" / output_name

        doc.save(str(output_path))

        session.output_path = output_path
        session.status = "completed"

        return output_path

    def _deanonymize_text(self, text: str, profile: TempAnonymProfile) -> str:
        """De-anonymisiert Text (Tarnname -> Original)."""
        reverse_mappings = profile.get_reverse_mappings()
        sorted_mappings = sorted(reverse_mappings.items(), key=lambda x: len(x[0]), reverse=True)

        result = text
        for tarn, original in sorted_mappings:
            result = result.replace(tarn, original)

        return result

    def _build_placeholders(
        self,
        session: WorkflowSession,
        llm_text: str,
        report_data: Optional[Dict] = None
    ) -> Dict[str, str]:
        """
        Erstellt Platzhalter-Mapping aus Session und LLM-Text.

        Die Platzhalter im Template haben das Format {{NAME}}.
        """
        import re

        placeholders = {}

        # Basis-Daten aus Profil
        if session.profile:
            original_name = session.profile.original_name
            # Name aufteilen
            name_parts = original_name.split()
            if len(name_parts) >= 2:
                vorname = name_parts[0]
                nachname = " ".join(name_parts[1:])
            else:
                vorname = original_name
                nachname = ""

            # Bei "Nachname, Vorname" Format
            if ", " in original_name:
                parts = original_name.split(", ")
                nachname = parts[0]
                vorname = parts[1] if len(parts) > 1 else ""

            placeholders["NAME"] = original_name
            placeholders["VORNAME"] = vorname
            placeholders["NACHNAME"] = nachname

            # Geburtsdatum (aus Mappings extrahieren)
            for orig, tarn in session.profile.mappings.get("dates", {}).items():
                placeholders["GEBURTSDATUM"] = orig
                break

        # Datum/Zeitraum
        placeholders["DATUM"] = datetime.now().strftime("%d.%m.%Y")
        placeholders["BERICHTSZEITRAUM"] = "01.01.2025 - 31.12.2025"

        # Wenn strukturierte Daten übergeben wurden
        if report_data:
            placeholders.update(report_data)

        # LLM-Text parsen und Abschnitte extrahieren
        sections = self._parse_llm_sections(llm_text)
        placeholders.update(sections)

        # Erweiterte Platzhalter aus dem Generator (damit sie im Fallback zumindest leer sind)
        # und nicht als {{TAG}} im Dokument bleiben
        extended_defaults = {
            "{{SACHBEARBEITER}}": "",
            "{{KOSTENZUSAGE_ENDE}}": "",
            "{{WEITERBEWILLIGUNG_AB}}": "",
            "{{FOERDERUNGSBEGINN}}": "",
            "{{DIAGNOSE_1}}": "",
            "{{THERAPEUT_EINZEL}}": "",  # User-specific: set your name here
            "{{THERAPEUT_GRUPPE}}": "",
            "{{KOPIE_AN}}": "",
            "{{UMFANG}}": "",
            "{{SONSTIGE_EMPFEHLUNG}}": "",
            "{{Weiterbewilligung oder Beendigung}}": "Weiterbewilligung",
            "{{Empfehlung}}": "die Weiterbewilligung",
            "{{Zeichen Amtschreiben}}": "",
            "{{AKTUELLE_ENTWICKLUNGEN}}": sections.get("AKTUELLE_ENTWICKLUNGEN", ""),
            "{{BEDINGUNGSMODELL}}": "",
            "{{BESONDERE_FAEHIGKEITEN}}": "",
            "{{AMT_ADRESSE}}": "",
            "{{LANDKREIS}}": "",
            "{{BEREICH_MOBILITAET}}": "",
            "{{BEREICH_SELBSTVERSORGUNG}}": sections.get("BEREICH_SELBSTVERSORGUNG", ""),
            "{{BEREICH_HAUSHALT}}": "",
            "{{BEREICH_LEBENSBEREICHE}}": "",
            "{{BEREICH_GESELLSCHAFT_FREIZEIT}}": "",
            "{{UMWELT_WAHRNEHMUNG}}": "",
            "{{UMWELT_KOMMUNIKATION}}": sections.get("BEREICH_KOMMUNIKATION", ""), # Mapping auf Kommunikation
            "{{UMWELT_KOGNITION_MOTORIK}}": "",
            "{{UMWELT_MEDIZINISCHES}}": ""
        }
        
        # Nur Keys uebernehmen die noch nicht existieren
        for key, val in extended_defaults.items():
            if key not in placeholders and key.strip("{}") not in placeholders:
                # Key ist z.B. "{{SACHBEARBEITER}}", placeholders hat oft nur "SACHBEARBEITER" (ohne klammern)
                # WordTemplateService._replace_placeholders_in_doc erwartet keys MIT klammern fuer replace_in_text
                # ABER die `placeholders` map hier mischt keys mit und ohne Klammern?
                # Schauen wir _replace_placeholders_in_doc an: 
                #   key = match.group(1) -> WITHOUT braces
                #   placeholders.get(key)
                
                # Also muessen die Keys OHNE Klammern in die Map!
                clean_key = key.replace("{", "").replace("}", "")
                placeholders[clean_key] = val

        return placeholders

    def _parse_llm_sections(self, llm_text: str) -> Dict[str, str]:
        """
        Parst den LLM-Text und extrahiert Abschnitte fuer Platzhalter.

        Erkennt Markdown-Ueberschriften und ordnet sie Template-Platzhaltern zu.
        """
        import re

        sections = {}

        # Mapping von Ueberschriften zu Platzhaltern
        section_mapping = {
            "aktuelle situation": "AKTUELLE_SITUATION",
            "lebenssituation": "AKTUELLE_ENTWICKLUNGEN",
            "schulische": "AKTUELLE_ENTWICKLUNGEN",
            "familiäre": "AKTUELLE_ENTWICKLUNGEN",
            "d3": "BEREICH_KOMMUNIKATION",
            "kommunikation": "BEREICH_KOMMUNIKATION",
            "d5": "BEREICH_SELBSTVERSORGUNG",
            "selbstversorgung": "BEREICH_SELBSTVERSORGUNG",
            "d7": "BEREICH_INTERAKTION",
            "interpersonell": "BEREICH_INTERAKTION",
            "beziehung": "BEREICH_INTERAKTION",
            "d2": "BEREICH_AUFGABEN",
            "allgemeine aufgaben": "BEREICH_AUFGABEN",
            "zusammenfassung": "ZUSAMMENFASSUNG",
            "empfehlung": "Empfehlung",
            "gesamteinschätzung": "ZUSAMMENFASSUNG",
        }

        # Aktuellen Abschnitt tracken
        current_section = None
        current_content = []

        for line in llm_text.split("\n"):
            # Markdown-Ueberschrift erkennen
            if line.startswith("#"):
                # Vorherigen Abschnitt speichern
                if current_section and current_content:
                    sections[current_section] = "\n".join(current_content).strip()
                    current_content = []

                # Neuen Abschnitt bestimmen
                header = re.sub(r"^#+\s*", "", line).lower().strip()
                # Fix: Nur führende Ziffern/Punkte entfernen (z.B. "1. " oder "D3. ")
                header_clean = re.sub(r"^[\d\w]+\.\s*", "", header).strip()

                current_section = None
                for key, placeholder in section_mapping.items():
                    # Prüfe sowohl den originalen Header als auch den bereinigten
                    if key in header or key in header_clean:
                        current_section = placeholder
                        break
            elif current_section:
                current_content.append(line)

        # Letzten Abschnitt speichern
        if current_section and current_content:
            sections[current_section] = "\n".join(current_content).strip()

        # Gesamttext als Fallback
        sections["BERICHT_TEXT"] = llm_text

        return sections

    def _replace_placeholders_in_doc(self, doc, placeholders: Dict[str, str]):
        """
        Ersetzt {{PLATZHALTER}} im gesamten Word-Dokument.
        """
        import re

        def replace_in_text(text: str) -> str:
            """Ersetzt alle {{XXX}} Platzhalter (auch mit Leerzeichen)."""
            def replacer(match):
                key = match.group(1).strip()
                return placeholders.get(key, match.group(0))  # Original behalten wenn nicht gefunden

            # FIX: [^}]+ statt \w+ erlaubt Leerzeichen und Sonderzeichen in Platzhaltern
            return re.sub(r"\{\{([^}]+)\}\}", replacer, text)

        def replace_in_paragraph(paragraph):
            """Ersetzt Platzhalter in einem Paragraph."""
            # Volltext des Paragraphs
            full_text = paragraph.text
            if "{{" not in full_text:
                return

            new_text = replace_in_text(full_text)
            if new_text != full_text:
                # Text ersetzen (einfache Methode)
                # Alle Runs löschen und neuen Text setzen
                for run in paragraph.runs:
                    run.text = ""
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                else:
                    paragraph.add_run(new_text)

        def replace_in_table(table):
            """Ersetzt Platzhalter in Tabellenzellen."""
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph)

        # Hauptdokument
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph)

        # Tabellen
        for table in doc.tables:
            replace_in_table(table)

        # Header und Footer
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for paragraph in header.paragraphs:
                        replace_in_paragraph(paragraph)
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for paragraph in footer.paragraphs:
                        replace_in_paragraph(paragraph)

    # ─────────────────────────────────────────────────────────────
    # Schritt 7: Cleanup
    # ─────────────────────────────────────────────────────────────

    def _deep_cleanup(self, directory: Path, delete_root: bool = False):
        """
        Löscht rekursiv alle Dateien und leere Ordner in einem Verzeichnis.

        Im Gegensatz zu shutil.rmtree funktioniert dies auch bei OneDrive-Sync,
        da zuerst alle Dateien gelöscht werden und dann bottom-up die leeren
        Ordner entfernt werden.

        Args:
            directory: Zu bereinigendes Verzeichnis
            delete_root: Auch das Root-Verzeichnis löschen (Default: False)
        """
        if not directory.exists():
            return

        # Phase 1: Alle Dateien löschen (rekursiv)
        for f in directory.rglob("*"):
            if f.is_file():
                try:
                    f.unlink()
                except (PermissionError, OSError):
                    pass  # OneDrive sync kann Dateien sperren

        # Phase 2: Leere Ordner löschen (bottom-up)
        # Sortiere nach Pfadlänge absteigend, damit tiefste Ordner zuerst gelöscht werden
        all_dirs = sorted(
            [d for d in directory.rglob("*") if d.is_dir()],
            key=lambda p: len(p.parts),
            reverse=True
        )
        for d in all_dirs:
            try:
                if d.exists() and not any(d.iterdir()):
                    d.rmdir()
            except (PermissionError, OSError):
                pass

        # Phase 3: Root-Verzeichnis löschen falls gewünscht
        if delete_root:
            try:
                if directory.exists() and not any(directory.iterdir()):
                    directory.rmdir()
            except (PermissionError, OSError):
                pass

    def cleanup(self, session: WorkflowSession, keep_output: bool = True):
        """
        Räumt temporäre Daten auf (NICHT data/ - wird manuell verwaltet).

        Args:
            session: Aktive Session
            keep_output: Output-Ordner behalten (Default: True)
        """
        # HINWEIS: data/ wird NICHT automatisch geleert - Benutzer verwaltet das selbst

        # bundles/ leeren
        for subdir in ["core", "extended"]:
            bundle_dir = self.base_path / "bundles" / subdir
            if bundle_dir.exists():
                for f in bundle_dir.iterdir():
                    try:
                        if f.is_file():
                            f.unlink()
                    except PermissionError:
                        pass

        # Profil-Cleanup
        if session.profile:
            session.profile.cleanup()

        # Session entfernen
        if session.session_id in self._sessions:
            del self._sessions[session.session_id]

    def cleanup_all(self):
        """Räumt alle Sessions und temporären Daten auf."""
        for session_id in list(self._sessions.keys()):
            session = self._sessions[session_id]
            self.cleanup(session)


# ═══════════════════════════════════════════════════════════════
# Convenience-Funktionen
# ═══════════════════════════════════════════════════════════════

def quick_report(
    input_folder: Path,
    client_name: str,
    geburtsdatum: str = "01.01.2010",
    berichtszeitraum: str = "01.01.2025 - 31.12.2025"
) -> Tuple[str, Path]:
    """
    Schneller Workflow: Ordner -> Prompt (ohne LLM-Aufruf).

    Args:
        input_folder: Ordner mit Quelldokumenten
        client_name: Name des Klienten
        geburtsdatum: Geburtsdatum
        berichtszeitraum: Berichtszeitraum

    Returns:
        (prompt, data_path) - Prompt und Pfad zu den Daten
    """
    service = ReportWorkflowService()
    session = service.start_session()

    # Prüfen ob input_folder bereits in data/ liegt
    data_dir = service.base_path / "data"
    is_in_data = False
    try:
        input_folder.relative_to(data_dir)
        is_in_data = True
    except ValueError:
        pass

    if is_in_data:
        # Daten liegen bereits in data/ - nicht kopieren!
        print(f"[INFO] Daten bereits in data/ - ueberspringe Import")
        session.status = "importing"
        session.input_files = list(input_folder.rglob("*"))
    else:
        # Import aus externem Ordner
        service.import_from_folder(session, input_folder)

    # Profil
    service.create_temp_profile(session, client_name, geburtsdatum, scan_folder=input_folder)

    # Anonymisieren - scan direkt den input_folder wenn in data/
    if is_in_data:
        result = service.anonymize_to_bundles_from_folder(session, input_folder)
    else:
        result = service.anonymize_to_bundles(session)

    # Prompt generieren
    prompt = service.generate_prompt(session, berichtszeitraum=berichtszeitraum)

    return prompt, service.base_path


def list_pending_reports() -> List[Dict]:
    """Listet alle Ordner in der Quarantine die verarbeitet werden koennten."""
    quarantine_dir = BACH_BASE / "user" / "foerderplaner" / "_incoming_quarantine"

    if not quarantine_dir.exists():
        return []

    results = []
    for folder in quarantine_dir.iterdir():
        if folder.is_dir() and not folder.name.startswith("."):
            file_count = len(list(folder.rglob("*")))
            results.append({
                "name": folder.name,
                "path": str(folder),
                "files": file_count
            })

    return results


# ═══════════════════════════════════════════════════════════════
# CLI Interface
# ═══════════════════════════════════════════════════════════════

def main():
    """CLI Einstiegspunkt."""
    import sys

    if len(sys.argv) < 2:
        print("Report Workflow Service v1.0.0")
        print()
        print("Usage:")
        print("  python report_workflow_service.py prompt <ordner> <name> [geb]   - Nur Prompt generieren")
        print("  python report_workflow_service.py generate <ordner> <name> [geb] - Kompletter Workflow (LLM + Bericht)")
        print("  python report_workflow_service.py list                           - Wartende Ordner anzeigen")
        print("  python report_workflow_service.py test                           - Schnelltest")
        print()
        print("Beispiel:")
        print("  python report_workflow_service.py generate \"data/Mustermann, Max\" \"Max Mustermann\" \"01.01.2015\"")
        return

    cmd = sys.argv[1]

    if cmd == "prompt":
        if len(sys.argv) < 4:
            print("Fehler: Ordner und Name angeben")
            return

        folder = Path(sys.argv[2])
        name = sys.argv[3]
        geb = sys.argv[4] if len(sys.argv) > 4 else "01.01.2010"

        print(f"Generiere Prompt fuer: {name}")
        prompt, data_path = quick_report(folder, name, geb)

        print(f"\n{'='*60}")
        print(prompt[:2000])
        print(f"{'='*60}")
        print(f"\nVollstaendiger Prompt: {len(prompt)} Zeichen")
        print(f"Daten unter: {data_path}")

    elif cmd == "list":
        pending = list_pending_reports()
        if pending:
            print("Wartende Ordner:")
            for item in pending:
                print(f"  - {item['name']} ({item['files']} Dateien)")
        else:
            print("Keine wartenden Ordner")

    elif cmd == "test":
        print("[TEST] Report Workflow Service")

        service = ReportWorkflowService()
        session = service.start_session()
        print(f"  Session: {session.session_id}")

        # Profil erstellen
        profile = service.create_temp_profile(
            session,
            "Max Mustermann",
            "15.03.2016",
            additional_terms=["Musterstr. 5", "07761/123456"]
        )
        print(f"  Tarnname: {profile.tarnname}")
        print(f"  Fake-Geb: {profile.fake_geburtsdatum}")
        print(f"  Mappings: {len(profile.get_all_mappings())} Eintraege")

        # Cleanup
        service.cleanup(session)
        print("\n[OK] Test abgeschlossen")

    elif cmd == "generate":
        # Kompletter Workflow: Prompt -> LLM -> Bericht
        if len(sys.argv) < 4:
            print("Fehler: Ordner und Name angeben")
            print("Usage: python report_workflow_service.py generate <ordner> <name> [geburtsdatum]")
            return

        folder = Path(sys.argv[2])
        name = sys.argv[3]
        geb = sys.argv[4] if len(sys.argv) > 4 else "01.01.2010"

        print(f"=== BERICHTERSTELLUNG fuer: {name} ===\n")

        # Schritt 1: Prompt generieren
        print("[1/3] Generiere Prompt und Bundles...")
        prompt, data_path = quick_report(folder, name, geb)
        print(f"      Prompt: {len(prompt)} Zeichen")

        # Prompt in Datei speichern
        prompt_file = data_path / "output" / "prompt.txt"
        prompt_file.parent.mkdir(parents=True, exist_ok=True)
        prompt_file.write_text(prompt, encoding="utf-8")
        print(f"      Gespeichert: {prompt_file}")

        # Schritt 2: Auf LLM-Antwort warten
        # HINWEIS: Ollama deaktiviert - verwende stattdessen Claude-Agent
        print("\n[2/3] LLM-Verarbeitung...")
        print("      Ollama deaktiviert - Prompt muss an Claude-Agent uebergeben werden.")
        print(f"      Prompt-Datei: {prompt_file}")

        response_file = data_path / "output" / "llm_response.txt"

        # Pruefen ob bereits eine LLM-Antwort existiert
        if response_file.exists() and response_file.stat().st_size > 100:
            print(f"      Vorhandene LLM-Antwort gefunden: {response_file}")
            llm_response = response_file.read_text(encoding="utf-8")
            print(f"      LLM-Antwort: {len(llm_response)} Zeichen")
        else:
            print("\n      AKTION ERFORDERLICH:")
            print("      1. Oeffne den Prompt in: " + str(prompt_file))
            print("      2. Sende den Inhalt an Claude (z.B. via Claude Code oder Web)")
            print("      3. Speichere die Antwort in: " + str(response_file))
            print("      4. Fuehre diesen Befehl erneut aus")
            print("\n      Alternative: bach bericht finalize <ordner> <llm_response.txt>")
            return

        # Schritt 3: Bericht generieren
        print("\n[3/3] Erstelle Word-Dokument...")
        service = ReportWorkflowService()
        session = service.start_session()
        service.create_temp_profile(session, name, geb, scan_folder=folder)

        output_path = service.generate_report(session, llm_response)

        if output_path and output_path.exists():
            print(f"\n=== FERTIG ===")
            print(f"Bericht erstellt: {output_path}")
        else:
            print(f"\n=== FEHLER ===")
            print(f"Fehler: {session.error_message}")

    else:
        print(f"Unbekannter Befehl: {cmd}")


if __name__ == "__main__":
    main()
