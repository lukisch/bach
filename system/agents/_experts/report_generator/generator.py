#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Copyright (c) 2026 Lukas Geiger

Permission is hereby granted, free of charge, to any person obtaining a copy
of this software and associated documentation files (the "Software"), to deal
in the Software without restriction, including without limitation the rights
to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
copies of the Software, and to permit persons to whom the Software is
furnished to do so, subject to the following conditions:

The above copyright notice and this permission notice shall be included in all
copies or substantial portions of the Software.

THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
SOFTWARE.
"""

"""
Report Generator - Kernlogik
=============================

Verbindet Dokumenten-Extraktion, LLM-Generierung und Word-Template-Befuellung.

Ablauf:
  1. Quelldokumente lesen (Word, PDF, Text)
  2. LLM-Prompt zusammenbauen mit JSON-Schema
  3. JSON-Output in Word-Vorlage einfuegen

Version: 1.0.0
Erstellt: 2026-01-27
"""

import json
import sys
import io
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple

# UTF-8 Output erzwingen
if sys.stdout.encoding != 'utf-8':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

# ═══════════════════════════════════════════════════════════════
# Universal Import: bach_paths.py (Single Source of Truth)
# ═══════════════════════════════════════════════════════════════
# Dieses Pattern funktioniert von ueberall im System

_current = Path(__file__).resolve()
for _parent in [_current] + list(_current.parents):
    _hub = _parent / "system" / "hub"
    if _hub.exists():
        if str(_hub) not in sys.path:
            sys.path.insert(0, str(_hub))
        break

# Jetzt koennen wir bach_paths importieren
try:
    from bach_paths import get_path, BACH_ROOT, SYSTEM_ROOT
    _USE_BACH_PATHS = True
except ImportError:
    _USE_BACH_PATHS = False

# Pfade - aus bach_paths.py oder Fallback
EXPERT_DIR = Path(__file__).parent
SCHEMAS_DIR = EXPERT_DIR / "schemas"

if _USE_BACH_PATHS:
    BACH_DIR = SYSTEM_ROOT
    TEMPLATES_DIR = get_path("templates")
else:
    # Fallback
    BACH_DIR = EXPERT_DIR.parent.parent.parent  # agents/_experts/report_generator -> skills -> system
    TEMPLATES_DIR = BACH_DIR / "skills" / "_templates"

# Services importieren - mehrere Pfade für verschiedene Aufruf-Kontexte
sys.path.insert(0, str(BACH_DIR))
sys.path.insert(0, str(BACH_DIR / "hub" / "_services" / "document"))
try:
    from hub._services.document.word_template_service import WordTemplateService
    WORD_SERVICE_AVAILABLE = True
except ImportError:
    try:
        from word_template_service import WordTemplateService
        WORD_SERVICE_AVAILABLE = True
    except ImportError:
        WORD_SERVICE_AVAILABLE = False

try:
    from docx import Document
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Textbausteine für Bereiche und Umweltfaktoren
try:
    from .textbausteine import get_bereich_text, get_umwelt_text
    TEXTBAUSTEINE_AVAILABLE = True
except ImportError:
    try:
        from textbausteine import get_bereich_text, get_umwelt_text
        TEXTBAUSTEINE_AVAILABLE = True
    except ImportError:
        TEXTBAUSTEINE_AVAILABLE = False

# Bereich-Codes für strukturierte LLM-Generierung
try:
    from .bereich_codes import (
        BEREICH_STRUKTUR, UMWELT_STRUKTUR,
        build_prompt_code_reference, get_bereich_name, get_umwelt_name
    )
    BEREICH_CODES_AVAILABLE = True
except ImportError:
    try:
        from bereich_codes import (
            BEREICH_STRUKTUR, UMWELT_STRUKTUR,
            build_prompt_code_reference, get_bereich_name, get_umwelt_name
        )
        BEREICH_CODES_AVAILABLE = True
    except ImportError:
        BEREICH_CODES_AVAILABLE = False

# Anonymized Workflow für LLM-Schutz
try:
    from .anonymized_workflow import AnonymizedWorkflow, BundleInfo
    ANONYMIZED_WORKFLOW_AVAILABLE = True
except ImportError:
    try:
        from anonymized_workflow import AnonymizedWorkflow, BundleInfo
        ANONYMIZED_WORKFLOW_AVAILABLE = True
    except ImportError:
        ANONYMIZED_WORKFLOW_AVAILABLE = False


# ═══════════════════════════════════════════════════════════════
# Datenklassen
# ═══════════════════════════════════════════════════════════════

@dataclass
class GeneratorResult:
    """Ergebnis der Berichtsgenerierung."""
    success: bool = False
    output_path: str = ""
    source_text_length: int = 0
    json_generated: bool = False
    template_filled: bool = False
    errors: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)


# ═══════════════════════════════════════════════════════════════
# Phase 1: Dokument-Extraktion
# ═══════════════════════════════════════════════════════════════

def extract_text_from_docx(filepath: str) -> str:
    """Extrahiert Text aus einem Word-Dokument."""
    if not DOCX_AVAILABLE:
        return f"[FEHLER: python-docx nicht installiert, kann {filepath} nicht lesen]"

    doc = Document(filepath)
    parts = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_text_from_txt(filepath: str) -> str:
    """Liest eine Textdatei."""
    path = Path(filepath)
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="latin-1")


def extract_text_from_pdf(filepath: str) -> str:
    """Extrahiert Text aus einem PDF. Nutzt pypdf (MIT) primaer, PyMuPDF optional."""
    # Primaer: pypdf (MIT-Lizenz)
    try:
        from pypdf import PdfReader
        reader = PdfReader(filepath)
        text = ""
        for page in reader.pages:
            text += (page.extract_text() or "") + "\n"
        if text.strip():
            return text
    except ImportError:
        pass
    except Exception:
        pass

    # Fallback: pdfplumber (MIT-Lizenz)
    try:
        import pdfplumber
        text = ""
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                text += (page.extract_text() or "") + "\n"
        if text.strip():
            return text
    except ImportError:
        pass
    except Exception:
        pass

    # Optional: PyMuPDF (AGPL) -- nur wenn installiert
    try:
        import fitz  # PyMuPDF optional
        doc = fitz.open(filepath)
        text = ""
        for page in doc:
            text += page.get_text() + "\n"
        doc.close()
        return text
    except ImportError:
        return f"[PDF-Extraktion nicht moeglich: pypdf, pdfplumber und PyMuPDF nicht installiert. Datei: {filepath}]"


def extract_text_from_excel(filepath: str) -> str:
    """Extrahiert Text aus einer Excel-Datei (.xlsx, .xls)."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(filepath, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            parts.append(f"[Tabelle: {sheet_name}]")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(cells):
                    parts.append(" | ".join(cells))
        wb.close()
        return "\n".join(parts)
    except ImportError:
        return f"[Excel-Extraktion nicht moeglich: openpyxl nicht installiert. Datei: {filepath}]"
    except Exception as e:
        return f"[Excel-Fehler: {e}. Datei: {filepath}]"


def extract_text_from_msg(filepath: str) -> str:
    """Extrahiert Text aus einer Outlook .msg Datei."""
    try:
        import extract_msg
        msg = extract_msg.Message(filepath)
        parts = []

        # Header-Infos
        if msg.date:
            parts.append(f"Datum: {msg.date}")
        if msg.sender:
            parts.append(f"Von: {msg.sender}")
        if msg.to:
            parts.append(f"An: {msg.to}")
        if msg.subject:
            parts.append(f"Betreff: {msg.subject}")

        parts.append("")  # Leerzeile

        # Body
        if msg.body:
            parts.append(msg.body)

        msg.close()
        return "\n".join(parts)
    except ImportError:
        return f"[MSG-Extraktion nicht moeglich: extract-msg nicht installiert. Datei: {filepath}]"
    except Exception as e:
        return f"[MSG-Fehler: {e}. Datei: {filepath}]"


def extract_text_from_eml(filepath: str) -> str:
    """Extrahiert Text aus einer .eml Datei."""
    try:
        import email
        from email import policy
        from email.parser import BytesParser

        with open(filepath, 'rb') as f:
            msg = BytesParser(policy=policy.default).parse(f)

        parts = []

        # Header-Infos
        if msg['Date']:
            parts.append(f"Datum: {msg['Date']}")
        if msg['From']:
            parts.append(f"Von: {msg['From']}")
        if msg['To']:
            parts.append(f"An: {msg['To']}")
        if msg['Subject']:
            parts.append(f"Betreff: {msg['Subject']}")

        parts.append("")  # Leerzeile

        # Body extrahieren
        if msg.is_multipart():
            for part in msg.walk():
                content_type = part.get_content_type()
                if content_type == 'text/plain':
                    charset = part.get_content_charset() or 'utf-8'
                    try:
                        body = part.get_payload(decode=True).decode(charset)
                        parts.append(body)
                        break
                    except:
                        pass
        else:
            charset = msg.get_content_charset() or 'utf-8'
            try:
                body = msg.get_payload(decode=True).decode(charset)
                parts.append(body)
            except:
                parts.append(str(msg.get_payload()))

        return "\n".join(parts)
    except Exception as e:
        return f"[EML-Fehler: {e}. Datei: {filepath}]"


def extract_all_sources(source_folder: str, recursive: bool = True) -> str:
    """
    Liest alle Quelldokumente aus einem Ordner.
    Unterstuetzt: .docx, .txt, .md, .pdf, .xlsx, .xls, .msg, .eml

    Args:
        source_folder: Pfad zum Klienten-Ordner
        recursive: Auch Unterordner durchsuchen (default: True)
    """
    folder = Path(source_folder)
    if not folder.exists():
        return f"[FEHLER: Ordner nicht gefunden: {source_folder}]"

    all_text = []
    supported = {".docx", ".txt", ".md", ".pdf", ".xlsx", ".xls", ".msg", ".eml"}

    # Rekursiv oder nur Hauptordner
    if recursive:
        files = sorted(folder.rglob("*"))
    else:
        files = sorted(folder.iterdir())

    for filepath in files:
        # Ueberspringe versteckte Dateien, output-Ordner und Profil-Dateien
        if filepath.name.startswith(".") or filepath.name.startswith("_"):
            continue
        if "output" in filepath.parts:
            continue
        if not filepath.is_file():
            continue
        if filepath.suffix.lower() not in supported:
            continue

        suffix = filepath.suffix.lower()
        rel_path = filepath.relative_to(folder)

        if suffix == ".docx":
            text = extract_text_from_docx(str(filepath))
        elif suffix in (".txt", ".md"):
            text = extract_text_from_txt(str(filepath))
        elif suffix == ".pdf":
            text = extract_text_from_pdf(str(filepath))
        elif suffix in (".xlsx", ".xls"):
            text = extract_text_from_excel(str(filepath))
        elif suffix == ".msg":
            text = extract_text_from_msg(str(filepath))
        elif suffix == ".eml":
            text = extract_text_from_eml(str(filepath))
        else:
            continue

        all_text.append(f"--- Quelle: {rel_path} ---\n{text}")

    if not all_text:
        return "[FEHLER: Keine Quelldokumente gefunden]"

    return "\n\n".join(all_text)


# ═══════════════════════════════════════════════════════════════
# Phase 2: LLM-Prompt Erstellung
# ═══════════════════════════════════════════════════════════════

# ICF-KAPITEL - Automatische Zuordnung basierend auf erstem Digit
# Alle D1xxx Codes gehören zu Kapitel 1, alle D7xxx zu Kapitel 7, etc.
ICF_KAPITEL = {
    "1": "Lernen und Wissensanwendung",
    "2": "Allgemeine Aufgaben und Anforderungen",
    "3": "Kommunikation",
    "4": "Mobilität",
    "5": "Selbstversorgung",
    "6": "Häusliches Leben",
    "7": "Interpersonelle Interaktionen und Beziehungen",
    "8": "Bedeutende Lebensbereiche",
    "9": "Gemeinschafts-, soziales und staatsbürgerliches Leben"
}


def get_icf_kapitel(icf_code: str) -> str:
    """
    Ermittelt das ICF-Kapitel automatisch aus dem Code.
    D1330 -> "Lernen und Wissensanwendung"
    D7503 -> "Interpersonelle Interaktionen und Beziehungen"
    """
    code = icf_code.upper().strip()
    if code.startswith("D") and len(code) > 1:
        digit = code[1]  # Erste Ziffer nach D
        return ICF_KAPITEL.get(digit, "")
    return ""


# GÜLTIGE ICF-CODES - Diese existieren im Word-Template als Zeilen
# Jeder Code hat Platzhalter wie {D1330-Ziel}, {D1330-Ist}, etc.
# ACHTUNG: Diese Liste muss EXAKT mit dem Template übereinstimmen!
# Codes wie D1, D2, D3, D7, D240, D7103, D880 existieren NICHT im Template!
VALID_ICF_CODES = {
    # Lernen und Wissensanwendung
    "D1310", "D1314", "D1630", "D1370", "D1371", "D1631", "D1632",
    "D160", "D161", "D177", "D135", "D1330", "D1331", "D1332",
    "D1402", "D1400", "D1401", "D1450", "D1451", "D1452", "D1660", "D1661",
    # Allgemeine Aufgaben (KEIN D240!)
    "D210", "D220", "D230", "D250",
    # Kommunikation
    "D331", "D310", "D330", "D315", "D335", "D3350", "D3351", "D3352",
    "D350", "D3500", "D3501", "D3502", "D3503", "D3504", "D3600",
    # Interpersonelle Interaktionen (KEIN D7103!)
    "D7100", "D7101", "D7102", "D7104", "D7105", "D7106",
    "D7200", "D7201", "D7202", "D7203", "D7204",
    "D730", "D740", "D7500", "D7501", "D7502", "D7503", "D760", "D770"
}


def validate_and_fix_icf_codes(report_data: dict) -> Tuple[dict, List[str]]:
    """
    Validiert ICF-Codes im JSON und entfernt ungültige Einträge.

    Returns:
        (fixed_data, warnings) - Bereinigte Daten und Liste von Warnungen
    """
    warnings = []

    # Förderziele prüfen
    foerderziele = report_data.get("foerderziele", []) or report_data.get("förderziele", [])
    valid_ziele = []

    for ziel in foerderziele:
        code = ziel.get("icf_code", "").upper()
        if code in VALID_ICF_CODES:
            ziel["icf_code"] = code  # Normalisieren auf Großbuchstaben
            valid_ziele.append(ziel)
        else:
            warnings.append(f"WARNUNG: ICF-Code '{code}' existiert nicht im Template - Ziel entfernt")

    # Neue Ziele prüfen
    neue_ziele = report_data.get("neue_ziele", [])
    valid_neue = []

    for ziel in neue_ziele:
        code = ziel.get("icf_code", "").upper()
        if code in VALID_ICF_CODES:
            ziel["icf_code"] = code
            valid_neue.append(ziel)
        else:
            warnings.append(f"WARNUNG: ICF-Code '{code}' in neue_ziele existiert nicht - entfernt")

    # Bereinigte Daten zurückgeben
    fixed_data = report_data.copy()
    fixed_data["foerderziele"] = valid_ziele
    if "förderziele" in fixed_data:
        fixed_data["förderziele"] = valid_ziele
    fixed_data["neue_ziele"] = valid_neue

    if not valid_ziele:
        warnings.append("FEHLER: Keine gültigen Förderziele nach Validierung!")

    return fixed_data, warnings


def validate_stammdaten(report_data: dict) -> List[str]:
    """
    Validiert Stammdaten auf Vollstaendigkeit und Format.
    Returns: Liste von Warnungen/Fehlern.
    """
    warnings = []
    stamm = report_data.get("stammdaten", {})

    if not stamm:
        warnings.append("FEHLER: Keine Stammdaten vorhanden!")
        return warnings

    # Pflichtfelder pruefen
    required = {"name": "Name", "geburtsdatum": "Geburtsdatum",
                "landkreis": "Landkreis", "kostenzusage_ende": "Kostenzusage-Ende"}
    for key, label in required.items():
        val = stamm.get(key, "")
        if not val or not str(val).strip():
            warnings.append(f"WARNUNG: Stammdaten.{label} fehlt oder ist leer")

    # Landkreis-Enum pruefen
    landkreis = stamm.get("landkreis", "")
    if landkreis and landkreis.lower() not in ("loerrach", "lörrach", "waldshut"):
        warnings.append(f"WARNUNG: Landkreis '{landkreis}' unbekannt - nur Loerrach/Waldshut unterstuetzt")

    # Datumsformat pruefen (dd.mm.yyyy)
    import re
    date_pattern = re.compile(r'^\d{2}\.\d{2}\.\d{4}$')
    date_fields = {"geburtsdatum": "Geburtsdatum", "kostenzusage_ende": "Kostenzusage-Ende",
                   "weiterbewilligung_ab": "Weiterbewilligung ab", "foerderungsbeginn": "Foerderungsbeginn"}
    for key, label in date_fields.items():
        val = stamm.get(key, "")
        if val and not date_pattern.match(str(val)):
            warnings.append(f"WARNUNG: {label} hat ungueltiges Format '{val}' (erwartet: dd.mm.yyyy)")

    # Diagnosen pruefen
    diagnosen = stamm.get("diagnosen", [])
    if not diagnosen:
        warnings.append("WARNUNG: Keine Diagnosen angegeben")
    else:
        icd_pattern = re.compile(r'^[A-Z]\d{2}(\.\d{1,2})?$')
        for i, diag in enumerate(diagnosen):
            if isinstance(diag, dict):
                icd = diag.get("icd_code", "")
                if icd and not icd_pattern.match(icd):
                    warnings.append(f"WARNUNG: Diagnose {i+1}: ICD-Code '{icd}' hat ungueltiges Format")

    return warnings


def validate_text_quality(report_data: dict) -> List[str]:
    """
    Prueft Textfelder auf Mindestlaenge und Qualitaet.
    Returns: Liste von Warnungen.
    """
    warnings = []
    MIN_SENTENCES = 2

    def count_sentences(text: str) -> int:
        if not text:
            return 0
        return len([s for s in text.replace("...", ".").split('.') if s.strip()])

    # Foerderziele: ist_stand pruefen
    ziele = report_data.get("foerderziele", []) or report_data.get("förderziele", [])
    for i, ziel in enumerate(ziele):
        ist = ziel.get("ist_stand", "")
        if ist and count_sentences(ist) < MIN_SENTENCES:
            code = ziel.get("icf_code", f"#{i+1}")
            warnings.append(f"WARNUNG: ist_stand fuer {code} zu kurz ({count_sentences(ist)} Satz, min. {MIN_SENTENCES})")

    # Abschluss-Texte
    abschluss = report_data.get("abschluss", {})
    for key, label, min_s in [
        ("aktuelle_entwicklungen", "Aktuelle Entwicklungen", 2),
        ("bedingungsmodell", "Bedingungsmodell", 2),
    ]:
        text = abschluss.get(key, "")
        if text and count_sentences(text) < min_s:
            warnings.append(f"WARNUNG: {label} zu kurz ({count_sentences(text)} Satz, min. {min_s})")

    # Besondere Faehigkeiten
    bf = report_data.get("besondere_faehigkeiten", "")
    if not bf or not bf.strip():
        warnings.append("WARNUNG: Besondere Faehigkeiten fehlt")

    return warnings


def load_schema(schema_name: str = "foerderbericht") -> dict:
    """Laedt ein JSON-Schema."""
    schema_path = SCHEMAS_DIR / f"{schema_name}.json"
    if not schema_path.exists():
        raise FileNotFoundError(f"Schema nicht gefunden: {schema_path}")
    return json.loads(schema_path.read_text(encoding="utf-8"))


def build_icf_reference() -> str:
    """
    Baut eine ICF-Referenz aus der Word-Vorlage fuer den LLM-Prompt.
    Listet NUR die ICF-Codes auf, die im Template als Zeilen existieren.
    """
    # ICF-Codes die EXAKT im Word-Template existieren (als Tabellenzeilen)
    # WICHTIG: Nur diese Codes verwenden! Andere werden vom Generator abgelehnt.
    icf_codes = {
        # Lernen und Wissensanwendung
        "D1310": "Foerderung des Symbolspiels",
        "D1314": "Foerderung des So-tun-als-ob-Spiels",
        "D1630": "Rollenspiele, Theory of Mind",
        "D1370": "Verstaendnis Groesse/Form/Menge (TEACCH)",
        "D1371": "Klassifizieren, Gruppieren, Reihen",
        "D1631": "Logische Zusammenhaenge erkennen",
        "D1632": "Ursache-Wirkung vorhersagen",
        "D160": "Aufmerksamkeit umlenken",
        "D161": "Aufmerksamkeit aufrechterhalten",
        "D177": "Auswahl treffen",
        "D135": "Wiederholen als Lerntechnik",
        "D1330": "Woerter sprechen / Symbole nutzen",
        "D1331": "Einfache Saetze produzieren",
        "D1332": "Komplexe Saetze sprechen",
        "D1402": "Woerter und Saetze verstehen",
        "D1400": "Buchstaben/Symbole erkennen (TEACCH)",
        "D1401": "Buchstaben aussprechen",
        "D1450": "Handmotorik foerdern",
        "D1451": "Morphem-Graphem Zuordnung",
        "D1452": "Gesprochene Woerter schreiben",
        "D1660": "Lesestrategien entwickeln",
        "D1661": "Lautes und stilles Lesen",
        # Allgemeine Aufgaben und Anforderungen
        "D210": "Einzelaufgabe uebernehmen",
        "D220": "Mehrfachaufgaben uebernehmen",
        "D230": "Tagesroutine, Plaene, Flexibilitaet",
        "D250": "Verhaltensregulation, Emotionsregulation",
        # Kommunikation
        "D331": "Praeverbale Aeusserungen (Lautieren)",
        "D310": "Verbalsprache verstehen",
        "D330": "Verbales Sprechen",
        "D315": "Nonverbale Botschaften verstehen",
        "D335": "Nonverbale Botschaften senden",
        "D3350": "Koerpersprache, Gesten, Mimik",
        "D3351": "Zeichen und Symbole verwenden",
        "D3352": "Piktogramme zur Kommunikation",
        "D350": "Konversation fuehren",
        "D3500": "Unterhaltung beginnen (Small Talk, Blickkontakt)",
        "D3501": "Unterhaltung aufrechterhalten",
        "D3502": "Unterhaltung beenden",
        "D3503": "Dialog fuehren",
        "D3504": "Gruppengespraech fuehren",
        "D3600": "Kommunikationsgeraete/-techniken",
        # Interpersonelle Interaktionen
        "D7100": "Respekt und Waerme in Beziehungen",
        "D7101": "Anerkennung geben/erhalten",
        "D7102": "Toleranz in Beziehungen",
        "D7104": "Soziale Zeichen verstehen/senden",
        "D7105": "Koerperliche Naehe/Distanz regulieren",
        "D7106": "Verhalten an Vertrautheit anpassen",
        "D7200": "Beziehungen eingehen und fuehren",
        "D7201": "Beziehungen beenden (Trennung, Verlust)",
        "D7202": "Emotionsregulation in Beziehungen",
        "D7203": "Sozialen Regeln gemaess interagieren",
        "D7204": "Sozialen Abstand wahren",
        "D730": "Mit Fremden umgehen",
        "D740": "Formelle Beziehungen / Hierarchien",
        "D7500": "Freundschaften pflegen",
        "D7501": "Mit Nachbarn zurechtkommen",
        "D7502": "Beziehungen zu Bekannten",
        "D7503": "Mit Gleichaltrigen umgehen",
        "D760": "Familienbeziehungen",
        "D770": "Liebe und Sexualitaet"
    }

    lines = [
        "=== GUELTIGE ICF-CODES (NUR DIESE VERWENDEN!) ===",
        "ACHTUNG: Verwende AUSSCHLIESSLICH die folgenden Codes.",
        "Andere Codes wie D1, D2, D3, D7, D240, D7103, D880 etc. existieren NICHT im Template!",
        ""
    ]
    for code, desc in icf_codes.items():
        lines.append(f"  {code}: {desc}")
    return "\n".join(lines)


def build_llm_prompt(source_text: str, schema: dict) -> str:
    """
    Baut den vollstaendigen LLM-Prompt zusammen.

    Der Prompt instruiert das LLM:
    - Beobachtungen zu analysieren
    - Relevante ICF-Bereiche zu identifizieren
    - Strukturierten JSON-Output zu generieren
    """
    icf_ref = build_icf_reference()

    # Bereich-Codes Referenz für strukturierte Textgenerierung
    if BEREICH_CODES_AVAILABLE:
        bereich_ref = build_prompt_code_reference()
    else:
        bereich_ref = ""

    prompt = f"""Du bist ein erfahrener Foerderplaner fuer Autismus-Foerderung.
Basierend auf den folgenden Beobachtungsnotizen erstellst du einen strukturierten
Entwicklungsbericht im JSON-Format.

AUFGABE:
1. Lies die Beobachtungen sorgfaeltig
2. Identifiziere die relevanten ICF-Bereiche
3. Formuliere fuer jedes aktive Foerderziel einen Ist-Stand (2-3 Saetze)
4. Bewerte ob die Ziele erreicht wurden (1=nicht, 2=teilweise, 3=erreicht)
5. Identifiziere welche Satzbausteine-Bereiche aktiv sein sollten
6. Formuliere individualisierte Texte fuer die aktiven Bereiche
7. Schlage neue Ziele vor falls angebracht
8. Schreibe die Abschluss-Bemerkung mit Bedingungsmodell

WICHTIG (ICF-CODES):
- Verwende AUSSCHLIESSLICH die unten gelisteten ICF-Codes (d1-d9).
- D880 existiert NICHT. Verwende D1310 oder D1314 für Spielverhalten/So-tun-als-ob.
- Erfinde KEINE eigenen Codes.

WICHTIG (FORMAT & STIL):
- Verwende den Namen des Klienten wie in den Beobachtungen angegeben
- Formuliere im fachlichen Berichtsstil (3. Person, sachlich)
- Ist-Stand-Beschreibungen sollen konkret sein, nicht allgemein
- Beziehe dich auf konkrete Beobachtungen und Fortschritte
- Bei Satzbausteine-Bereichen: Passe den Vorlagentext an den Klienten an
  (ersetze [Name] durch den echten Namen, streiche nicht zutreffende Teile)
- Gib NUR valides JSON zurueck, keine Erklaerungen davor oder danach

STAMMDATEN-EXTRAKTION (WICHTIG - Suche aktiv in den Dokumenten!):
- landkreis: Aus Kostenzusage/Hilfeplan ablesen - "Loerrach" oder "Waldshut"
  * Suche nach "Landratsamt Loerrach" oder "Landratsamt Waldshut"
- zeichen_amt: Aktenzeichen des Amts (z.B. "51.2.1-XXX" oder "Z1-512.41-...")
  * Suche nach "Ihr Zeichen", "Az.", "Aktenzeichen", "Geschaeftszeichen"
- foerderungsbeginn: Wann begann die Foerderung bei proAutismus? (aus Aktendeckblatt)
- kostenzusage_ende: Datum Bewilligungsende (dd.mm.yyyy) - DIESES FELD IST KRITISCH!
  * Suche nach "Bewilligungszeitraum", "bewilligt bis", "Kostenzusage bis", "Massnahme endet"
  * Wenn Text wie "bis Ende Februar 2026", konvertiere zu "28.02.2026"
  * Wenn Text wie "bis Ende September 2025", konvertiere zu "30.09.2025"
  * Wenn Text wie "31.12.2025" oder "30.06.2026", uebernehme direkt
  * NIEMALS leer lassen! Suche in allen Dokumenten (auch E-Mail-Anhaenge)
- sachbearbeiter: Name der zustaendigen Person beim Jugendamt
  * Suche nach "Sachbearbeiter", "Ansprechpartner", Unterschrift in Amtsschreiben
- diagnosen: Fuer JEDE Diagnose erfassen:
  * icd_code: z.B. "F84.5", "F84.0", "F84.1"
  * bezeichnung: z.B. "Asperger-Syndrom", "Fruehkindlicher Autismus"
  * diagnostiker: Wer hat diagnostiziert? (z.B. "Dr. Ritter-Gekeler", "Kreiskrankenhaus Loerrach")
  * datum: Wann? (Jahr oder volles Datum)
  * quelle: Aus welchem Dokument stammt die Info? (z.B. "Bericht vom 12.03.2019")

{icf_ref}

{bereich_ref}

=== BEREICHE UND UMWELTFAKTOREN (Strukturierte Textgenerierung) ===
Für jeden Bereich und Umweltfaktor gibt es Sub-Codes.
Schreibe für jeden RELEVANTEN Sub-Code einen kurzen Text (1-2 Sätze).
Wenn mindestens ein Sub-Code Text hat, wird der Bereich automatisch aufgenommen.
Codes die nicht relevant sind: null oder weglassen.

Beispiel für "bereiche":
{{
  "gesellschaft_freizeit": {{
    "D910": "Jaden zeigt Interesse an anderen Kindern und sucht Blickkontakt.",
    "D920": "In der Freizeitgestaltung bevorzugt Jaden strukturierte Aktivitäten.",
    "D9200": null
  }},
  "selbstversorgung": {{
    "D510": "Jaden zeigt Fortschritte bei der Körperpflege.",
    "D540": "Das An- und Ausziehen gelingt zunehmend selbstständig."
  }}
}}

Beispiel für "umweltfaktoren":
{{
  "wahrnehmung": {{
    "E240": "Jaden bevorzugt gedimmtes Licht und reagiert empfindlich auf Helligkeit.",
    "E250": null
  }}
}}

BEOBACHTUNGEN:
==============
{source_text}
==============

=== EMPFEHLUNG (Defaults und Nachfragen) ===

(1) VERLÄNGERUNG ODER BEENDIGUNG:
- DEFAULT: Verlängerung (verlaengerung: true)
- Falls aus den Dokumenten hervorgeht, dass Beendigung angedacht ist:
  * verlaengerung: false
  * beendigung_bedingungen: z.B. "Spätere Wiederaufnahme vereinbart" oder andere Bedingungen
- Falls unklar: Im Feld "rueckfragen" notieren: "Verlängerung oder Beendigung?"

(2) SONSTIGE EMPFEHLUNGEN:
- Feld: sonstige_empfehlung
- Aus Dokumenten ableiten oder leer lassen
- Falls unklar aber relevant: Im Feld "rueckfragen" notieren

(3) UMFANG DER AUTISMUSTHERAPIE:
- Feld: umfang (String mit vollständiger Beschreibung)

DEFAULT (wenn nichts anderes angegeben):
  "Autismusspezifische Einzelförderung: 75 Minuten wöchentlich
   Eltern-/Umfeldarbeit: bis zu 1 Stunde monatlich"

SONDERREGELUNGEN (nur wenn aus Dokumenten ersichtlich):

a) Mit Gruppenförderung:
  "Autismusspezifische Einzelförderung: 3 x 75 Minuten monatlich
   Gruppenförderung: 1 x 75 Minuten monatlich
   Eltern-/Umfeldarbeit: bis zu 1 Stunde monatlich"

b) Nur Gruppe:
  "Gruppenförderung: 1 x 75 Minuten monatlich"

c) Abrufbudget (bei Beendigung):
  "Abrufbudget: [X] Stunden für Übergangszeit bis [DATUM]"
  (0-10 Stunden für definierte Übergangszeit)

- Falls Umfang unklar: Im Feld "rueckfragen" notieren: "Welcher Förderumfang?"

=== RÜCKFRAGEN ===
Falls wichtige Informationen fehlen, erstelle ein Array "rueckfragen" mit offenen Fragen:
Beispiel: ["Verlängerung oder Beendigung?", "Welcher Förderumfang?"]

AUSGABE-FORMAT (JSON):
Erzeuge ein JSON-Objekt gemaess folgendem Schema.
Nur die relevanten ICF-Codes in foerderziele aufnehmen.
Bereiche und Umweltfaktoren nur aktivieren wenn in den Beobachtungen erwaehnt.

Schema-Zusammenfassung:
- stammdaten:
  * name, geburtsdatum, sachbearbeiter
  * landkreis ("Loerrach" oder "Waldshut")
  * zeichen_amt (Aktenzeichen)
  * foerderungsbeginn (bei proAutismus)
  * kostenzusage_ende (Datum, auch aus Text konvertiert)
  * diagnosen: Array von {{icd_code, bezeichnung, diagnostiker, datum, quelle}}
  * therapieangebot, netzwerkarbeit
- foerderziele: Array mit icf_code, zielformulierung, ist_stand, erreicht (1-3), grund_nichterreichung
- bereiche: mobilitaet, selbstversorgung, haushalt, lebensbereiche, gesellschaft_freizeit (je aktiv + text)
- umweltfaktoren: wahrnehmung, kommunikation, kognition_motorik, medizinisches (je aktiv + text)
- besondere_faehigkeiten: Staerken des Klienten
- neue_ziele: Array mit icf_code, zielformulierung, ist_stand
- empfehlung:
  * verlaengerung (bool, default: true)
  * begruendung (Array): ["therapieziele_nicht_erreicht", "uebergang_lebensbereich", "krise_erhoehter_bedarf"]
  * beendigung_bedingungen (String, nur bei verlaengerung=false): z.B. "Spätere Wiederaufnahme vereinbart"
  * umfang (String): Vollständige Beschreibung des Förderumfangs (siehe Defaults oben)
  * sonstige_empfehlung (String): Weitere Empfehlungen
- abschluss: aktuelle_entwicklungen, bedingungsmodell
- rueckfragen (Array, optional): Offene Fragen falls wichtige Infos fehlen

JSON:
"""
    return prompt


# ═══════════════════════════════════════════════════════════════
# Phase 3: Template-Befuellung
# ═══════════════════════════════════════════════════════════════

def _replace_complex_placeholder_in_xml(doc: Document, search_text: str, replacement: str) -> bool:
    """
    Ersetzt komplexe Platzhalter die ueber mehrere XML-Elemente verteilt sind.
    Sucht nach dem Text im XML und ersetzt ihn direkt.

    Args:
        doc: Word-Dokument
        search_text: Text der gesucht wird (z.B. "Landratsamt")
        replacement: Ersetzungstext

    Returns:
        True wenn ersetzt, False sonst
    """
    from docx.oxml.ns import qn

    replaced = False

    # Suche in allen Paragraphen (auch in Tabellenzellen)
    for para in doc.paragraphs:
        # Sammle alle Text-Elemente
        t_elements = list(para._element.iter(qn("w:t")))
        full_text = "".join(t.text or "" for t in t_elements)

        if search_text in full_text:
            # Finde Start und Ende
            start_idx = full_text.find(search_text)
            end_idx = start_idx + len(search_text)

            # Finde die zugehoerigen t-Elemente und ersetze
            char_idx = 0
            first_t = None
            for t in t_elements:
                t_text = t.text or ""
                t_start = char_idx
                t_end = char_idx + len(t_text)

                if t_start <= start_idx < t_end and first_t is None:
                    first_t = t
                    # Ersetze in diesem Element
                    prefix = t_text[:start_idx - t_start]
                    suffix_start = min(end_idx - t_start, len(t_text))
                    suffix = t_text[suffix_start:] if suffix_start < len(t_text) else ""
                    t.text = prefix + replacement + suffix
                    replaced = True
                elif first_t is not None and t_start < end_idx:
                    # Nachfolgende Elemente die Teil des Platzhalters sind
                    if t_end <= end_idx:
                        t.text = ""  # Gesamtes Element war Teil
                    else:
                        # Teilweise - behalte Rest
                        t.text = t_text[end_idx - t_start:]

                char_idx = t_end

    # Auch in Tabellen suchen
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    t_elements = list(para._element.iter(qn("w:t")))
                    full_text = "".join(t.text or "" for t in t_elements)

                    if search_text in full_text:
                        start_idx = full_text.find(search_text)
                        end_idx = start_idx + len(search_text)

                        char_idx = 0
                        first_t = None
                        for t in t_elements:
                            t_text = t.text or ""
                            t_start = char_idx
                            t_end = char_idx + len(t_text)

                            if t_start <= start_idx < t_end and first_t is None:
                                first_t = t
                                prefix = t_text[:start_idx - t_start]
                                suffix_start = min(end_idx - t_start, len(t_text))
                                suffix = t_text[suffix_start:] if suffix_start < len(t_text) else ""
                                t.text = prefix + replacement + suffix
                                replaced = True
                            elif first_t is not None and t_start < end_idx:
                                if t_end <= end_idx:
                                    t.text = ""
                                else:
                                    t.text = t_text[end_idx - t_start:]

                            char_idx = t_end

    return replaced


def fill_template(
    template_path: str,
    report_data: dict,
    output_path: str
) -> GeneratorResult:
    """
    Fuellt die Word-Vorlage mit den generierten Daten.

    Args:
        template_path: Pfad zur Word-Vorlage
        report_data: Generierter JSON-Output vom LLM
        output_path: Pfad fuer das Ergebnis-Dokument
    """
    result = GeneratorResult()

    if not WORD_SERVICE_AVAILABLE:
        result.errors.append("WordTemplateService nicht verfuegbar")
        return result

    svc = WordTemplateService()

    try:
        doc = svc.load_template(template_path)
    except FileNotFoundError as e:
        result.errors.append(str(e))
        return result

    # === STAMMDATEN VALIDIERUNG ===
    stamm_warnings = validate_stammdaten(report_data)
    for warning in stamm_warnings:
        result.warnings.append(warning)
        print(f"[VALIDIERUNG] {warning}")

    # === TEXT-QUALITAET VALIDIERUNG ===
    text_warnings = validate_text_quality(report_data)
    for warning in text_warnings:
        result.warnings.append(warning)
        print(f"[VALIDIERUNG] {warning}")

    # === ICF-CODE VALIDIERUNG ===
    # Prüfe und bereinige ungültige ICF-Codes BEVOR sie verarbeitet werden
    report_data, validation_warnings = validate_and_fix_icf_codes(report_data)
    for warning in validation_warnings:
        result.warnings.append(warning)
        print(f"[VALIDIERUNG] {warning}")

    stamm = report_data.get("stammdaten", {})
    # Unterstütze beide Varianten: förderziele (mit Umlaut) und foerderziele (ohne)
    ziele = report_data.get("förderziele", []) or report_data.get("foerderziele", [])
    bereiche = report_data.get("bereiche", {})
    umwelt = report_data.get("umweltfaktoren", {})
    neue = report_data.get("neue_ziele", [])
    empf = report_data.get("empfehlung", {})
    abschluss = report_data.get("abschluss", {})

    # --- STAMMDATEN via {{PLACEHOLDER}} ---
    # Empfehlung-Text generieren
    empf_text = "die Weiterbewilligung" if empf.get("verlängerung", True) else "die Beendigung"
    weiterbew_text = "Weiterbewilligung" if empf.get("verlängerung", True) else "Beendigung"

    # Landkreis-Adresse ermitteln
    landkreis = stamm.get("landkreis", "")
    if landkreis.lower() in ("loerrach", "lörrach"):
        amt_adresse = "Landratsamt Lörrach\nJugend & Familie SD IV\nHebelstraße 11\n79650 Schopfheim"
    elif landkreis.lower() == "waldshut":
        amt_adresse = "Landratsamt Waldshut\nJugendamt\nKaiserstraße 110\n79761 Waldshut-Tiengen"
    else:
        amt_adresse = ""  # Unbekannt - manuell eintragen

    # Komplexer Amt-Adresse-Platzhalter aus Template
    # Template hat: {{Landratsamt Lörrach / Waldshut...}} mit beiden Adressen als Platzhalter
    amt_platzhalter_loerrach = (
        "{{Landratsamt Lörrach / Waldshut\n"
        "Jugend & Familie SD IV / Jugendamt\n"
        "Hebelstraße 11 / Kaiserstraße 110\n"
        "79650 Schopfheim / 79761 Waldshut-Tiengen}}"
    )

    placeholders = {
        "{{NAME}}": stamm.get("name", ""),
        "{{GEBURTSDATUM}}": stamm.get("geburtsdatum", ""),
        "{{SACHBEARBEITER}}": stamm.get("sachbearbeiter", ""),
        "{{KOSTENZUSAGE_ENDE}}": stamm.get("kostenzusage_ende", ""),
        "{{WEITERBEWILLIGUNG_AB}}": stamm.get("weiterbewilligung_ab", ""),
        "{{FOERDERUNGSBEGINN}}": stamm.get("foerderungsbeginn", ""),
        "{{DIAGNOSE_1}}": "",
        "{{THERAPEUT_EINZEL}}": stamm.get("therapeut_einzel", ""),  # User-specific: set default in profile
        "{{THERAPEUT_GRUPPE}}": stamm.get("therapeut_gruppe", ""),
        "{{KOPIE_AN}}": stamm.get("kopie_an", ""),
        "{{UMFANG}}": empf.get("umfang", ""),
        "{{SONSTIGE_EMPFEHLUNG}}": empf.get("sonstige_empfehlung", "") or "",
        # Neue Platzhalter
        "{{Weiterbewilligung oder Beendigung}}": weiterbew_text,
        "{{Empfehlung}}": empf_text,
        "{{Zeichen Amtschreiben}}": stamm.get("zeichen_amt", ""),
        "{{AKTUELLE_ENTWICKLUNGEN}}": abschluss.get("aktuelle_entwicklungen", ""),
        "{{BEDINGUNGSMODELL}}": abschluss.get("bedingungsmodell", ""),
        "{{BESONDERE_FAEHIGKEITEN}}": report_data.get("besondere_faehigkeiten", ""),
        # Amt-Adresse basierend auf Landkreis (einfacher Platzhalter)
        "{{AMT_ADRESSE}}": amt_adresse,
        "{{LANDKREIS}}": landkreis,
        # Komplexer Amt-Adresse-Platzhalter (für Templates mit beiden Adressen)
        amt_platzhalter_loerrach: amt_adresse,
    }

    # Diagnosen zusammenbauen (neues Format mit Details)
    diagnosen = stamm.get("diagnosen", [])
    if diagnosen:
        diagnose_lines = []
        for diag in diagnosen:
            if isinstance(diag, dict):
                # Neues Format: {icd_code, bezeichnung, diagnostiker, datum, quelle}
                line = f"{diag.get('icd_code', '')} {diag.get('bezeichnung', '')}"
                extras = []
                if diag.get('diagnostiker'):
                    extras.append(f"Diagnose {diag['diagnostiker']}")
                if diag.get('datum'):
                    extras.append(diag['datum'])
                if extras:
                    line += f" ({', '.join(extras)})"
                diagnose_lines.append(line)
            else:
                # Altes Format: einfacher String
                diagnose_lines.append(str(diag))
        placeholders["{{DIAGNOSE_1}}"] = "\n".join(diagnose_lines)

    svc.replace_placeholders(doc, placeholders)

    # --- KOMPLEXER LANDRATSAMT-PLATZHALTER (ueber mehrere Runs verteilt) ---
    # Der Platzhalter im Template ist: {{Landratsamt Lörrach / Waldshut...}}
    # mit Zeilenumbruechen dazwischen - muss direkt im XML ersetzt werden
    landratsamt_search = "{{Landratsamt"  # Anfang des komplexen Platzhalters
    if _replace_complex_placeholder_in_xml(doc, landratsamt_search, amt_adresse):
        # Entferne auch das abschliessende }}
        _replace_complex_placeholder_in_xml(doc, "79761 Waldshut-Tiengen}}", "")

    # --- CHECKBOXEN ---
    therapie = stamm.get("therapieangebot", {})
    svc.set_checkbox(doc, "Einzelförderung", therapie.get("einzelfoerderung", True))
    svc.set_checkbox(doc, "Gruppenförderung", therapie.get("gruppenfoerderung", False))
    svc.set_checkbox(doc, "Co-Förderung", therapie.get("co_foerderung", False))
    svc.set_checkbox(doc, "Förderdiagnostik", therapie.get("foerderdiagnostik", False))

    netzwerk = stamm.get("netzwerkarbeit", {})
    svc.set_checkbox(doc, "Elternarbeit", netzwerk.get("elternarbeit", True))
    svc.set_checkbox(doc, "Kindergarten", netzwerk.get("kindergarten", False))
    svc.set_checkbox(doc, "Schule", netzwerk.get("schule", False))
    svc.set_checkbox(doc, "Sonstige Stellen", netzwerk.get("sonstige", False))

    # --- EMPFEHLUNGS-CHECKBOXEN ---
    verlaengerung = empf.get("verlängerung", True)
    begruendungen = empf.get("begründung", [])

    if verlaengerung:
        # Verlängerungs-Gründe
        svc.set_checkbox(doc, "Therapieziele noch nicht erreicht", "therapieziele_nicht_erreicht" in begruendungen)
        svc.set_checkbox(doc, "Übergang in den nächsten Lebensbereich", "übergang_lebensbereich" in begruendungen)
        svc.set_checkbox(doc, "Krise", "krise_erhoehter_bedarf" in begruendungen)
        svc.set_checkbox(doc, "erhöhter Bedarf", "krise_erhoehter_bedarf" in begruendungen)
    else:
        # Beendigungs-Gründe
        svc.set_checkbox(doc, "Ziele erreicht", "ziele_erreicht" in begruendungen)
        svc.set_checkbox(doc, "andere Therapie vorrangig", "andere_therapie_vorrangig" in begruendungen)
        svc.set_checkbox(doc, "Pausierung", "pausierung" in begruendungen)

    # --- ICF-KERNTABELLE MIT PLATZHALTERN (Tabelle 3) ---
    # Neue Methode: Zeilen mit Platzhaltern {CODE-Ziel}, {CODE-Ist}, {CODE-E}, {CODE-G}
    # werden gefuellt, Zeilen ohne Daten geloescht
    if len(doc.tables) >= 3:
        icf_table = doc.tables[2]  # 0-basiert, Tabelle 3

        if ziele:
            # Versuche zuerst die neue Platzhalter-Methode
            replacements, deleted = svc.fill_icf_placeholders_and_cleanup(
                icf_table, ziele, header_rows=1
            )

            # Fallback: Wenn keine Platzhalter gefunden wurden, alte Methode verwenden
            if replacements == 0:
                aktive_codes = [z["icf_code"].upper() for z in ziele]
                _filter_icf_table(icf_table, aktive_codes)

                # Ist-Stand und Bewertung einfuegen
                for ziel in ziele:
                    row_idx = svc.find_table_row_by_text(icf_table, ziel["icf_code"], column=0)
                    if row_idx is not None:
                        # Zielformulierung (Spalte 2)
                        if ziel.get("zielformulierung"):
                            svc.set_cell_text(icf_table, row_idx, 2, ziel["zielformulierung"])
                        # Ist-Stand (Spalte 3)
                        if ziel.get("ist_stand"):
                            svc.set_cell_text(icf_table, row_idx, 3, ziel["ist_stand"])
                        # Erreicht (Spalte 4)
                        if ziel.get("erreicht"):
                            erreicht_text = {1: "(1)", 2: "(2)", 3: "(3)"}.get(ziel["erreicht"], "")
                            svc.set_cell_text(icf_table, row_idx, 4, erreicht_text)
                        # Grund (Spalte 5)
                        if ziel.get("grund_nichterreichung"):
                            grund_text = f"({ziel['grund_nichterreichung']})"
                            svc.set_cell_text(icf_table, row_idx, 5, grund_text)

    # --- SATZBAUSTEINE-BEREICHE (als {{PLACEHOLDER}}) ---
    # Klientenname für [Name]-Ersetzung in Vorlagen
    klient_name = stamm.get("name", "").split()[0] if stamm.get("name") else "der Klient"

    bereich_map = {
        "mobilitaet": "{{BEREICH_MOBILITAET}}",
        "selbstversorgung": "{{BEREICH_SELBSTVERSORGUNG}}",
        "haushalt": "{{BEREICH_HAUSHALT}}",
        "lebensbereiche": "{{BEREICH_LEBENSBEREICHE}}",
        "gesellschaft_freizeit": "{{BEREICH_GESELLSCHAFT_FREIZEIT}}"
    }

    # Sammle alle BEREICH-Texte, die eingefügt werden sollen
    bereich_texte = []
    for key, placeholder in bereich_map.items():
        bereich = bereiche.get(key, {})

        # Neues flaches Format: Codes direkt im Bereich ohne "aktiv" Flag
        # {"gesellschaft_freizeit": {"D910": "Text...", "D920": "Text...", "D9200": null}}
        # Wenn mindestens ein Code Text hat → Bereich ist aktiv
        if BEREICH_CODES_AVAILABLE and isinstance(bereich, dict):
            # Prüfe ob es Codes direkt im Bereich gibt (flaches Format)
            from bereich_codes import BEREICH_STRUKTUR, build_bereich_text_from_codes
            if key in BEREICH_STRUKTUR:
                valid_codes = set(BEREICH_STRUKTUR[key]["codes"].keys())
                # Filtere: nur gültige Codes mit nicht-leerem Text
                code_texte = {k: v for k, v in bereich.items()
                             if k in valid_codes and v and v.strip()}
                if code_texte:
                    text, used_codes = build_bereich_text_from_codes(key, code_texte, klient_name)
                    if text:
                        bereich_texte.append(text)
                    # Platzhalter entfernen und weiter
                    svc.remove_textblock(doc, placeholder, remove_paragraph=False)
                    continue

        # Altes Format mit "aktiv" Flag
        if bereich.get("aktiv"):
            if TEXTBAUSTEINE_AVAILABLE:
                text = get_bereich_text(key, klient_name, bereich.get("text"))
            else:
                text = bereich.get("text", "").replace("[Name]", klient_name)
            if text:
                bereich_texte.append(text)
        # Platzhalter entfernen (NICHT den ganzen Paragraph!)
        svc.remove_textblock(doc, placeholder, remove_paragraph=False)

    # Alle gesammelten BEREICH-Texte zusammen einfügen
    if bereich_texte:
        combined_text = "\n\n".join(bereich_texte)
        # Füge nach "Bemerkungen zum Förderverlauf" ein
        for i, para in enumerate(doc.paragraphs):
            if "Bemerkungen zum Förderverlauf" in para.text:
                # Neuen Paragraph nach dem Header einfügen
                new_para = doc.paragraphs[i + 1] if i + 1 < len(doc.paragraphs) else para
                if new_para.runs:
                    new_para.runs[0].text = combined_text
                else:
                    new_para.add_run(combined_text)
                break

    # --- UMWELTFAKTOREN (als {{PLACEHOLDER}}) ---
    umwelt_map = {
        "wahrnehmung": "{{UMWELT_WAHRNEHMUNG}}",
        "kommunikation": "{{UMWELT_KOMMUNIKATION}}",
        "kognition_motorik": "{{UMWELT_KOGNITION_MOTORIK}}",
        "medizinisches": "{{UMWELT_MEDIZINISCHES}}"
    }

    # Sammle alle UMWELT-Texte, die eingefügt werden sollen
    umwelt_texte = []
    for key, placeholder in umwelt_map.items():
        faktor = umwelt.get(key, {})

        # Neues flaches Format: Codes direkt im Faktor ohne "aktiv" Flag
        # {"wahrnehmung": {"E240": "Text...", "E250": "Text...", "B1562": null}}
        # Wenn mindestens ein Code Text hat → Faktor ist aktiv
        if BEREICH_CODES_AVAILABLE and isinstance(faktor, dict):
            from bereich_codes import UMWELT_STRUKTUR, build_umwelt_text_from_codes
            if key in UMWELT_STRUKTUR:
                valid_codes = set(UMWELT_STRUKTUR[key]["codes"].keys())
                # Filtere: nur gültige Codes mit nicht-leerem Text
                code_texte = {k: v for k, v in faktor.items()
                             if k in valid_codes and v and v.strip()}
                if code_texte:
                    text, used_codes = build_umwelt_text_from_codes(key, code_texte, klient_name)
                    if text:
                        umwelt_texte.append(text)
                    svc.remove_textblock(doc, placeholder, remove_paragraph=False)
                    continue

        # Altes Format mit "aktiv" Flag
        if faktor.get("aktiv"):
            if TEXTBAUSTEINE_AVAILABLE:
                text = get_umwelt_text(key, klient_name, faktor.get("text"))
            else:
                text = faktor.get("text", "").replace("[Name]", klient_name)
            if text:
                umwelt_texte.append(text)
        # Platzhalter entfernen (NICHT den ganzen Paragraph!)
        svc.remove_textblock(doc, placeholder, remove_paragraph=False)

    # Alle gesammelten UMWELT-Texte zusammen einfügen
    if umwelt_texte:
        combined_text = "\n\n".join(umwelt_texte)
        # Füge nach "Umweltfaktoren" ein
        for i, para in enumerate(doc.paragraphs):
            if "Umweltfaktoren" in para.text and "{{" not in para.text:
                # Neuen Paragraph nach dem Header einfügen
                new_para = doc.paragraphs[i + 1] if i + 1 < len(doc.paragraphs) else para
                if new_para.runs:
                    new_para.runs[0].text = combined_text
                else:
                    new_para.add_run(combined_text)
                break

    # --- BESONDERE FAEHIGKEITEN ---
    if report_data.get("besondere_faehigkeiten"):
        svc.activate_textblock(
            doc,
            "{{BESONDERE_FAEHIGKEITEN}}",
            report_data["besondere_faehigkeiten"]
        )
    else:
        svc.remove_textblock(doc, "{{BESONDERE_FAEHIGKEITEN}}", remove_paragraph=True)

    # --- NEUE ZIELE (Tabelle 4) ---
    if len(doc.tables) >= 4 and neue:
        neue_tabelle = doc.tables[3]  # 0-basiert
        for i, ziel in enumerate(neue):
            row_idx = i + 1  # +1 fuer Header
            if row_idx < len(neue_tabelle.rows):
                icf_code = ziel.get("icf_code", "")
                # Kapitel automatisch aus Code ableiten wenn nicht angegeben
                kapitel = ziel.get("kapitel") or get_icf_kapitel(icf_code)
                svc.set_cell_text(neue_tabelle, row_idx, 0, icf_code)
                svc.set_cell_text(neue_tabelle, row_idx, 1, kapitel)
                svc.set_cell_text(neue_tabelle, row_idx, 2, ziel.get("zielformulierung", ""))
                svc.set_cell_text(neue_tabelle, row_idx, 3, ziel.get("ist_stand", ""))

    # --- EMPFEHLUNG (Tabellen 5-8) ---
    if empf.get("verlaengerung"):
        begruendungen = empf.get("begruendung", [])
        svc.set_checkbox(doc, "Therapieziele noch nicht erreicht",
                         "therapieziele_nicht_erreicht" in begruendungen)
        svc.set_checkbox(doc, "Übergang in den nächsten Lebensbereich",
                         "uebergang_lebensbereich" in begruendungen)
        svc.set_checkbox(doc, "Krise/ erhöhter Bedarf",
                         "krise_erhoehter_bedarf" in begruendungen)

    if empf.get("beendigung"):
        svc.set_checkbox(doc, "Ziele erreicht",
                         "ziele_erreicht" in empf.get("beendigung_begruendung", []))
        svc.set_checkbox(doc, "andere Therapie vorrangig",
                         "andere_therapie" in empf.get("beendigung_begruendung", []))
        svc.set_checkbox(doc, "Pausierung",
                         "pausierung" in empf.get("beendigung_begruendung", []))

    # --- ABSCHLUSS ---
    if abschluss.get("aktuelle_entwicklungen"):
        svc.activate_textblock(
            doc,
            "{{AKTUELLE_ENTWICKLUNGEN}}",
            abschluss["aktuelle_entwicklungen"]
        )
    else:
        svc.remove_textblock(doc, "{{AKTUELLE_ENTWICKLUNGEN}}", remove_paragraph=True)

    if abschluss.get("bedingungsmodell"):
        svc.activate_textblock(
            doc,
            "{{BEDINGUNGSMODELL}}",
            abschluss["bedingungsmodell"]
        )
    else:
        svc.remove_textblock(doc, "{{BEDINGUNGSMODELL}}", remove_paragraph=True)

    # --- SPEICHERN ---
    try:
        saved = svc.save(doc, output_path)
        result.success = True
        result.output_path = str(saved)
        result.template_filled = True
    except Exception as e:
        result.errors.append(f"Speichern fehlgeschlagen: {e}")

    return result


def _filter_icf_table(table, aktive_codes: List[str]):
    """
    Filtert die ICF-Tabelle unter Beruecksichtigung von Sub-Items und Kapitel-Headern.

    Struktur der ICF-Tabelle:
    - Row mit Code in C0 = normaler ICF-Code
    - Row mit leerem C0 + vMerge=continue = Sub-Item (gehoert zum vorherigen Code)
    - Row mit leerem C0 + vMerge=restart = Parent einer Sub-Gruppe
    - Row mit leerem C0 + kein vMerge + Text in C1 = Kapitel-Header

    Kapitel-Header werden nur behalten, wenn mindestens ein nachfolgender
    Code im selben Kapitel aktiv ist.
    """
    keep_values_set = set(c.strip().upper() for c in aktive_codes)

    # Phase 1: Zeilen klassifizieren
    row_info = []  # (row, type, code) - type: 'header','code','subitem','section_header'
    for i, row in enumerate(table.rows):
        if i == 0:
            row_info.append((row, "header", ""))
            continue

        row_elem = row._element
        tc_elems = row_elem.findall(qn("w:tc"))
        if not tc_elems:
            row_info.append((row, "unknown", ""))
            continue

        tc0 = tc_elems[0]
        tc_pr = tc0.find(qn("w:tcPr"))
        vm = tc_pr.find(qn("w:vMerge")) if tc_pr is not None else None
        is_vmerge_continue = vm is not None and vm.get(qn("w:val")) is None
        is_vmerge_restart = vm is not None and vm.get(qn("w:val")) == "restart"

        cell_texts = [t.text or "" for t in tc0.iter(qn("w:t"))]
        cell_text = "".join(cell_texts).strip().upper()

        if is_vmerge_continue:
            row_info.append((row, "subitem", ""))
        elif cell_text:
            row_info.append((row, "code", cell_text))
        else:
            # Leere C0: entweder section_header oder vMerge-restart parent
            row_info.append((row, "section_header", ""))

    # Phase 2: Entscheiden was behalten wird
    rows_to_remove = []
    current_code_active = False

    for idx in range(len(row_info)):
        row, rtype, code = row_info[idx]

        if rtype == "header":
            continue
        elif rtype == "code":
            current_code_active = code in keep_values_set
            if not current_code_active:
                rows_to_remove.append(row)
        elif rtype == "subitem":
            if not current_code_active:
                rows_to_remove.append(row)
        elif rtype == "section_header":
            # Behalten nur wenn ein nachfolgender Code (vor dem naechsten
            # section_header) aktiv ist
            has_active_follower = False
            for j in range(idx + 1, len(row_info)):
                _, jtype, jcode = row_info[j]
                if jtype == "section_header":
                    break  # Naechster Abschnitt, stopp
                if jtype == "code" and jcode in keep_values_set:
                    has_active_follower = True
                    break
            if not has_active_follower:
                rows_to_remove.append(row)

    # Phase 3: Von hinten loeschen (Index-Stabilitaet)
    for row in reversed(rows_to_remove):
        table._tbl.remove(row._tr)


# ═══════════════════════════════════════════════════════════════
# Haupt-Workflow
# ═══════════════════════════════════════════════════════════════

class ReportGenerator:
    """
    Haupt-Klasse fuer die Berichtsgenerierung.

    Verwendung:
        gen = ReportGenerator()

        # Vollautomatisch (mit LLM):
        result = gen.generate(
            template_path="vorlage.docx",
            source_folder="klienten/K_0042/",
            output_path="klienten/K_0042/output/bericht.docx"
        )

        # Manuell (mit vorbereitetem JSON):
        result = gen.generate_from_json(
            template_path="vorlage.docx",
            json_path="bericht_data.json",
            output_path="bericht.docx"
        )
    """

    def __init__(self):
        self.default_template = str(TEMPLATES_DIR / "bericht_template_geiger_universal.docx")
        self.default_schema = "foerderbericht"

    def extract_sources(self, source_folder: str) -> str:
        """Phase 1: Quelldokumente extrahieren."""
        return extract_all_sources(source_folder)

    def build_prompt(self, source_text: str, schema_name: str = None) -> str:
        """Phase 2: LLM-Prompt zusammenbauen."""
        schema = load_schema(schema_name or self.default_schema)
        return build_llm_prompt(source_text, schema)

    def fill_report(
        self,
        report_data: dict,
        template_path: str = None,
        output_path: str = None
    ) -> GeneratorResult:
        """Phase 3: Word-Template befuellen."""
        return fill_template(
            template_path or self.default_template,
            report_data,
            output_path or "bericht_output.docx"
        )

    def generate_from_json(
        self,
        json_path: str,
        template_path: str = None,
        output_path: str = None
    ) -> GeneratorResult:
        """Generiert Bericht aus vorbereiteter JSON-Datei."""
        path = Path(json_path)
        if not path.exists():
            result = GeneratorResult()
            result.errors.append(f"JSON nicht gefunden: {json_path}")
            return result

        data = json.loads(path.read_text(encoding="utf-8"))
        result = self.fill_report(data, template_path, output_path)
        result.json_generated = True
        return result

    def generate(
        self,
        template_path: str = None,
        source_folder: str = None,
        output_path: str = None,
        schema_name: str = None
    ) -> GeneratorResult:
        """
        Vollstaendiger Generierungsprozess.

        HINWEIS: Die LLM-Generierung (Phase 2) geschieht AUSSERHALB dieses Skripts.
        Dieser Befehl:
        1. Extrahiert die Quelltexte
        2. Gibt den fertigen Prompt aus (oder speichert ihn)
        3. Erwartet den JSON-Output als Eingabe

        Fuer die automatische Variante muss der aufrufende Workflow
        den Prompt an ein LLM senden und den JSON-Output zurueckfuehren.
        """
        result = GeneratorResult()

        # Phase 1: Extraktion
        if source_folder:
            source_text = self.extract_sources(source_folder)
            result.source_text_length = len(source_text)
        else:
            result.errors.append("source_folder nicht angegeben")
            return result

        # Phase 2: Prompt erstellen
        prompt = self.build_prompt(source_text, schema_name)

        # Prompt speichern fuer manuelle/externe LLM-Nutzung
        if output_path:
            prompt_path = Path(output_path).parent / "_prompt.txt"
            prompt_path.write_text(prompt, encoding="utf-8")
            result.warnings.append(f"Prompt gespeichert: {prompt_path}")

        # Phase 3 muss extern getriggert werden mit dem JSON-Output
        result.warnings.append(
            "LLM-Generierung muss extern erfolgen. "
            "Nutze generate_from_json() mit dem LLM-Output."
        )

        return result

    # ═══════════════════════════════════════════════════════════════
    # ANONYMISIERTER WORKFLOW
    # ═══════════════════════════════════════════════════════════════

    def extract_sources_anonymized(
        self,
        klient_name: str,
        password: str,
        force_new_bundle: bool = False
    ) -> Tuple[str, "BundleInfo"]:
        """
        Extrahiert Quelltexte aus einem anonymisierten Bundle.

        Das LLM sieht NUR Tarnnamen, niemals echte Namen.

        Args:
            klient_name: Name des Klienten-Ordners
            password: Passwort für Schlüssel-Erstellung
            force_new_bundle: Bundle komplett neu erstellen

        Returns:
            (anonymisierter_text, bundle_info)
        """
        if not ANONYMIZED_WORKFLOW_AVAILABLE:
            raise RuntimeError("Anonymized Workflow nicht verfügbar")

        workflow = AnonymizedWorkflow()

        # Bundle erstellen/laden
        bundle_info = workflow.create_or_update_bundle(
            klient_ordner=klient_name,
            password=password,
            force_recreate=force_new_bundle
        )

        # Anonymisierten Text extrahieren
        source_text = workflow.get_bundle_content(bundle_info.client_id)

        return source_text, bundle_info

    def generate_anonymized(
        self,
        klient_name: str,
        password: str,
        template_path: str = None,
        schema_name: str = None
    ) -> Tuple[GeneratorResult, "BundleInfo"]:
        """
        Generiert einen Prompt aus anonymisierten Quelldaten.

        WICHTIG: Das LLM sieht nur Tarnnamen.
        Nach der JSON-Generierung muss save_deanonymized() aufgerufen werden.

        Args:
            klient_name: Name des Klienten-Ordners
            password: Passwort für Schlüssel
            template_path: Word-Vorlage
            schema_name: JSON-Schema Name

        Returns:
            (GeneratorResult mit Prompt, BundleInfo für spätere De-Anonymisierung)
        """
        result = GeneratorResult()

        try:
            # Anonymisierte Quellen extrahieren
            source_text, bundle_info = self.extract_sources_anonymized(
                klient_name=klient_name,
                password=password
            )
            result.source_text_length = len(source_text)

            # Prompt erstellen
            prompt = self.build_prompt(source_text, schema_name)

            # Prompt in Bundle-Output speichern
            prompt_path = bundle_info.bundle_path / "output" / "_prompt.txt"
            prompt_path.write_text(prompt, encoding="utf-8")

            result.warnings.append(f"Client-ID: {bundle_info.client_id}")
            result.warnings.append(f"Tarnname: {bundle_info.tarnname}")
            result.warnings.append(f"Prompt gespeichert: {prompt_path}")
            result.warnings.append(
                "LLM-Generierung muss extern erfolgen. "
                "Danach: generate_from_json_anonymized() aufrufen."
            )

            return result, bundle_info

        except Exception as e:
            result.errors.append(str(e))
            return result, None

    def generate_from_json_anonymized(
        self,
        json_path: str,
        client_id: str,
        password: str,
        template_path: str = None
    ) -> GeneratorResult:
        """
        Generiert Bericht aus JSON und de-anonymisiert automatisch.

        Args:
            json_path: Pfad zur JSON-Datei (mit Tarnnamen)
            client_id: Client-ID für De-Anonymisierung
            password: Passwort für Schlüssel
            template_path: Word-Vorlage

        Returns:
            GeneratorResult mit Pfad zum de-anonymisierten Bericht
        """
        if not ANONYMIZED_WORKFLOW_AVAILABLE:
            result = GeneratorResult()
            result.errors.append("Anonymized Workflow nicht verfügbar")
            return result

        workflow = AnonymizedWorkflow()

        # Temporären Output-Pfad im Bundle
        temp_output = workflow.bundles_dir / client_id / "output" / "bericht_temp.docx"

        # JSON laden und Bericht erstellen (mit Tarnnamen)
        result = self.generate_from_json(
            json_path=json_path,
            template_path=template_path,
            output_path=str(temp_output)
        )

        if not result.success:
            return result

        # De-anonymisieren
        deanon_result = workflow.save_deanonymized(
            source_path=str(temp_output),
            client_id=client_id,
            password=password
        )

        if deanon_result.success:
            result.output_path = deanon_result.output_path
            result.warnings.append(f"De-anonymisiert: {deanon_result.message}")
        else:
            result.errors.extend(deanon_result.errors)
            result.success = False

        return result


# ═══════════════════════════════════════════════════════════════
# Export-Pipeline
# ═══════════════════════════════════════════════════════════════

# Anonymizer-Service importieren
try:
    from skills._services.document.anonymizer_service import (
        DocumentDeanonymizer, decrypt_key_file, get_key_path
    )
    ANONYMIZER_AVAILABLE = True
except ImportError:
    ANONYMIZER_AVAILABLE = False


@dataclass
class ExportResult:
    """Ergebnis eines Export-Vorgangs."""
    success: bool = False
    output_folder: str = ""
    real_name: str = ""
    files_processed: int = 0
    replacements: int = 0
    errors: List[str] = field(default_factory=list)


class ExportPipeline:
    """
    Export-Pipeline fuer Foerderberichte.

    Ablauf:
        1. prepare(): Kopiert Output-Dateien nach _prepare_for_export/
        2. deanonymize(): Ersetzt Tarnnamen durch echte Namen
        3. finalize(): Verschiebt nach _ready_for_export/<EchterName>/
        4. archive() (optional): Verschiebt nach _archive/

    Kann auch als Einschritt-Pipeline genutzt werden:
        result = pipeline.export(client_folder, password)
    """

    def __init__(self, base_dir: str = None):
        if base_dir:
            self.base_dir = Path(base_dir)
        else:
            self.base_dir = BACH_DIR / "user" / "foerderplaner"

        self.klienten_dir = self.base_dir / "klienten"
        self.prepare_dir = self.base_dir / "_prepare_for_export"
        self.ready_dir = self.base_dir / "_ready_for_export"
        self.archive_dir = self.base_dir / "_archive"

    def _load_profil(self, client_folder: Path) -> Optional[dict]:
        """Laedt .profil.json aus dem Klienten-Ordner."""
        profil_path = client_folder / ".profil.json"
        if profil_path.exists():
            return json.loads(profil_path.read_text(encoding="utf-8"))
        return None

    def _get_real_name(self, client_id: str, password: str) -> Optional[str]:
        """Ermittelt den echten Namen aus dem Schluessel."""
        try:
            key_path = get_key_path(client_id)
            if not key_path.exists():
                return None
            profile = decrypt_key_file(str(key_path), password)
            names = profile.mappings.get("names", {})
            # Erster Eintrag = Hauptname (Original -> Tarnname)
            if names:
                return list(names.keys())[0]
        except Exception:
            pass
        return None

    def export(
        self,
        client_folder: str,
        password: str,
        output_subfolder: str = "output"
    ) -> ExportResult:
        """
        Vollstaendiger Export-Vorgang (Einschritt-Pipeline).

        Args:
            client_folder: Klienten-Ordner (z.B. "klienten/K_0042" oder absoluter Pfad)
            password: Passwort fuer den Anonymisierungs-Schluessel
            output_subfolder: Unterordner mit den Berichten (default: "output")

        Returns:
            ExportResult mit Ergebnis-Details
        """
        result = ExportResult()

        if not ANONYMIZER_AVAILABLE:
            result.errors.append("Anonymizer-Service nicht verfuegbar (cryptography installiert?)")
            return result

        # Klienten-Ordner aufloesen
        client_path = Path(client_folder)
        if not client_path.is_absolute():
            client_path = self.base_dir / client_folder

        if not client_path.exists():
            result.errors.append(f"Klienten-Ordner nicht gefunden: {client_path}")
            return result

        # Profil laden
        profil = self._load_profil(client_path)
        client_id = profil.get("client_id") if profil else client_path.name

        # Echten Namen ermitteln
        real_name = self._get_real_name(client_id, password)
        if not real_name:
            result.errors.append(
                f"Echter Name konnte nicht ermittelt werden fuer {client_id}. "
                f"Schluessel vorhanden? Passwort korrekt?"
            )
            return result

        result.real_name = real_name
        safe_name = real_name.replace(" ", "_")

        # Quellordner (nur output/ mit generierten Berichten)
        source_dir = client_path / output_subfolder
        if not source_dir.exists():
            result.errors.append(f"Output-Ordner nicht gefunden: {source_dir}")
            return result

        # Phase 1: Prepare — Dateien nach _prepare_for_export kopieren
        prepare_path = self.prepare_dir / client_id
        prepare_path.mkdir(parents=True, exist_ok=True)

        import shutil

        export_files = [
            f for f in source_dir.iterdir()
            if f.is_file()
            and not f.name.startswith("_")
            and not f.name.startswith(".")
            and f.suffix.lower() in (".docx", ".txt", ".md", ".pdf")
        ]

        if not export_files:
            result.errors.append(f"Keine exportierbaren Dateien in {source_dir}")
            return result

        for f in export_files:
            shutil.copy2(f, prepare_path / f.name)

        # Phase 2: De-Anonymisierung
        ready_path = self.ready_dir / safe_name
        deanon = DocumentDeanonymizer()
        deanon_result = deanon.deanonymize_folder(
            folder=str(prepare_path),
            schluessel_path=None,
            password=password,
            output_folder=str(ready_path),
            client_id=client_id
        )

        result.files_processed = deanon_result.processed_files
        result.replacements = deanon_result.replacements_total

        if deanon_result.errors:
            result.errors.extend(deanon_result.errors)

        # Phase 3: Aufraumen — _prepare_for_export leeren
        shutil.rmtree(prepare_path, ignore_errors=True)

        result.success = True
        result.output_folder = str(ready_path)

        return result

    def archive(self, client_folder: str, ready_name: str = None) -> bool:
        """
        Verschiebt exportierten Bericht in _archive/.

        Args:
            client_folder: Klienten-Ordner (fuer client_id)
            ready_name: Name des Ordners in _ready_for_export/ (oder auto-detect)
        """
        import shutil

        if ready_name:
            src = self.ready_dir / ready_name
        else:
            # Auto-detect: Suche in _ready_for_export/
            candidates = list(self.ready_dir.iterdir()) if self.ready_dir.exists() else []
            if len(candidates) == 1:
                src = candidates[0]
            else:
                return False

        if not src.exists():
            return False

        timestamp = datetime.now().strftime("%Y-%m")
        dest = self.archive_dir / f"{src.name}_{timestamp}"
        dest.mkdir(parents=True, exist_ok=True)

        for f in src.iterdir():
            if f.is_file():
                shutil.move(str(f), str(dest / f.name))

        # Leeren Ordner entfernen
        shutil.rmtree(src, ignore_errors=True)
        return True


# ═══════════════════════════════════════════════════════════════
# CLI Interface
# ═══════════════════════════════════════════════════════════════

def main():
    """CLI fuer den Report Generator."""
    import argparse

    parser = argparse.ArgumentParser(description="Report Generator CLI")
    subparsers = parser.add_subparsers(dest="command")

    # extract: Quelltexte extrahieren
    extract_p = subparsers.add_parser("extract", help="Quelltexte extrahieren")
    extract_p.add_argument("folder", help="Quellordner")
    extract_p.add_argument("--output", "-o", help="Output-Datei (optional)")

    # prompt: LLM-Prompt generieren
    prompt_p = subparsers.add_parser("prompt", help="LLM-Prompt erstellen")
    prompt_p.add_argument("folder", help="Quellordner")
    prompt_p.add_argument("--output", "-o", help="Prompt-Ausgabedatei")
    prompt_p.add_argument("--schema", "-s", default="foerderbericht", help="Schema-Name")

    # fill: Vorlage mit JSON fuellen
    fill_p = subparsers.add_parser("fill", help="Word-Vorlage mit JSON fuellen")
    fill_p.add_argument("json_file", help="JSON-Datei mit Berichtsdaten")
    fill_p.add_argument("--template", "-t", help="Word-Vorlage")
    fill_p.add_argument("--output", "-o", required=True, help="Ausgabe-Datei")

    # export: Bericht exportieren (de-anonymisieren)
    export_p = subparsers.add_parser("export", help="Bericht exportieren (de-anonymisieren)")
    export_p.add_argument("client_folder", help="Klienten-Ordner (z.B. klienten/K_0042)")
    export_p.add_argument("--password", "-p", required=True, help="Schluessel-Passwort")
    export_p.add_argument("--subfolder", default="output", help="Output-Unterordner (default: output)")

    # archive: Exportierten Bericht archivieren
    archive_p = subparsers.add_parser("archive", help="Exportierten Bericht archivieren")
    archive_p.add_argument("name", nargs="?", help="Name in _ready_for_export/ (oder auto)")

    # info: Template-Info anzeigen
    info_p = subparsers.add_parser("info", help="Template-Info anzeigen")
    info_p.add_argument("--template", "-t", help="Word-Vorlage")

    args = parser.parse_args()

    gen = ReportGenerator()

    if args.command == "extract":
        text = gen.extract_sources(args.folder)
        if args.output:
            Path(args.output).write_text(text, encoding="utf-8")
            print(f"Extrahiert: {len(text)} Zeichen -> {args.output}")
        else:
            print(text)

    elif args.command == "prompt":
        text = gen.extract_sources(args.folder)
        prompt = gen.build_prompt(text, args.schema)
        if args.output:
            Path(args.output).write_text(prompt, encoding="utf-8")
            print(f"Prompt: {len(prompt)} Zeichen -> {args.output}")
        else:
            print(prompt)

    elif args.command == "fill":
        result = gen.generate_from_json(
            args.json_file,
            template_path=args.template,
            output_path=args.output
        )
        if result.success:
            print(f"[OK] Bericht erstellt: {result.output_path}")
        else:
            print(f"[FEHLER] {result.errors}")

    elif args.command == "export":
        pipeline = ExportPipeline()
        print(f"Exportiere: {args.client_folder}")
        result = pipeline.export(
            client_folder=args.client_folder,
            password=args.password,
            output_subfolder=args.subfolder
        )
        if result.success:
            print(f"[OK] Export erfolgreich: {result.output_folder}")
            print(f"     Name: {result.real_name}")
            print(f"     Dateien: {result.files_processed}, Ersetzungen: {result.replacements}")
        else:
            print(f"[FEHLER] {result.errors}")

    elif args.command == "archive":
        pipeline = ExportPipeline()
        success = pipeline.archive(
            client_folder="",
            ready_name=args.name if hasattr(args, "name") else None
        )
        if success:
            print("[OK] Archivierung erfolgreich")
        else:
            print("[FEHLER] Archivierung fehlgeschlagen (Ordner nicht gefunden?)")

    elif args.command == "info":
        template = args.template or gen.default_template
        print(f"Template: {template}")
        print(f"Schema-Dir: {SCHEMAS_DIR}")
        print(f"Schemas: {[f.stem for f in SCHEMAS_DIR.glob('*.json')]}")
        if Path(template).exists():
            if WORD_SERVICE_AVAILABLE:
                svc = WordTemplateService()
                doc = svc.load_template(template)
                print(f"Tabellen: {len(doc.tables)}")
                print(f"Paragraphen: {len(doc.paragraphs)}")

    else:
        parser.print_help()


if __name__ == "__main__":
    main()
